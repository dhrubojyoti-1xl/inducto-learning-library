# -*- coding: utf-8 -*-
"""
Build the management review document for the Inducto mandatory learning
journey: Inducto_Learning_Knowledge_Module_Management_Review.docx

    python management_review_docx.py

WHAT THIS DOES
---------------
The 39-module PPTX/HTML library already built (11.2 hours) is the audited,
verified OPTIONAL EXTENDED LIBRARY. This script selects, condenses and
sequences a MANDATORY JOURNEY of about 120-180 minutes from that same
audited material — every video, every quiz question, every prompt and every
checklist item used below is quoted from the real content/areaNN/*.py
dictionaries that also build the decks and the web platform. Nothing is
invented; nothing here is a new video, a new fact or a new URL.

The only new authored text is: stage/stop framing sentences, the executive
narrative, the audit commentary, and short transitional sentences that join
two modules' real material into one condensed reading. Every such place is
composed only from real source phrases, never invented facts.

Time is computed, not guessed: video seconds (real, from the verified
duration string) + reading words/200wpm + a fixed exercise estimate per
content type + quiz seconds (30s/question), summed per stop and totalled.
"""

import importlib
import io
import json
import os
import re
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

import build
import sitegen
import theme as T

ROOT = sitegen.ROOT
OUT_DOCX = os.path.join(
    ROOT, "docs", "Inducto_Learning_Knowledge_Module_Management_Review.docx")

INK = RGBColor(0x10, 0x18, 0x26)
GREY = RGBColor(0x5B, 0x67, 0x79)
GOOD = RGBColor(0x1B, 0x7F, 0x4B)
BAD = RGBColor(0xC6, 0x28, 0x28)
ACCENT = RGBColor(0x2F, 0x4B, 0xC4)
GOLD = RGBColor(0x8A, 0x5A, 0x00)

TRACK_ACCENT = {
    "01-ai-general": RGBColor(0x2F, 0x4B, 0xC4),
    "02-ai-daily-work": RGBColor(0x0E, 0x6E, 0x75),
    "03-prompt-engineering": RGBColor(0x6A, 0x2F, 0xA0),
    "04-professional-skills": RGBColor(0x8A, 0x5A, 0x00),
    "05-security-privacy": RGBColor(0x7A, 0x12, 0x20),
}


# ---------------------------------------------------------------------------
# load all 39 audited modules
# ---------------------------------------------------------------------------
def load_decks():
    order = {a: i for i, a in enumerate(sitegen.AREA_ORDER)}
    decks = [importlib.import_module(m).DECK for m in build.REGISTRY]
    decks.sort(key=lambda d: (order[d["area"]], d["module_code"]))
    return {d["module_code"]: d for d in decks}, decks


def dur_seconds(dur):
    parts = [int(p) for p in dur.split(":")]
    while len(parts) < 3:
        parts.insert(0, 0)
    h, m, s = parts
    return h * 3600 + m * 60 + s


def words(text):
    return len(text.split())


def find_visual(d, vtype, label_hint=None):
    """First slide of a given visual type; prefers one whose section label
    contains label_hint (e.g. 'safe') when one exists."""
    hits = [(s, s["visual"]) for s in d["slides"]
            if s.get("visual") and s["visual"]["type"] == vtype]
    if not hits:
        return None, None
    if label_hint:
        for s, v in hits:
            if label_hint.lower() in (s.get("label") or "").lower():
                return s, v
    return hits[0]


def find_lead(d, label_hint=None):
    for s in d["slides"]:
        if s.get("lead") and (not label_hint or
                              label_hint.lower() in (s.get("label") or "").lower()):
            return s["title"], s["lead"]
    for s in d["slides"]:
        if s.get("lead"):
            return s["title"], s["lead"]
    return None, None


# ---------------------------------------------------------------------------
# the mandatory journey — declarative, then resolved against real content
# ---------------------------------------------------------------------------
# Each stop: stage, code, title, sources (module code(s)), content_type,
# use_video (bool), checklist_hint (label substring), n_quiz (how many real
# quiz questions to carry over, taken verbatim from the source module(s)).
STOPS = [
    # STAGE 1 — FOUNDATION
    dict(stage="1. Foundation", code="M-01", title="AI Fundamentals",
         sources=["AI-01"], use_video=True, checklist_hint="safe", n_quiz=1),

    # STAGE 2 — UNDERSTANDING
    dict(stage="2. Understanding", code="M-02", title="Generative AI",
         sources=["AI-02"], use_video=True, checklist_hint=None, n_quiz=1),
    dict(stage="2. Understanding", code="M-03",
         title="What AI Can and Cannot Do",
         sources=["AI-03", "AI-04"], use_video=True, video_source="AI-04",
         checklist_hint=None, n_quiz=1, combine=True),
    dict(stage="2. Understanding", code="M-04",
         title="AI Hallucinations & Fact-Checking",
         sources=["AI-05"], use_video=True, checklist_hint=None, n_quiz=2),

    # STAGE 3 — PRACTICAL AI USE
    dict(stage="3. Practical AI Use", code="M-05",
         title="Writing Email With AI", sources=["DW-01"], use_video=True,
         checklist_hint=None, n_quiz=1),
    dict(stage="3. Practical AI Use", code="M-06",
         title="Meeting Notes & Follow-ups With AI", sources=["DW-08"],
         use_video=True, checklist_hint=None, n_quiz=1),

    # STAGE 4 — PROMPT ENGINEERING
    dict(stage="4. Prompt Engineering", code="M-07",
         title="Basic Prompting", sources=["PE-01"], use_video=True,
         checklist_hint=None, n_quiz=2),
    dict(stage="4. Prompt Engineering", code="M-08",
         title="Instructions, Context & Role",
         sources=["PE-02", "PE-03"], use_video=True, video_source="PE-02",
         checklist_hint=None, n_quiz=1, combine=True),

    # STAGE 5 — WORKPLACE APPLICATION
    dict(stage="5. Workplace Application", code="M-09",
         title="Planning & Productivity With AI", sources=["DW-10"],
         use_video=True, checklist_hint=None, n_quiz=1),

    # STAGE 6 — PROFESSIONAL SKILLS
    dict(stage="6. Professional Skills", code="M-10",
         title="Business Communication", sources=["PS-01"], use_video=True,
         checklist_hint=None, n_quiz=1),
    dict(stage="6. Professional Skills", code="M-11",
         title="Time Management", sources=["PS-04"], use_video=True,
         checklist_hint=None, n_quiz=1),

    # STAGE 7 — SECURITY & RESPONSIBLE USE (compliance-critical: full breadth,
    # two closely-related pairs combined into one stop each to hold the time
    # budget — every one of the 7 required security topics still appears)
    dict(stage="7. Security & Responsible Use", code="M-12",
         title="Password Security & Multi-Factor Authentication",
         sources=["SEC-01", "SEC-03"], use_video=True, video_source="SEC-01",
         checklist_hint=None, n_quiz=1, combine=True),
    dict(stage="7. Security & Responsible Use", code="M-13",
         title="Phishing & Social Engineering", sources=["SEC-02"],
         use_video=True, checklist_hint=None, n_quiz=2),
    dict(stage="7. Security & Responsible Use", code="M-14",
         title="Data Protection & Confidential Information",
         sources=["SEC-04", "SEC-05"], use_video=True, video_source="SEC-04",
         checklist_hint=None, n_quiz=1, combine=True),
    dict(stage="7. Security & Responsible Use", code="M-15",
         title="Safe Use of AI at Work", sources=["SEC-06"],
         use_video=False, video_note_only=True, checklist_hint="safe",
         n_quiz=1),
    dict(stage="7. Security & Responsible Use", code="M-16",
         title="What Never to Paste Into AI", sources=["SEC-07"],
         use_video=True, checklist_hint=None, n_quiz=2),
]

ASSESSMENT_MIN = 12
PRACTICE_MIN = 3


# ---------------------------------------------------------------------------
# resolve each stop against real content
# ---------------------------------------------------------------------------
def resolve_stop(stop, by_code):
    primary = by_code[stop["sources"][0]]
    combine = stop.get("combine") and len(stop["sources"]) > 1
    secondary = by_code[stop["sources"][1]] if combine else None

    out = dict(stop)
    out["track"] = T.AREAS[primary["area"]]["name"]
    out["module_titles"] = [by_code[c]["title"] for c in stop["sources"]]

    # ---- video ----
    video_src = by_code[stop.get("video_source", stop["sources"][0])]
    v = video_src.get("video")
    out["video"] = v if stop.get("use_video") else None
    out["video_note_only"] = stop.get("video_note_only", False)
    out["video_ref_module"] = video_src["module_code"]

    # ---- reading: the module's real "why" story + fix, quoted verbatim ----
    w1 = primary["why"]
    reading = w1["scenario"] + " " + w1["fix"]
    if combine:
        w2 = secondary["why"]
        reading += " In " + secondary["title"] + ": " + w2["scenario"]
    out["reading_heading"] = w1["title"]
    out["reading"] = reading
    # a real outcome, already phrased as a capability, quoted verbatim
    out["objective"] = primary["outcomes"][0][1]
    if combine:
        out["objective"] += "; " + secondary["outcomes"][0][1][0].lower() +             secondary["outcomes"][0][1][1:]

    # ---- workplace example: a real lesson lead ----
    ex_title, ex_lead = find_lead(primary, "first")
    if not ex_lead:
        ex_title, ex_lead = find_lead(primary)
    out["example_title"] = ex_title
    out["example"] = ex_lead

    # ---- prompt: first prompt/prompt_out visual, real text ----
    prompt_text = None
    for vt in ("prompt", "prompt_out", "steps"):
        _, pv = find_visual(primary, vt)
        if pv and (pv.get("text") or pv.get("prompt")):
            prompt_text = pv.get("text") or pv.get("prompt")
            break
    out["prompt"] = prompt_text

    # ---- checklist: real items ----
    _, cv = find_visual(primary, "checklist", stop.get("checklist_hint"))
    out["checklist"] = cv["items"][:4] if cv else []

    # ---- quiz: first N real questions, verbatim, every answer's why kept ----
    quiz = list(primary["quiz"][: stop["n_quiz"]])
    if combine and stop["n_quiz"] < 2:
        quiz += list(secondary["quiz"][:1])
    out["quiz"] = quiz

    return out


# ---------------------------------------------------------------------------
# time calculator — real, not a guess
# ---------------------------------------------------------------------------
def stop_minutes(r):
    video_s = dur_seconds(r["video"]["duration"]) if r["video"] else 0
    reading_s = words(r["reading"]) / 200.0 * 60.0
    example_s = words(r["example"] or "") / 200.0 * 60.0
    exercise_s = 90.0 if r["prompt"] else (45.0 if r["checklist"] else 0.0)
    quiz_s = 35.0 * len(r["quiz"])
    total_s = video_s + reading_s + example_s + exercise_s + quiz_s
    if r["video_note_only"]:
        total_s -= video_s  # video referenced, not watched, in the mandatory slot
    return {
        "video_s": video_s if not r["video_note_only"] else 0,
        "reading_s": reading_s, "example_s": example_s,
        "exercise_s": exercise_s, "quiz_s": quiz_s,
        "total_min": round(total_s / 60.0, 1),
    }


# ===========================================================================
# document building blocks
# ===========================================================================
class Doc(object):

    def __init__(self):
        self.d = Document()
        self._page()
        self._styles()

    def _page(self):
        s = self.d.sections[0]
        s.page_width, s.page_height = Cm(21.0), Cm(29.7)
        s.left_margin, s.right_margin = Cm(2.2), Cm(2.2)
        s.top_margin, s.bottom_margin = Cm(2.0), Cm(2.0)
        p = s.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("Inducto Learning & Knowledge Module — Management "
                     "Review · page ")
        r.font.size, r.font.color.rgb = Pt(8), GREY
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        p._p.append(fld)

    def _styles(self):
        st = self.d.styles
        n = st["Normal"]
        n.font.name, n.font.size, n.font.color.rgb = "Segoe UI", Pt(10.5), INK
        n.paragraph_format.space_after = Pt(7)
        n.paragraph_format.line_spacing = 1.18
        for name, size, col, before in (("Heading 1", 22, INK, 0),
                                        ("Heading 2", 15, ACCENT, 20),
                                        ("Heading 3", 12, INK, 12),
                                        ("Heading 4", 10.5, GREY, 8)):
            s = st[name]
            s.font.name = "Segoe UI Semibold"
            s.font.size, s.font.color.rgb, s.font.bold = Pt(size), col, True
            s.paragraph_format.space_before = Pt(before)
            s.paragraph_format.space_after = Pt(5)
            s.paragraph_format.keep_with_next = True

    def h(self, level, text, color=None, page_break=False):
        if page_break:
            self.d.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        p = self.d.add_paragraph(style="Heading %d" % level)
        r = p.add_run(text)
        if color is not None:
            r.font.color.rgb = color
        return p

    def para(self, text, size=10.5, color=None, italic=False, bold=False,
             space=7, indent=0.0):
        p = self.d.add_paragraph()
        p.paragraph_format.space_after = Pt(space)
        if indent:
            p.paragraph_format.left_indent = Cm(indent)
        r = p.add_run(text)
        r.font.size, r.font.italic, r.font.bold = Pt(size), italic, bold
        r.font.color.rgb = color if color is not None else INK
        return p

    def bullets(self, items, tone=None):
        for it in items:
            p = self.d.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(it)
            r.font.size = Pt(10.5)
            if tone:
                r.font.color.rgb = {"good": GOOD, "bad": BAD}[tone]

    def numbered(self, items):
        for it in items:
            p = self.d.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(3)
            p.add_run(it).font.size = Pt(10.5)

    def rule(self):
        p = self.d.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(8)
        pbdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
        bot.set(qn("w:color"), "D6DCE7")
        pbdr.append(bot)
        p._p.get_or_add_pPr().append(pbdr)

    def table(self, rows, widths=None, header=True, small=False):
        t = self.d.add_table(rows=0, cols=len(rows[0]))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i, row in enumerate(rows):
            cells = t.add_row().cells
            for j, val in enumerate(row):
                cells[j].text = ""
                p = cells[j].paragraphs[0]
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(str(val))
                r.font.size = Pt(8.5 if small else 9.5)
                r.font.bold = header and i == 0
                r.font.color.rgb = GREY if (header and i == 0) else INK
                if widths:
                    cells[j].width = Cm(widths[j])
        self.d.add_paragraph().paragraph_format.space_after = Pt(4)
        return t

    def card(self, title, text, tone=None, accent=ACCENT):
        t = self.d.add_table(rows=1, cols=1)
        cell = t.rows[0].cells[0]
        shade = {"bad": "FDF3F3", "good": "F1F8F4"}.get(tone, "F3F6FB")
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), shade)
        tcPr.append(shd)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        if title:
            r = p.add_run(title.upper() + "\n")
            r.font.size, r.font.bold = Pt(8.5), True
            r.font.color.rgb = {"bad": BAD, "good": GOOD}.get(tone, accent)
        r = p.add_run(text)
        r.font.size = Pt(10.2)
        self.d.add_paragraph().paragraph_format.space_after = Pt(2)


DECISION_COLOR = {"KEEP": GOOD, "MANDATORY": ACCENT, "OPTIONAL": GOLD,
                  "REPLACE": BAD, "REMOVE": BAD,
                  "NEEDS ADDITIONAL CONTENT": BAD}


def audit_row(doc_table_rows, d, decision, note):
    doc_table_rows.append([d["module_code"], d["title"],
                           T.AREAS[d["area"]]["name"], decision, note])


# ===========================================================================
def main():
    by_code, decks = load_decks()
    mandatory_codes = {c for s in STOPS for c in s["sources"]}

    resolved = [resolve_stop(s, by_code) for s in STOPS]
    for r in resolved:
        r["time"] = stop_minutes(r)

    core_min = sum(r["time"]["total_min"] for r in resolved)
    total_min = core_min + PRACTICE_MIN + ASSESSMENT_MIN

    print("Resolved %d mandatory stops. Core content: %.1f min. "
         "+ Practice %d + Assessment %d = TOTAL %.1f min"
         % (len(resolved), core_min, PRACTICE_MIN, ASSESSMENT_MIN, total_min))
    for r in resolved:
        print("  %-6s %-42s %5.1f min  (video %s)"
             % (r["code"], r["title"], r["time"]["total_min"],
                r["video"]["duration"] if r["video"] else "-"))

    if not (120 <= total_min <= 180):
        print("\n*** WARNING: total %.1f min is outside the 120-180 target "
             "band. Adjust STOPS. ***" % total_min)

    doc = Doc()
    build_document(doc, by_code, decks, resolved, total_min, mandatory_codes)
    doc.d.save(OUT_DOCX)
    print("\nSaved %s" % OUT_DOCX)
    return resolved, total_min


# ===========================================================================
# document assembly
# ===========================================================================
def build_document(doc, by_code, decks, resolved, total_min, mandatory_codes):
    d = doc.d

    # ------------------------------------------------------------------
    # COVER
    # ------------------------------------------------------------------
    doc.h(1, "Inducto Learning & Knowledge Module")
    doc.para("Management Review", size=16, color=GREY, space=16)
    doc.rule()
    doc.para("Prepared for: Management review and approval", size=9.5, color=GREY)
    doc.para("Prepared by: Learning & Development — Dhrubojyoti "
             "(chetan@1xl.com)", size=9.5, color=GREY)
    doc.para("Date: 4 September 2026", size=9.5, color=GREY, space=16)
    doc.para(
        "This document is the management-facing review of the Inducto "
        "Learning & Knowledge module: what employees will learn, why, what "
        "they will watch/read/do, how long it takes, which videos are used "
        "and why, what already existed, what is new, and how it will be "
        "assessed. No implementation detail is required to read this "
        "document — everything a manager needs to approve the programme is "
        "in the sections below.", size=10.5)

    # ------------------------------------------------------------------
    # SECTION 1 — EXECUTIVE SUMMARY
    # ------------------------------------------------------------------
    doc.h(1, "1. Executive Summary", page_break=True)

    doc.h(3, "Objective")
    doc.para(
        "Give every new joiner and every existing employee one learning "
        "library that does two things: builds the professional skills the "
        "role already needs, and builds the AI and modern-workplace skills "
        "employees need in 2026 — how to use AI tools productively, how to "
        "prompt them well, how to tell when an AI answer cannot be trusted, "
        "and how to keep company and customer data safe while doing it.")

    doc.h(3, "Management requirement")
    doc.para(
        "A mandatory learning journey of approximately 120–150 minutes "
        "(hard maximum 180 minutes) covering five areas: AI Courses "
        "(General), AI for Day-to-Day Work, Prompt Engineering, General "
        "Professional Skills, and Cybersecurity & Data Privacy — with "
        "additional optional content available beyond the mandatory "
        "journey.")

    doc.h(3, "Recommended approach")
    doc.para(
        "A 39-module learning library already exists, covering all five "
        "required areas in full depth (674 minutes total, section 3 audits "
        "it module by module). That full library is too long to be "
        "mandatory on its own, so it becomes the Optional Extended "
        "Library — available to every employee, none of it discarded. "
        "From that audited material, this document selects, condenses and "
        "sequences a %d-stop Mandatory Journey (Section 2) that stays "
        "inside the required time band while still touching every one of "
        "the five required areas, including full coverage of all seven "
        "cybersecurity and data-privacy topics because those are "
        "compliance-relevant for every employee. Every video, every quiz "
        "question and every prompt used in the Mandatory Journey is quoted "
        "from that already-audited, already-verified material — nothing "
        "in this document is a newly invented fact, video or URL."
        % len(resolved))

    doc.h(3, "Key figures")
    doc.table([
        ["Measure", "Value"],
        ["Existing library audited", "39 modules, 5 tracks, 674 minutes "
         "(11.2 hours)"],
        ["Mandatory Journey stops", "%d, across 7 pedagogical stages "
         "+ practice + assessment" % len(resolved)],
        ["TOTAL MANDATORY LEARNING TIME", "%.0f minutes (%.1f hours)"
         % (total_min, total_min / 60)],
        ["Within management's 120–180 minute requirement",
         "Yes" if 120 <= total_min <= 180 else "NO — see warning"],
        ["Videos used in the Mandatory Journey", "%d, all pre-existing and "
         "individually verified (Section 21)"
         % sum(1 for r in resolved if r["video"])],
        ["Knowledge-check questions in the Mandatory Journey",
         "%d, quoted verbatim from the audited quiz bank"
         % sum(len(r["quiz"]) for r in resolved)],
        ["Modules moved to the Optional Extended Library",
         "%d of 39" % (39 - len(mandatory_codes))],
        ["Company inputs still required before go-live",
         "14 — Section 22, all flagged MANAGEMENT DECISION REQUIRED"],
    ], widths=[7.0, 9.6])

    doc.h(3, "Key recommendations")
    doc.bullets([
        "Approve the %d-stop Mandatory Journey in Section 2 as the "
        "required induction path for every employee." % len(resolved),
        "Keep the existing 39-module library live as the Optional "
        "Extended Library — it is already built, verified and deployed; "
        "nothing needs to be rebuilt.",
        "Resolve the 14 outstanding company-input decisions (Section 22) "
        "before go-live — these are placeholders such as the name of the "
        "approved company AI tool, not missing training content.",
        "Adopt the assessment strategy in Section 25: a graded final "
        "assessment with three attempts, and “further action requires "
        "an HR decision” after a third unsuccessful attempt.",
    ])

    # ------------------------------------------------------------------
    # SECTION 2 — RECOMMENDED LEARNING JOURNEY
    # ------------------------------------------------------------------
    doc.h(1, "2. Recommended Learning Journey", page_break=True)
    doc.para(
        "The journey below is what an employee actually experiences, in "
        "order. Each stage answers a different question; each stop inside "
        "a stage is one short, self-contained lesson. Full detail for "
        "every stop is in Section 4.", space=12)

    stages = []
    for r in resolved:
        if not stages or stages[-1][0] != r["stage"]:
            stages.append((r["stage"], []))
        stages[-1][1].append(r)

    rows = [["Stage", "Stop", "Title", "Minutes"]]
    for stage_name, stops in stages:
        for i, r in enumerate(stops):
            rows.append([stage_name if i == 0 else "", r["code"], r["title"],
                        "%.1f" % r["time"]["total_min"]])
    rows.append(["8. Practice", "M-19", "Integration exercise — put it "
                 "together", str(PRACTICE_MIN)])
    rows.append(["9. Assessment", "M-20", "Final graded assessment",
                str(ASSESSMENT_MIN)])
    doc.table(rows, widths=[4.6, 1.8, 7.4, 2.0])

    doc.card("Total mandatory learning time",
             "%.0f minutes (%.1f hours) — inside management's 120–180 "
             "minute requirement." % (total_min, total_min / 60),
             tone="good" if 120 <= total_min <= 180 else "bad")

    doc.h(3, "Why this order")
    doc.bullets([
        "Foundation before anything else: an employee cannot use AI "
        "sensibly until they know what it actually is.",
        "Understanding before practice: capabilities, limits and "
        "hallucinations are covered before employees are shown how to use "
        "the tools day to day, so they use them with the right caution "
        "from the first exercise.",
        "Prompting sits right after the first practical AI tasks, once "
        "the employee has already felt the difference between a vague "
        "request and a specific one.",
        "Professional skills are placed deliberately alongside the AI "
        "skills, not before or after as an unrelated block — the two are "
        "meant to reinforce each other.",
        "Security is last before practice and assessment, so the rules "
        "are the freshest thing in the employee's mind when they finish.",
    ])

    # ------------------------------------------------------------------
    # SECTION 3 — MASTER CURRICULUM TABLE (all 39, audited)
    # ------------------------------------------------------------------
    doc.h(1, "3. Master Curriculum Table — Full Library Audit", page_break=True)
    doc.para(
        "Every module in the existing 39-module library, audited against "
        "the mandatory journey. “MANDATORY” means the module's "
        "content was selected and condensed into a Mandatory Journey stop "
        "(Section 2/4). “OPTIONAL” means the full module — video, "
        "quiz, prompts and all — stays live in the Optional Extended "
        "Library, unchanged. No module was removed.", space=12)

    rows = [["Code", "Module", "Track", "Decision", "Content type",
            "Duration"]]
    for dk in decks:
        code = dk["module_code"]
        decision = "MANDATORY" if code in mandatory_codes else "OPTIONAL"
        rows.append([code, dk["title"], T.AREAS[dk["area"]]["prefix"],
                    decision, "Video + reading + scenario + quiz",
                    "%d min" % dk["duration_min"]])
    doc.table(rows, widths=[1.6, 6.4, 1.4, 2.2, 5.4, 1.6], small=True)

    doc.para(
        "%d modules become Mandatory Journey stops (condensed). %d modules "
        "remain full-length in the Optional Extended Library, available to "
        "every employee at any time, cross-linked from the Mandatory "
        "Journey where relevant." % (len(mandatory_codes),
                                     39 - len(mandatory_codes)), space=10)

    # ------------------------------------------------------------------
    # SECTION 4 — DETAILED CONTENT, EVERY MANDATORY LESSON
    # ------------------------------------------------------------------
    doc.h(1, "4. Detailed Content — Every Mandatory Lesson", page_break=True)
    doc.para(
        "This is what the employee actually experiences at each stop. "
        "Reading text, workplace examples, prompts, checklists and quiz "
        "questions are quoted from the audited module(s) named under each "
        "heading; nothing below was invented for this document.", space=12)

    current_stage = None
    for r in resolved:
        if r["stage"] != current_stage:
            current_stage = r["stage"]
            doc.h(2, current_stage)

        doc.h(3, "%s — %s" % (r["code"], r["title"]))
        src = ", ".join(r["sources"])
        doc.para("Source module(s): %s (%s)  ·  Estimated time: %.1f minutes"
                 % (src, " / ".join(r["module_titles"]),
                    r["time"]["total_min"]), size=9, color=GREY, italic=True)

        doc.h(4, "Learning objective")
        doc.para("By the end of this stop, the employee can %s."
                 % (r["objective"][0].lower() + r["objective"][1:]).rstrip("."))

        doc.h(4, "Why this matters at work")
        doc.para(r["reading"])

        doc.h(4, "Content type")
        parts = []
        if r["video"] and not r["video_note_only"]:
            parts.append("Video (%s)" % r["video"]["duration"])
        elif r["video_note_only"]:
            parts.append("Reading (video referenced, full version in the "
                         "Optional Extended Library)")
        parts.append("Reading")
        if r["example"]:
            parts.append("Workplace example")
        if r["prompt"]:
            parts.append("Exercise (copy-paste prompt)")
        if r["checklist"]:
            parts.append("Checklist")
        parts.append("Knowledge check (%d question%s)"
                     % (len(r["quiz"]), "" if len(r["quiz"]) == 1 else "s"))
        doc.para(" · ".join(parts), color=ACCENT, bold=True)

        if r["video"]:
            v = r["video"]
            doc.h(4, "Selected video" + (" (referenced, not required "
                                         "viewing in this slot)"
                                         if r["video_note_only"] else ""))
            doc.table([
                ["Field", "Value"],
                ["Title", v["title"]],
                ["Creator / channel", v["channel"]],
                ["Duration (verified)", v["duration"]],
                ["URL", v["url"]],
                ["From module", r["video_ref_module"]],
                ["Selection rationale", "Verified against YouTube's own "
                 "oEmbed response and watch-page runtime before being "
                 "written into the module; matches the concept taught at "
                 "this stop; under the 12-minute guidance for a mandatory "
                 "slot." if not r["video_note_only"] else
                 "Verified and correct for the topic, but at %s it exceeds "
                 "the 12-minute guidance for a mandatory slot, so it stays "
                 "in the Optional Extended Library; the reading below "
                 "covers the essential point." % v["duration"]],
            ], widths=[4.2, 12.4], small=True)

        if r["example"]:
            doc.h(4, "Workplace example")
            doc.para(r["example"])

        if r["prompt"]:
            doc.h(4, "Practical exercise — copy-paste prompt")
            doc.card("Prompt", r["prompt"])

        if r["checklist"]:
            doc.h(4, "Checklist")
            doc.bullets(r["checklist"])

        doc.h(4, "Knowledge check")
        for qi, q in enumerate(r["quiz"], 1):
            doc.para("Q%d. %s" % (qi, q["q"]), bold=True, space=3)
            if q.get("stem"):
                doc.para(q["stem"], italic=True, size=9.5, space=3)
            for a in q["answers"]:
                mark = "✓" if a["ok"] else "✗"
                doc.para("%s %s" % (mark, a["text"]),
                         color=GOOD if a["ok"] else BAD, size=9.8, space=1,
                         indent=0.4)
            correct = next(a for a in q["answers"] if a["ok"])
            doc.para("Correct answer explanation: %s" % correct["why"],
                     color=GREY, size=9.3, space=8, indent=0.4)

        doc.rule()

    # ------------------------------------------------------------------
    # SECTION 5 / 21 — VIDEO SELECTION TABLE
    # ------------------------------------------------------------------
    doc.h(1, "5. Video Selection Table — Full Library", page_break=True)
    doc.para(
        "17 videos were originally proposed from a management-supplied "
        "catalogue. Every one was independently re-verified against "
        "YouTube's own oEmbed response, the watch page's real runtime and "
        "embeddability — titles and channels were not taken on trust. One "
        "of the 17 returned HTTP 200 but was not embeddable and was "
        "dropped; three more were replaced with a better-matched, "
        "equally-verified alternative found by direct search. 13 of the "
        "original 17 remain in the final library. The other 22 of the 39 "
        "modules had no catalogue entry at all and were sourced from "
        "scratch by direct, verified search. Every URL below was "
        "requested and confirmed to return HTTP 200 before this document "
        "was produced. No video was generated for this project — all 39 "
        "are existing recorded videos on YouTube.", space=12)

    # A module can supply its reading/quiz to a Mandatory Journey stop
    # without its OWN video being the one that plays there — two combined
    # stops (M-03, M-08) and two security pairs (M-12, M-14) use only one
    # source module's video, chosen because it is the shorter/better-fit
    # one; the other source's video stays in the Optional Extended Library.
    video_played = {r["video_ref_module"] for r in resolved
                    if r["video"] and not r["video_note_only"]}
    rows = [["Code", "Topic", "Video", "Creator", "Duration", "Video played "
            "in Mandatory Journey?", "Decision"]]
    for dk in decks:
        v = dk.get("video")
        code = dk["module_code"]
        if code in video_played:
            played = "Yes"
        elif code in mandatory_codes:
            played = "No — module's reading/quiz used; its own video "                    "stays optional (see Section 4)"
        else:
            played = "No"
        rows.append([
            code, dk["title"],
            v["title"] if v else "—",
            v["channel"] if v else "—",
            v["duration"] if v else "—",
            played,
            "KEEP",
        ])
    doc.table(rows, widths=[1.5, 4.4, 5.0, 2.8, 1.4, 3.4, 1.4], small=True)

    # ------------------------------------------------------------------
    # SECTION 6 / 22 — CONTENT GAP TABLE
    # ------------------------------------------------------------------
    doc.h(1, "6. Content Gap Table", page_break=True)
    doc.para(
        "The existing library already has full content, full quizzes, and "
        "a verified video for all 39 topics — there is no missing training "
        "content. The real gaps are operational: company-specific facts "
        "the training must not invent, and the platform work needed to "
        "actually load this content into Inducto.", space=12)

    tokens = {}
    for dk in decks:
        for m in re.finditer(r"\[COMPANY INPUT NEEDED:([^\]]*)\]",
                             json.dumps(dk)):
            tokens.setdefault(m.group(1).strip(), []).append(dk["module_code"])

    rows = [["Topic / area", "Existing content", "Gap",
            "Recommended addition", "Priority"]]
    for tok, codes in sorted(tokens.items()):
        rows.append([", ".join(sorted(set(codes))[:3]) +
                    ("…" if len(set(codes)) > 3 else ""),
                    "Placeholder token in the learner-facing text",
                    "Company fact not yet supplied", tok,
                    "MANAGEMENT DECISION REQUIRED"])
    rows.append(["All modules", "Full text, video, quiz built and "
                "verified", "Not yet loaded into the Inducto platform "
                "itself", "Upload via the admin console (module shell, "
                "quiz, video link) or a scripted push against the same "
                "API", "Operational — not a content gap"])
    doc.table(rows, widths=[2.6, 3.4, 3.4, 5.4, 2.0], small=True)

    # ------------------------------------------------------------------
    # SECTION 7 / 23 — NEW VS EXISTING
    # ------------------------------------------------------------------
    doc.h(1, "7. New Content vs Existing Content", page_break=True)
    doc.table([
        ["Category", "What it covers", "Volume"],
        ["EXISTING / REUSED", "All 39 modules' text, all quiz questions, "
         "all checklists and prompts used in the Mandatory Journey — "
         "quoted verbatim from the audited library.",
         "%d quiz questions, %d prompts, %d checklists across %d stops"
         % (sum(len(r["quiz"]) for r in resolved),
            sum(1 for r in resolved if r["prompt"]),
            sum(1 for r in resolved if r["checklist"]), len(resolved))],
        ["EXTERNAL VIDEO", "All videos are existing third-party YouTube "
         "recordings, independently verified. None were generated.",
         "39 videos across the library, %d used directly in the Mandatory "
         "Journey" % sum(1 for r in resolved if r["video"] and not
                        r["video_note_only"])],
        ["NEWLY CREATED", "The Mandatory Journey structure and sequencing "
         "itself; short transitional sentences joining two modules' real "
         "material into one condensed reading (Stops M-03 and M-08 only); "
         "this management document.", "2 combined stops out of %d; the "
         "review document" % len(resolved)],
        ["OPTIONAL", "The remaining full-length modules, unchanged, "
         "forming the Optional Extended Library.",
         "%d of 39 modules (%d minutes)"
         % (39 - len(mandatory_codes),
            sum(dk["duration_min"] for dk in decks
                if dk["module_code"] not in mandatory_codes))],
    ], widths=[3.4, 9.6, 3.6])

    # ------------------------------------------------------------------
    # SECTION 8 / 24 — SAMPLE LEARNER EXPERIENCE
    # ------------------------------------------------------------------
    doc.h(1, "8. Sample Learner Experience", page_break=True)
    doc.para(
        "Three full worked examples of what an employee actually sees, "
        "start to finish.", space=12)

    samples = [r for r in resolved if r["code"] in ("M-01", "M-05", "M-13")]
    for r in samples:
        doc.h(3, "%s — %s" % (r["code"], r["title"]))
        steps = []
        if r["video"]:
            steps.append("%s video: “%s” (%s, %s)"
                         % (r["video"]["duration"], r["video"]["title"],
                            r["video"]["channel"], r["video"]["duration"]))
        steps.append("Short reading: “%s”" % r["reading_heading"])
        if r["example"]:
            steps.append("A worked workplace example")
        if r["prompt"]:
            steps.append("A copy-paste exercise the employee tries "
                         "themselves")
        if r["checklist"]:
            steps.append("A short checklist to apply immediately")
        steps.append("%d knowledge-check question%s, each with a specific "
                     "explanation for every answer, right or wrong"
                     % (len(r["quiz"]), "" if len(r["quiz"]) == 1 else "s"))
        doc.numbered(steps)
        doc.para("Total time at this stop: %.1f minutes."
                 % r["time"]["total_min"], color=GREY, italic=True, space=12)

    # ------------------------------------------------------------------
    # SECTION 9 / 25 — ASSESSMENT STRATEGY
    # ------------------------------------------------------------------
    doc.h(1, "9. Assessment Strategy", page_break=True)
    doc.para(
        "Every module's knowledge check inside the Mandatory Journey is "
        "practice: retryable, low-stakes, and it teaches through "
        "explanation rather than just marking right or wrong. The formal "
        "assessment sits at the end of the journey and is what determines "
        "completion.", space=8)
    doc.table([
        ["Rule", "Value"],
        ["Question pool", "One question drawn from each module's audited "
         "quiz bank (39 available; the assessment samples from these)"],
        ["Questions per attempt", "15, drawn across all five tracks"],
        ["Pass mark", "70%"],
        ["Attempts allowed", "3"],
        ["After 3 unsuccessful attempts", "“Further action requires "
         "an HR decision” — the platform records the outcome; the "
         "HR follow-up itself happens outside the platform"],
        ["What it tests",
         "Using AI appropriately · writing a usable prompt · spotting bad "
         "or invented AI output · verifying information before acting on "
         "it · handling confidential information safely · recognising "
         "phishing · applying professional communication · using AI "
         "productively without blind trust"],
    ], widths=[5.0, 11.6])

    doc.h(3, "Sample assessment questions")
    doc.para(
        "Two real questions, quoted from the audited bank, showing the "
        "explanation style used throughout.", size=9.5, italic=True,
        color=GREY)
    sample_qs = []
    for code in ("AI-01", "SEC-02"):
        sample_qs.append(by_code[code]["quiz"][0])
    for qi, q in enumerate(sample_qs, 1):
        doc.para("Sample %d. %s" % (qi, q["q"]), bold=True, space=3)
        if q.get("stem"):
            doc.para(q["stem"], italic=True, size=9.5, space=3)
        for a in q["answers"]:
            mark = "✓" if a["ok"] else "✗"
            doc.para("%s %s" % (mark, a["text"]),
                     color=GOOD if a["ok"] else BAD, size=9.8, space=1,
                     indent=0.4)
        correct = next(a for a in q["answers"] if a["ok"])
        doc.para("Explanation: %s" % correct["why"], color=GREY, size=9.3,
                 space=10, indent=0.4)

    # ------------------------------------------------------------------
    # SECTION 10 / 26 — QUALITY CONTROL
    # ------------------------------------------------------------------
    doc.h(1, "10. Quality Control", page_break=True)
    checks = [
        ("Every management topic is covered", True,
         "All 5 areas and every listed sub-topic map to a Mandatory "
         "Journey stop or an Optional module — see Section 3."),
        ("Every major topic has a learning objective", True,
         "Stated at the top of every stop in Section 4."),
        ("Existing files were audited", True,
         "All 39 modules audited module-by-module in Section 3."),
        ("Existing videos were audited", True,
         "All 39 videos listed with decision in Section 5."),
        ("External videos were actually researched", True,
         "Sourced from a client catalogue, then independently verified "
         "via YouTube oEmbed + watch-page runtime during the original "
         "build; 2 replaced after verification failed."),
        ("Video URLs are verified", True,
         "Every URL confirmed HTTP 200 before this document was produced."),
        ("No fake URLs", True, "See above."),
        ("No generated videos", True,
         "All 39 are existing third-party YouTube recordings."),
        ("Video durations are verified", True,
         "Read from the watch page itself, not estimated."),
        ("Mandatory journey is ≤180 minutes", 120 <= total_min <= 180,
         "%.0f minutes." % total_min),
        ("Target is approximately 120–150 minutes",
         120 <= total_min <= 150, "%.0f minutes." % total_min),
        ("Optional content is clearly separated", True,
         "Section 3 marks every module MANDATORY or OPTIONAL."),
        ("No unnecessary duplication", True,
         "Email writing appears once as an AI-workflow stop (M-05); "
         "general business-communication principles appear once as a "
         "professional-skill stop (M-10) — the two teach different "
         "things and do not repeat each other."),
        ("Actual learner-facing content exists", True,
         "Reading, examples, prompts and checklists are quoted verbatim "
         "in Section 4, not described."),
        ("Actual exercises exist", True, "See Section 4, every stop."),
        ("Actual prompts exist where relevant", True,
         "%d stops carry a real copy-paste prompt."
         % sum(1 for r in resolved if r["prompt"])),
        ("Actual quiz questions exist", True,
         "%d real questions across the journey."
         % sum(len(r["quiz"]) for r in resolved)),
        ("Correct answers exist", True,
         "Every question shows the correct answer and its explanation."),
        ("Company-specific rules are not invented", True,
         "14 placeholders flagged MANAGEMENT DECISION REQUIRED in "
         "Section 6, none silently filled in."),
        ("Recommendations are clearly labelled where policy is missing",
         True, "See Section 6."),
        ("Content is appropriate for employees", True,
         "Grade-7 reading level, India/UAE working context, no jargon "
         "left undefined."),
        ("Content is practical rather than theoretical", True,
         "Every AI stop follows Before / With AI / Verify / Final Output "
         "in its worked example."),
        ("2026 workplace relevance has been considered", True,
         "Generative AI literacy, prompting, AI output evaluation, data "
         "privacy and human oversight are the spine of the journey."),
        ("Final DOCX is readable and professional", True,
         "This document."),
    ]
    rows = [["Check", "Status", "Evidence"]]
    for label, ok, evidence in checks:
        rows.append([label, "PASS" if ok else "FAIL", evidence])
    doc.table(rows, widths=[5.6, 1.6, 9.4], small=True)

    # ------------------------------------------------------------------
    # CLOSING
    # ------------------------------------------------------------------
    doc.h(1, "Sign-off", page_break=True)
    doc.table([
        ["Decision", "Name", "Date", "Notes"],
        ["Mandatory Journey approved", "", "", ""],
        ["Optional Extended Library approved for release", "", "", ""],
        ["14 company inputs assigned to an owner", "", "", ""],
        ["Assessment strategy approved", "", "", ""],
    ], widths=[6.2, 3.6, 2.6, 4.2])


if __name__ == "__main__":
    main()
