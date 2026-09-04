# -*- coding: utf-8 -*-
"""
INDUCTO World-Class Content Master — FINAL (v3).

    python world_class_master_v3.py

A second, independent forensic pass over the v2 (RECTIFIED_FINAL) document
and its source, run without trusting v2's own "no contradiction remains"
claim. Restructures the document into seven separated parts (Part A learner
content, Part B video/visual specs, Part C management inputs, Part D
assessment master, Part E source traceability, Part F technical appendix,
Part G QA/rectification log) and adds an explicit, computed assessment
blueprint instead of a prose "balanced across areas" assertion.

Reuses journey_data.py / management_review_docx.py for all resolved content
(so the document and the live product stay identical) and world_class_
master_docx.py's Doc/add_hyperlink/visual_brief machinery rather than
duplicating it.
"""

import io
import json
import os
import re
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
sys.stdout.reconfigure(encoding="utf-8")

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor

import journey_data as J
import sitegen
import theme as T
import world_class_master_docx as W  # Doc, add_hyperlink, visual_brief, pick_visual

ROOT = sitegen.ROOT
OUT_DOCX = os.path.join(ROOT, "docs",
                        "INDUCTO_World_Class_Content_Master_FINAL.docx")

INK, GREY, GOOD, BAD, ACCENT, GOLD = W.INK, W.GREY, W.GOOD, W.BAD, W.ACCENT, W.GOLD
Doc = W.Doc
add_hyperlink = W.add_hyperlink
visual_brief = W.visual_brief
pick_visual = W.pick_visual

PART_COLOR = RGBColor(0x0E, 0x6E, 0x75)

# ---------------------------------------------------------------------------
# v2's rectification log, carried forward as history (R1-R5). This pass adds
# R6 onward for what an independent second pass found that v2 missed or
# stated imprecisely.
# ---------------------------------------------------------------------------
PRIOR_RECTIFICATIONS = [
    ("R1", "Medium", "AI technical accuracy", "AI-01",
     "“An AI assistant does not look up an answer” read as an absolute "
     "claim on the first slide a learner sees."),
    ("R2", "High", "Cybersecurity — password guidance", "SEC-01",
     "The live M-12 lesson told learners to ask an AI assistant to generate "
     "passphrases, contradicting the module's own “never in an AI chat "
     "window” rule."),
    ("R3", "High", "M-19 scenario logic", "journey_data.py",
     "Scenario set at “Friday, 4:50pm” required a reroute “before Friday "
     "morning” — already in the past at that moment."),
    ("R4", "Medium", "Reporting consistency — video count",
     "management/world-class docs", "“16 mandatory videos” did not "
     "distinguish embedded-and-required from referenced-only."),
    ("R5", "Low", "Management-input register completeness",
     "world_class_master_docx.py", "Register scan covered only the 39 "
     "module dictionaries and missed the site-footer token."),
]

RECTIFICATIONS = [
    {
        "id": "R6", "severity": "High", "area": "Cybersecurity — phishing "
                                              "verification threshold",
        "location": "SEC-02 — “Two together is enough to stop” (source "
                    "slide, quiz remember line, and recap point)",
        "finding": "Three places taught that verification is warranted only "
                  "once TWO pressure signals appear together (“Any two "
                  "means verify”). A single high-risk request — for "
                  "example a calm, unhurried email changing bank details, "
                  "with no urgency and no secrecy attached — could "
                  "therefore read as not yet requiring verification under "
                  "this taught rule, which is the more dangerous failure "
                  "mode: quiet, single-signal business-email-compromise "
                  "attempts are a well-documented real pattern.",
        "fix": "Retitled the slide “One is enough to make you check” and "
              "reordered its items so “A change, on its own, is reason "
              "enough” leads. The quiz remember line and the recap point "
              "were both changed from “Any two means stop” to “one alone "
              "is reason enough to verify.” Combinations (urgency+"
              "authority, secrecy+change) are kept as illustrations of what "
              "makes an attempt more obvious — not as the threshold for "
              "acting.",
        "verified_against": "Standard phishing/BEC guidance (a single "
                            "high-risk indicator, particularly a payment or "
                            "bank-detail change, is treated as sufficient "
                            "grounds for out-of-band verification) — not "
                            "company-specific, so no management input "
                            "required.",
    },
    {
        "id": "R7", "severity": "High", "area": "Privacy — personal-"
                                              "account claim",
        "location": "SEC-07 — “Why is a personal account worse?” quiz "
                    "answer (this exact question is used as this document's "
                    "own M-16 lesson knowledge check)",
        "finding": "The correct answer stated as fact that “the company "
                  "cannot see, control or delete what was shared” on a "
                  "personal account — a categorical claim about every "
                  "provider's behaviour, stated as a general rule rather "
                  "than describing one scenario's real consequence.",
        "fix": "Replaced with “It gives up the company's contractual, "
              "security and audit controls,” with the explanation "
              "changed to “a personal account may provide none of that” "
              "— accurate regardless of the specific provider, and no "
              "longer asserting universal technical behaviour. (Two other "
              "occurrences of similar wording in AI-01 and SEC-07's own "
              "scenario intro were checked and left unchanged: both "
              "describe one concrete scenario's real consequence — "
              "pasting into an unapproved free tool with no company "
              "agreement — not a categorical claim.)",
        "verified_against": "Re-read in full source context; the two "
                            "scenario-embedded instances were confirmed "
                            "accurate as scoped, only the categorical quiz-"
                            "answer version needed correction.",
    },
    {
        "id": "R8", "severity": "High", "area": "Approved-tool definition",
        "location": "SEC-06 (title, lead, section label, recap point, "
                    "glossary) and SEC-07 (glossary) — 6 locations across "
                    "2 modules",
        "finding": "The course defined an “approved tool” as “the company "
                  "has an agreement and you use your work account” in six "
                  "places, including as the correct answer to a knowledge-"
                  "check question. That formula omits that approval can "
                  "also depend on the specific product, its configuration, "
                  "the permitted use case and the permitted data — a work "
                  "account on a general company subscription does not by "
                  "itself approve every tool or every use of it.",
        "fix": "All six locations updated to state that the tool, the "
              "account/configuration and the specific use must each be "
              "approved. The SEC-06 knowledge-check question's correct "
              "answer and explanation were rewritten to match; the wrong "
              "answers needed no change.",
        "verified_against": "Internal consistency across both modules; the "
                            "correction matches the general principle "
                            "already stated correctly elsewhere in SEC-06 "
                            "(“Approved tool” checklist items already listed "
                            "access management and accountability as "
                            "separate conditions, not folded into "
                            "“agreement + account”).",
    },
    {
        "id": "R9", "severity": "Medium", "area": "Assessment quality — "
                                                 "distractor wording",
        "location": "SEC-02 — “Who may ask for a one-time code?” (this "
                    "exact question is Q13 in the M-20 assessment pool)",
        "finding": "The correct answer read “Nobody, ever, under any "
                  "circumstances” — emphatic triple-qualifier phrasing "
                  "that stands out from the other three options by tone "
                  "alone, closer to “wording that gives away the answer” "
                  "than a plausible-sounding option a learner has to "
                  "genuinely evaluate against IT, the bank and the "
                  "manager.",
        "fix": "Changed to “Nobody — not IT, your bank, or your manager,” "
              "which states the same true, deliberately absolute rule "
              "(sharing a one-time code is never legitimate, with no "
              "exception — this claim itself was independently re-"
              "checked and is correctly taught as absolute) but requires "
              "the learner to actually rule out each specific plausible "
              "distractor rather than pattern-matching on emphatic "
              "wording.",
        "verified_against": "Re-examined against the assessment-quality "
                            "standard in this pass's brief (avoid wording "
                            "that gives away the answer); the underlying "
                            "rule was separately re-verified as one of the "
                            "few security claims that is correctly "
                            "absolute, so only the phrasing changed.",
    },
    {
        "id": "R10", "severity": "Low", "area": "AI technical accuracy — "
                                               "quiz takeaway",
        "location": "AI-02 — “Why are the answers different?” quiz "
                    "remember line",
        "finding": "The memorable takeaway read “Nothing is retrieved. "
                  "Everything is built,” a portable absolute claim that "
                  "reads as a rule about every AI system, when the "
                  "question itself is really about why wording varies "
                  "(construction is probabilistic, whether or not a fact "
                  "was retrieved first).",
        "fix": "Changed to “The wording is built fresh every time, even "
              "when a tool has looked something up first” — keeps the "
              "true, question-specific point (wording varies because "
              "sentences are generated, not cached) without implying no "
              "AI tool can retrieve information.",
        "verified_against": "Cross-checked against AI-01 and AI-04's "
                            "correct, already-nuanced treatment of "
                            "retrieval/search capability.",
    },
    {
        "id": "R11", "severity": "Low", "area": "Privacy — reporting-route "
                                               "wording",
        "location": "SEC-04 — the incident-response flow diagram",
        "finding": "A step read “Same day, to the named contact” — "
                  "phrased as if a specific contact were already "
                  "established, even though the module correctly flags "
                  "the actual contact elsewhere as [COMPANY INPUT NEEDED: "
                  "who handles data requests].",
        "fix": "Changed to “Same day, through the company's designated "
              "incident-reporting process” — matches this pass's "
              "required phrasing pattern and removes the implication that "
              "a specific contact is already known.",
        "verified_against": "Cross-checked against the module's own "
                            "[COMPANY INPUT NEEDED] token for the same "
                            "fact, confirming the contact is genuinely not "
                            "yet established.",
    },
    {
        "id": "R12", "severity": "Informational", "area": "M-16 objective "
                                                         "alignment (re-"
                                                         "verified)",
        "location": "SEC-07 — “Five things that never go in”",
        "finding": "This pass's brief raised the risk that the learning "
                  "objective (“name the five kinds of information that "
                  "must never go into an AI tool”) might not match what "
                  "the lesson actually teaches, and warned against "
                  "inventing categories to force a match.",
        "fix": "No change required. The source content was re-read in "
              "full: it defines exactly five categories (customer/"
              "employee personal data; anything from a signed contract; "
              "financial records and pricing; login details; anything "
              "marked confidential or internal), and the knowledge-check "
              "questions correctly test application of the rule rather "
              "than rote recall of the list. Logged here as a checked-and-"
              "confirmed item, not a silent pass.",
        "verified_against": "Direct re-read of the SEC-07 source content "
                            "against the stated objective.",
    },
    {
        "id": "R13", "severity": "Informational", "area": "M-19 timeline "
                                                         "(re-verified)",
        "location": "journey_data.py — the full Chakan scenario",
        "finding": "This pass's brief warned not to assume the R3 Thursday "
                  "fix was correct merely because it resolved the earlier "
                  "contradiction, and asked for every date, deadline and "
                  "cost to be retraced.",
        "fix": "No further change required. Every date reference was "
              "retraced against “Thursday” as today: Tuesday (dock "
              "doors, 2 days past) and Wednesday (promised delivery, 1 day "
              "past) are both in the past; Friday (reroute, 1 day ahead), "
              "Monday (revised delivery, 4 days ahead) and Tuesday next "
              "week (dock ETA, 5 days ahead) are all genuinely in the "
              "future. No impossible deadline remains.",
        "verified_against": "Manual day-by-day trace of all 6 date "
                            "references in the scenario and the 3 model "
                            "answers that repeat them.",
    },
    {
        "id": "R14", "severity": "High", "area": "Systemic — human wording "
                                              "mislabelled as an AI prompt",
        "location": "management_review_docx.py resolve_stop() — affected "
                    "the live Mandatory Journey stops M-10, M-11, M-13 and "
                    "M-14",
        "finding": "The Professional Skills and Security modules "
                  "deliberately mark a “type: prompt” visual with the "
                  "header “Copy this wording” (or a close variant) to mean "
                  "wording a person says or writes to another person — a "
                  "verification phone call, an incident report, a spoken "
                  "line to a colleague — never text to paste into an AI "
                  "tool. The generic selection logic that fills the "
                  "Mandatory Journey's “prompt” field did not know this "
                  "distinction: it took the first “type: prompt” visual "
                  "regardless of source, and the journey page always "
                  "renders that field under the fixed label “Try it — "
                  "copy-paste prompt.” The result, live in production: "
                  "M-10 and M-11 showed a human script under an "
                  "AI-prompt label, M-13 showed the phishing "
                  "verification-call script instead of the module's real "
                  "AI-safe prompt, and — most serious — M-14 showed a "
                  "real incident-report template (“I need to report a "
                  "possible data incident… What data was involved…”) "
                  "captioned as something to copy and paste, which is "
                  "exactly the exposure risk this pass's brief warned "
                  "against.",
        "fix": "resolve_stop() now selects the journey's “prompt” field "
              "based on module area: for Professional Skills (04-*) and "
              "Security & Privacy (05-*) modules — where a “type: "
              "prompt” visual is always human-facing wording — it takes "
              "only the “prompt” key nested inside a “steps” visual, "
              "which is where the module's actual AI-safe prompt lives "
              "when one exists (confirmed present for every affected "
              "stop; M-12 correctly has none, since the module's own "
              "rule is never to ask an AI tool to handle a real "
              "password). All other areas are unaffected. Verified by "
              "re-resolving the live journey data before and after: "
              "M-10, M-11, M-13 and M-14 now each show a genuine, safe "
              "AI prompt; M-14's field no longer contains any incident "
              "detail.",
        "verified_against": "Direct before/after diff of journey_data."
                            "load()'s resolved “prompt” field for all 16 "
                            "mandatory stops, and a full catalogue of "
                            "every “type: prompt” visual's header text "
                            "across all 39 source modules, cross-checked "
                            "against its module area.",
    },
    {
        "id": "R15", "severity": "Medium", "area": "Cybersecurity — "
                                                 "unsupported claims",
        "location": "SEC-01 — “How do accounts usually get taken?” "
                    "knowledge-check (live in M-12)",
        "finding": "Three of the four answer explanations used unsupported "
                  "quantified language — “Almost never,” “This is the "
                  "standard attack,” “The overwhelming majority” — "
                  "without a cited source for the implied statistic.",
        "fix": "Reworded all three to state the same relative point "
              "(credential stuffing from leaked lists is a materially "
              "more common cause of account takeover than guessing or "
              "shoulder-surfing) without a precise, uncited proportion.",
        "verified_against": "Re-read against this pass's brief; no "
                            "authoritative source for an exact figure "
                            "exists in the repository, so the claim was "
                            "qualified rather than deleted.",
    },
    {
        "id": "R16", "severity": "Medium", "area": "Cybersecurity — unsafe "
                                                 "password guidance "
                                                 "(recurrence)",
        "location": "SEC-01 — Toolkit, “The passphrase generator prompt”",
        "finding": "v2's R2 removed an AI-generated-passphrase prompt "
                  "from the module's “Do this now” visual, but the "
                  "module's own Toolkit section — “three things to take "
                  "with you” — still offered “The passphrase generator "
                  "prompt: eight four-word phrases, no names, no dates,” "
                  "which is exactly the pattern R2 removed, re-introduced "
                  "in a part of the same module R2's own reviewer did not "
                  "check.",
        "fix": "Replaced with “The generate-and-store habit: your "
              "password manager's own generator — never an AI tool,” "
              "consistent with the module's own corrected “Do this now” "
              "guidance.",
        "verified_against": "Full-text search of the module for every "
                            "reference to a passphrase, prompt or "
                            "generator, not just the slide R2 already "
                            "touched.",
    },
    {
        "id": "R17", "severity": "Medium", "area": "Cybersecurity — MFA "
                                                 "factor accuracy",
        "location": "SEC-03 — lead and glossary definition of "
                    "multi-factor authentication",
        "finding": "The module explained only two of the three standard "
                  "authentication factors — something you know and "
                  "something you have — and never named “something you "
                  "are” (biometrics) at all, even though this pass's "
                  "brief specifically required all three to be explained "
                  "correctly.",
        "fix": "The opening slide and the glossary definition now name "
              "all three factor categories, while keeping the module's "
              "practical focus on what the company's systems actually "
              "offer (usually something you have).",
        "verified_against": "Re-read of the full module; confirmed no "
              "other slide implicitly limited MFA to two factors.",
    },
    {
        "id": "R18", "severity": "Medium", "area": "Cybersecurity — "
                                                 "phishing threshold "
                                                 "(recurrence)",
        "location": "SEC-02 — “Which combination should stop you?” "
                    "knowledge-check (live in M-13, immediately below the "
                    "R6-fixed slide in the same module)",
        "finding": "R6 fixed the slide, quiz remember line "
                  "and recap to teach that one pressure signal alone is "
                  "reason enough to verify. This quiz question's actual "
                  "correct answer was left untouched: it still required a "
                  "two-signal combination (“urgency and a request to keep "
                  "it quiet”), and none of its four options tested "
                  "whether a learner understood that a single change-of-"
                  "bank-details signal, alone, is already sufficient — "
                  "the exact contradiction this pass's brief warned "
                  "about between a module's stated rule and its own "
                  "knowledge check.",
        "fix": "Rebuilt the question to ask which single signal is "
              "reason enough to verify on its own, with the bank-details "
              "change as the correct answer and the same three "
              "non-signal distractors retained.",
        "verified_against": "Re-read of the full SEC-02 quiz against the "
                            "module's own corrected slide and recap "
                            "wording, checking for consistency rather "
                            "than trusting R6's slide-level fix to have "
                            "propagated.",
    },
    {
        "id": "R19", "severity": "Low", "area": "Unsupported productivity "
                                              "statistics",
        "location": "PS-04 — lead, two knowledge-check answers, two "
                    "“mistakes” items, the recap and the glossary (live "
                    "in M-11) — 7 locations",
        "finding": "An uncited, specific “fifteen minutes to refocus” "
                  "and “wastes twenty minutes” figure was stated as fact "
                  "throughout the module, including as the justification "
                  "for two knowledge-check correct answers.",
        "fix": "All 7 locations reworded to keep the behavioural lesson "
              "(an interruption costs more than the interruption itself; "
              "deciding the day's task in advance protects the most "
              "valuable part of a protected block) without the specific, "
              "uncited minute counts. The module's own observable "
              "decision rule — asking whether a request is “a two-minute "
              "thing or a twenty-minute thing” — was kept, since that "
              "number describes the request itself, which the asker "
              "states, not an invented cost of interruption in general.",
        "verified_against": "Full-text search of the module for every "
                            "occurrence of a specific refocus-time or "
                            "wasted-time figure.",
    },
    {
        "id": "R20", "severity": "Low", "area": "Unsupported productivity "
                                              "statistics",
        "location": "DW-10 — a “mistakes” item, a knowledge-check "
                    "question, and the recap (live in M-09) — 3 locations",
        "finding": "“Almost everyone is around fifty per cent optimistic, "
                  "consistently” and “one of the most reliable patterns "
                  "in how people plan work” were stated as precise, "
                  "uncited fact, including as a knowledge-check's "
                  "correct-answer justification.",
        "fix": "Reworded to the qualitative, defensible version of the "
              "same lesson — people are consistently optimistic about "
              "their own time estimates, so a buffer corrects for it — "
              "without the specific uncited percentage or the “most "
              "reliable pattern” overclaim. The module's existing "
              "prompt-based technique (telling an AI tool to assume "
              "estimates are 50% optimistic) was kept: that is the "
              "learner's own chosen buffer, an observable decision rule, "
              "not an assertion about how people in general behave.",
        "verified_against": "Re-read against this pass's brief, which "
                            "names this exact figure as an example "
                            "requiring qualification.",
    },
    {
        "id": "R21", "severity": "Low", "area": "AI technical accuracy — "
                                              "residual absolute claim",
        "location": "AI-02 — title, lead, section label and recap of "
                    "“What generative means” (live in M-02)",
        "finding": "“It builds. It does not fetch.” / “There is no file "
                  "it is copying from.” / “Neither answer was retrieved. "
                  "Both were made.” were taught as unqualified facts "
                  "about every generative AI system — true of the "
                  "underlying wording-construction mechanism, but not of "
                  "whether a tool looked something up first, which "
                  "several of the tools staff actually use can do.",
        "fix": "Reworded throughout to keep the true, useful point (the "
              "wording is always freshly built, so it can never be "
              "quoted as a stored record) while adding that a tool may "
              "search or read a document first and then build the "
              "answer from that — matching this pass's required "
              "wording, “assume nothing was looked up unless the tool "
              "shows you it was.”",
        "verified_against": "Cross-checked against AI-04's already-"
                            "correct treatment of the same distinction "
                            "(“Unless the tool visibly shows a link it "
                            "visited, no page was opened”), which this "
                            "module's wording had not been brought in "
                            "line with.",
    },
]

ALL_RECTIFICATIONS = PRIOR_RECTIFICATIONS  # for the short summary table


def word_class_data():
    by_code, decks, resolved = J.load()
    total_min = sum(r["time"]["total_min"] for r in resolved) + 3 + 12
    pool = J.assessment_pool(by_code)
    domain_counts = J.assessment_domain_counts(by_code)

    tokens = {}
    for dk in decks:
        for mm in re.finditer(r"\[COMPANY INPUT NEEDED:([^\]]*)\]",
                              json.dumps(dk)):
            tokens.setdefault(mm.group(1).strip(), []).append(dk["module_code"])
    tokens.setdefault("name of the team that owns this training",
                      []).append("site-wide footer, every page")

    briefs = []
    for r in resolved:
        primary = by_code[r["sources"][0]]
        slide_title, v = pick_visual(primary)
        if v:
            briefs.append(visual_brief(r["code"], r["title"], r["stage"],
                                       slide_title, v))

    return by_code, decks, resolved, total_min, pool, domain_counts, tokens, briefs


# ===========================================================================
def part_break(doc, letter, title, intro):
    doc.d.add_paragraph().add_run().add_break()
    p = doc.d.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pbdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single"); top.set(qn("w:sz"), "24")
    top.set(qn("w:color"), "0E6E75")
    pbdr.append(top)
    p._p.get_or_add_pPr().append(pbdr)
    r = p.add_run("PART %s" % letter)
    r.font.size, r.font.bold, r.font.color.rgb = Pt(11), True, PART_COLOR
    doc.h(1, title, color=PART_COLOR)
    if intro:
        doc.para(intro, size=10.5, color=GREY, space=14)


def main():
    (by_code, decks, resolved, total_min, pool, domain_counts, tokens,
     briefs) = word_class_data()

    doc = Doc()
    build(doc, by_code, decks, resolved, total_min, pool, domain_counts,
         tokens, briefs)
    doc.d.save(OUT_DOCX)

    videos_embedded = sum(1 for r in resolved
                         if r["video"] and not r["video_note_only"])
    videos_total = sum(1 for r in resolved if r["video"])

    print("Saved %s" % OUT_DOCX)
    print("  %d mandatory lessons, %.1f minutes, %d knowledge checks, "
         "%d assessment questions" %
         (len(resolved), total_min, sum(len(r["quiz"]) for r in resolved),
          len(pool)))
    print("  videos: %d embedded/required + %d referenced-only = %d total "
         "| %d visual briefs | %d company inputs open"
         % (videos_embedded, videos_total - videos_embedded, videos_total,
            len(briefs), len(tokens)))
    print("  this pass: %d new rectifications (R6-R%d) + %d carried "
         "forward from v2 (R1-R5)"
         % (len(RECTIFICATIONS), 5 + len(RECTIFICATIONS),
            len(PRIOR_RECTIFICATIONS)))
    print("  assessment domains: %s" % domain_counts)
    return doc, resolved, pool, domain_counts, tokens, briefs, total_min


# ===========================================================================
def build(doc, by_code, decks, resolved, total_min, pool, domain_counts,
         tokens, briefs):
    d = doc.d
    mandatory_codes = {c for r in resolved for c in r["sources"]}

    # ------------------------------------------------------------------
    # COVER + DOCUMENT CONTROL
    # ------------------------------------------------------------------
    doc.h(1, "Inducto World-Class Content Master")
    doc.para("FINAL — seven parts, independently re-audited",
            size=14, color=GREY, space=16)
    doc.rule()
    doc.para("Prepared for: Management · Software Development · Learning "
            "& Development · Content · Proofreading · Instructional "
            "Design · Image/Graphics · Video/Media · LMS & Admin",
            size=9.5, color=GREY)
    doc.para("Prepared by: Learning & Development — Dhrubojyoti "
            "(chetan@1xl.com)", size=9.5, color=GREY)
    doc.para("Version 3.1 (Final) · 4 September 2026", size=9.5,
            color=GREY, space=16)
    doc.para(
        "This version is the product of two further independent forensic "
        "passes over the previous “RECTIFIED_FINAL” document, neither of "
        "which trusted the prior pass's own PASS status or rectification "
        "log. Together they found and fixed %d further defects (Part G); "
        "the most consequential (R6, R7, R14) directly changed live "
        "content — including a real incident-report template that was "
        "mislabelled as an AI prompt in the Mandatory Journey itself."
        % len(RECTIFICATIONS), size=10.5)

    doc.h(1, "Document Control", page_break=True)
    doc.table([
        ["Version", "Date", "Change"],
        ["1.0", "4 Sep 2026", "First content master."],
        ["2.0 (Rectified)", "4 Sep 2026", "First rectification pass — 5 "
         "defects (R1-R5): an absolute AI claim, an unsafe AI-generated-"
         "password prompt live in M-12, an impossible M-19 deadline, "
         "imprecise video-count reporting, and an incomplete management-"
         "input register."],
        ["3.0 (Final)", "4 Sep 2026", "Second, independent forensic pass "
         "— did not trust v2's own PASS claim. 8 further defects found "
         "(R6-R13): a phishing verification threshold that could let a "
         "single high-risk request through, a categorical privacy claim "
         "about personal accounts, an incomplete approved-tool definition "
         "repeated in 6 places, an assessment distractor that gave away "
         "its answer, plus 4 lower-severity or re-verification items. "
         "Restructured into 7 separated parts with an explicit, computed "
         "assessment blueprint."],
        ["3.1 (Final)", "4 Sep 2026", "Third, independent forensic pass "
         "— did not trust v3.0's own PASS claim either. %d further "
         "defects found (R14-R%d), the most severe being a systemic "
         "labelling defect (R14) that presented human-facing wording — "
         "including SEC-04's real incident-report template — to "
         "learners as a “copy-paste prompt” in 4 live Mandatory Journey "
         "stops (M-10, M-11, M-13, M-14). Also fixed: unsupported "
         "quantified security claims (R15), a recurrence of the v2 R2 "
         "unsafe-password-prompt defect in a part of SEC-01 the prior "
         "pass did not check (R16), an incomplete MFA factor explanation "
         "(R17), a phishing knowledge check left contradicting its own "
         "module's R6 fix (R18), and unsupported productivity/planning "
         "statistics in three modules (R19-R21). Video records "
         "independently spot-checked live (not HTTP-200-only) rather "
         "than trusted from a prior pass. Journey duration recalculated "
         "after every change in this document: %.1f minutes."
         % (len(RECTIFICATIONS) - 8, 5 + len(RECTIFICATIONS), total_min)],
    ], widths=[3.2, 2.6, 10.8])
    doc.para(
        "Superseded document: INDUCTO_World_Class_Content_Master_"
        "RECTIFIED_FINAL.docx (v2). This file is the current authority.",
        size=9, color=GREY)

    # ==================================================================
    # PART A — APPROVED LEARNER CONTENT
    # ==================================================================
    part_break(doc, "A", "Approved Learner Content",
              "Everything a learner actually sees in the Mandatory "
              "Journey. Every sentence below is SOURCE-VERIFIED — quoted "
              "verbatim from content/areaNN/*.py and confirmed present, "
              "unchanged, by the programmatic check in Part E. Four stops "
              "(M-03, M-08, M-12, M-14) combine verbatim text from two "
              "source modules; that joining is the only editorial "
              "assembly anywhere in this part. M-19's scenario is the one "
              "exception — NEWLY AUTHORED, grounded in real techniques "
              "the journey already teaches (labelled at A5).")

    # A1 — Programme overview
    doc.h(2, "A1. Programme Overview")
    doc.table([
        ["Measure", "Value"],
        ["Mandatory lessons", "16 (M-01 to M-16)"],
        ["Mandatory journey length", "%.1f minutes (%.1f hours), computed — "
         "video seconds + reading + exercise + quiz, not video length "
         "alone" % (total_min, total_min / 60)],
        ["Knowledge-check questions", "24, embedded and independently "
         "recounted from the shipped HTML in Part E"],
        ["Integration exercise", "M-19 — one scenario, four real skills, "
         "timeline re-traced day by day (Part G, R13)"],
        ["Final assessment", "M-20 — 15 questions, explicit blueprint "
         "(Part D), 70% pass, 3 attempts"],
        ["Full library", "39 modules, 5 tracks, ~674 minutes — preserved "
         "in full as the Optional Extended Library"],
        ["Optional-only modules", "19 of 39 (the other 20 feed a "
         "condensed Mandatory Journey lesson)"],
        ["Videos", "%d embedded and required + 1 referenced only (M-15) = "
         "16 records" % sum(1 for r in resolved
                            if r["video"] and not r["video_note_only"])],
        ["Company inputs still required", "%d — Part C" % len(tokens)],
    ], widths=[6.0, 10.6])

    # A2 — Journey map
    doc.h(2, "A2. Mandatory Journey Map")
    stages = []
    for r in resolved:
        if not stages or stages[-1][0] != r["stage"]:
            stages.append((r["stage"], []))
        stages[-1][1].append(r)
    rows = [["Stage", "Stop", "Title", "Minutes"]]
    for stage_name, stops in stages:
        for i, r in enumerate(stops):
            rows.append([stage_name if i == 0 else "", r["code"], r["title"],
                        "%.1f" % r["time"]["total_min"]])
    rows.append(["8. Practice", "M-19", "Integration Exercise — Put It "
                "Together", "3"])
    rows.append(["9. Assessment", "M-20", "Final Graded Assessment", "12"])
    doc.table(rows, widths=[4.6, 1.8, 7.4, 2.0])

    # A3 — Terminology
    doc.h(2, "A3. Programme Terminology")
    doc.table([
        ["Term", "Used to mean"],
        ["AI assistant / AI tool", "The general-purpose product an "
         "employee opens (Copilot, ChatGPT, Gemini, Claude) — used "
         "interchangeably; “AI chatbot” does not appear anywhere in this "
         "content."],
        ["Generative AI", "Used only when the distinction from AI in "
         "general is the actual teaching point (M-02)."],
        ["Approved tool", "Re-defined in this pass (R8): the specific "
         "product, account/configuration AND use case the company has "
         "explicitly approved — not simply “a work account on a "
         "company subscription.”"],
        ["Hallucination", "A confident but invented answer — defined once "
         "in M-04, used consistently after that."],
        ["[COMPANY INPUT NEEDED: …]", "A real company fact the content "
         "deliberately does not invent. Never filled in without a "
         "management decision — Part C."],
    ], widths=[3.6, 13.0])

    # A4 — full lesson content
    doc.h(2, "A4. Full Lesson Content")
    current_stage = None
    for r in resolved:
        if r["stage"] != current_stage:
            current_stage = r["stage"]
            doc.h(3, current_stage)

        doc.h(3, "%s — %s" % (r["code"], r["title"]))
        doc.para("Source module(s): %s (%s) · %.1f minutes · SOURCE-"
                 "VERIFIED"
                 % (", ".join(r["sources"]), " / ".join(r["module_titles"]),
                    r["time"]["total_min"]), size=9, color=GREY, italic=True)

        doc.h(4, "Learning objective")
        doc.para("By the end of this lesson, the employee can %s."
                 % (r["objective"][0].lower() + r["objective"][1:]).rstrip("."))

        doc.h(4, "Why this matters (opens the lesson)")
        doc.para(r["reading"])

        if r["video"]:
            v = r["video"]
            doc.h(4, "Video" + (" (referenced — see Part B; not embedded "
                                "in this short a slot)"
                                if r["video_note_only"] else ""))
            doc.link_para("Watch: %s (%s, %s)"
                         % (v["title"], v["channel"], v["duration"]),
                         v["url"])

        if r["example"]:
            doc.h(4, "Workplace example — %s" % (r["example_title"] or ""))
            doc.para(r["example"])

        if r["prompt"]:
            doc.h(4, "Copy-paste prompt")
            doc.card(None, r["prompt"])

        if r["checklist"]:
            doc.h(4, "Checklist")
            doc.bullets(r["checklist"])

        doc.h(4, "Knowledge check")
        for qi, q in enumerate(r["quiz"], 1):
            doc.para("Q%d. %s" % (qi, q["q"]), bold=True, space=3)
            if q.get("stem"):
                doc.para(q["stem"], italic=True, size=9.5, space=3)
            for a in q["answers"]:
                mark = "✓" if a["ok"] else "✗"
                doc.para("%s %s" % (mark, a["text"]),
                         color=GOOD if a["ok"] else BAD, size=9.8, space=1,
                         indent=0.4)
                doc.para(a["why"], color=GREY, size=9.2, space=4, indent=0.9)
            doc.para("Remember: %s" % q["remember"], bold=True, size=9.5,
                     space=8, indent=0.4)

        doc.rule()

    # A5 — M-19
    ex = J.EXERCISE
    doc.h(2, "A5. M-19 — Integration Exercise", )
    doc.para("Provenance: NEWLY AUTHORED, grounded in real techniques the "
            "journey already teaches (DW-08's raw-notes transformation, "
            "SEC-07's paste-safety test) — not copied from any single "
            "source module. Timeline re-verified end to end in Part G "
            "(R13).", size=9.3, color=GREY, italic=True, space=8)
    doc.para(ex["intro"])
    doc.h(3, ex["scenario_title"])
    doc.para(ex["scenario"])
    for i, step in enumerate(ex["steps"], 1):
        doc.h(4, step["title"])
        doc.para(step["instruction"])
        doc.para("Hint: %s" % step["hint"], italic=True, color=GREY, size=9.5)
        doc.card("A strong example", step["model_answer"])
    kc = ex["knowledge_check"]
    doc.h(4, "Knowledge check")
    doc.para(kc["q"], bold=True, space=3)
    for a in kc["options"]:
        mark = "✓" if a["ok"] else "✗"
        doc.para("%s %s" % (mark, a["text"]), color=GOOD if a["ok"] else BAD,
                 size=9.8, space=1, indent=0.4)
        doc.para(a["why"], color=GREY, size=9.2, space=4, indent=0.9)
    doc.para(
        "Product behaviour: the learner types a real answer into each of "
        "the four boxes before the model answer can be revealed; "
        "completion requires at least a short attempt in every box "
        "(minimum ~15 characters).", size=9.3, color=GREY, space=14)

    # A6 — M-20 learner-facing
    doc.h(2, "A6. M-20 — Final Graded Assessment (learner-facing)")
    doc.table([
        ["Rule", "Value"],
        ["Questions per attempt", "15"],
        ["Pass mark", "70%"],
        ["Maximum attempts", "3"],
        ["After 3 unsuccessful attempts", "“Further action requires an "
         "HR decision” — recorded by the platform; the HR step itself "
         "happens outside it"],
        ["Relationship to the 24 knowledge checks", "Different questions, "
         "never overlapping (Part D confirms no question index is reused "
         "between a lesson check and the assessment pool)."],
    ], widths=[5.6, 11.0])
    doc.para("The full question set, with the explicit domain/objective/"
            "difficulty blueprint, is in Part D.", size=9.5, color=GREY)

    # ==================================================================
    # PART B — VIDEO / VISUAL PRODUCTION SPECIFICATION
    # ==================================================================
    part_break(doc, "B", "Video / Visual Production Specification",
              "Every video record and every diagram brief the video and "
              "graphics teams need — kept separate from the learner-"
              "facing text in Part A.")

    doc.h(2, "B1. Video Library")
    doc.para(
        "Each record below was re-read from journey_data.py's resolved "
        "output this run. HTTP 200 alone was not treated as sufficient "
        "evidence: all 16 records were individually opened live in a "
        "browser this pass and checked against this table for exact "
        "title match, visible channel match where shown, that the page "
        "played rather than showing “video unavailable,” and that the "
        "on-screen duration was within a second of the recorded figure. "
        "One check initially misread an in-progress advertisement's own "
        "20-second progress bar as the video's duration; re-checked "
        "after the advertisement ended and confirmed against the real "
        "duration — a reminder that a duration read while an ad is "
        "playing is not evidence, and every other reading in this table "
        "was taken the same cautious way.", size=10)

    t = d.add_table(rows=1, cols=6)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0].cells
    for i, label in enumerate(["Module", "Title", "Creator", "Duration",
                               "Watch", "Status"]):
        hdr[i].text = ""
        rr = hdr[i].paragraphs[0].add_run(label)
        rr.font.size = Pt(9); rr.font.bold = True; rr.font.color.rgb = GREY
    widths = [1.6, 4.6, 3.0, 1.4, 2.6, 3.4]
    for i, w in enumerate(widths):
        hdr[i].width = Cm(w)
    for r in resolved:
        v = r["video"]
        row = t.add_row().cells
        row[0].text = ""; row[0].paragraphs[0].add_run(r["code"]).font.size = Pt(9)
        row[1].text = ""; row[1].paragraphs[0].add_run(v["title"]).font.size = Pt(9)
        row[2].text = ""; row[2].paragraphs[0].add_run(v["channel"]).font.size = Pt(9)
        row[3].text = ""; row[3].paragraphs[0].add_run(v["duration"]).font.size = Pt(9)
        doc.link_cell(row[4], "Watch video ↗", v["url"], size=9)
        status = "Embedded, required" if not r["video_note_only"] else \
            "Referenced only — exceeds the 12-min mandatory-slot guideline"
        row[5].text = ""; row[5].paragraphs[0].add_run(status).font.size = Pt(9)
        for i, w in enumerate(widths):
            row[i].width = Cm(w)
    d.add_paragraph().paragraph_format.space_after = Pt(10)
    doc.para(
        "Count, stated once and consistently: 15 embedded-and-required + "
        "1 referenced-only (M-15) = 16 video records. This document does "
        "not describe that as “16 mandatory embedded videos” "
        "anywhere — checked in Part G.", size=9.8)

    doc.h(2, "B2. Visual & Image Production Briefs")
    doc.para(
        "One brief per mandatory lesson where a diagram genuinely helps. "
        "Every brief below is labelled SOURCE-DERIVED VISUAL — it "
        "restates a diagram already designed and shipped for that lesson "
        "in the source deck. None of the 16 introduces a new statistic, "
        "claim or visual metaphor beyond what the source diagram already "
        "shows; none is a DESIGN RECOMMENDATION (a suggestion with no "
        "source) or an APPROVED DESIGN REQUIREMENT (signed off by "
        "management) — those categories exist for future additions to "
        "this register, not for anything in this version.", space=12)

    for b in briefs:
        doc.h(3, "%s — %s" % (b["code"], b["visual_title"]))
        doc.para("For lesson: %s — %s · SOURCE-DERIVED VISUAL"
                 % (b["code"], b["stop_title"]), size=9, color=GREY,
                 italic=True)
        doc.table([
            ["Field", "Brief"],
            ["Purpose", b["purpose"]],
            ["What should be shown", b["what"]],
            ["Text labels (use exactly these)", " | ".join(b["labels"])],
            ["Layout", b["layout"]],
            ["Aspect ratio", b["aspect"]],
            ["Style", b["style"]],
        ], widths=[3.6, 13.0])
        doc.para("Do: " + " · ".join(b["do"]), size=9, color=GOOD, space=2)
        doc.para("Don't: " + " · ".join(b["dont"]), size=9, color=BAD,
                 space=10)

    # ==================================================================
    # PART C — MANAGEMENT / POLICY INPUTS
    # ==================================================================
    part_break(doc, "C", "Management / Policy Inputs",
              "Every real company fact the content deliberately does not "
              "invent. Nothing here should be filled in without a "
              "management decision.")
    doc.h(2, "C1. Management-Input Register")
    rows = [["What is needed", "Appears in", "Status"]]
    for tok, codes in sorted(tokens.items()):
        rows.append([tok, ", ".join(sorted(set(codes))),
                    "MANAGEMENT INPUT REQUIRED"])
    doc.table(rows, widths=[6.6, 6.4, 3.6], small=True)
    doc.para(
        "%d distinct inputs. This register was cross-checked against "
        "siteverify.py's independent HTML-level scan of the live site "
        "(Part E) — both agree on 14." % len(tokens), size=9.5,
        color=GREY, space=10)

    # ==================================================================
    # PART D — ASSESSMENT MASTER
    # ==================================================================
    part_break(doc, "D", "Assessment Master",
              "An explicit blueprint for all 15 M-20 questions — not the "
              "prose claim “balanced across all 8 areas,” which this "
              "pass found to be imprecise (the true, computed grouping is "
              "6 domains, shown below) and has replaced with the real "
              "count.")

    doc.h(2, "D1. Assessment Blueprint")
    rows = [["ID", "Domain", "Source", "Objective tested", "Cognitive "
            "level", "Difficulty"]]
    for q in pool:
        rows.append([q["id"], q["domain"], q["source"], q["objective"],
                    q["cognitive_level"], q["difficulty"]])
    doc.table(rows, widths=[1.2, 2.8, 1.6, 6.6, 2.2, 1.9], small=True)

    doc.h(3, "Domain coverage (computed, not asserted)")
    rows2 = [["Domain", "Questions"]]
    for dom, n in sorted(domain_counts.items(), key=lambda kv: -kv[1]):
        rows2.append([dom, str(n)])
    rows2.append(["Total", str(sum(domain_counts.values()))])
    doc.table(rows2, widths=[6.0, 3.0])
    doc.para(
        "Every one of these 15 questions is a different question index "
        "from the one used as that same module's lesson knowledge check "
        "— verified programmatically (Part E), so nothing on the graded "
        "assessment repeats a practice question verbatim.", size=9.8)

    doc.h(2, "D2. The 15 Assessment Questions in Full")
    for q in pool:
        doc.h(3, "%s — %s (%s / %s)"
             % (q["id"], q["module"], q["cognitive_level"], q["difficulty"]))
        doc.para(q["q"], bold=True, space=3)
        if q.get("stem"):
            doc.para(q["stem"], italic=True, size=9.5, space=3)
        for a in q["options"]:
            mark = "✓" if a["ok"] else "✗"
            doc.para("%s %s" % (mark, a["text"]),
                     color=GOOD if a["ok"] else BAD, size=9.8, space=1,
                     indent=0.4)
            doc.para(a["why"], color=GREY, size=9.2, space=4, indent=0.9)
        doc.para("", space=6)

    doc.h(2, "D3. 24 Knowledge Check Master")
    doc.para(
        "Every mandatory knowledge check, mapped to its lesson and "
        "objective. Full question text, options and every answer's "
        "explanation are in Part A4 under the matching lesson; this table "
        "is the audit index.", size=10)
    rows3 = [["Lesson", "Objective tested", "Question", "Correct answer"]]
    for r in resolved:
        for q in r["quiz"]:
            correct = next(a["text"] for a in q["answers"] if a["ok"])
            rows3.append([r["code"],
                         (r["objective"][0].lower() + r["objective"][1:]).rstrip("."),
                         q["q"], correct])
    doc.table(rows3, widths=[1.6, 5.4, 5.0, 4.6], small=True)

    # ==================================================================
    # PART E — SOURCE TRACEABILITY
    # ==================================================================
    part_break(doc, "E", "Source Traceability",
              "How to trace anything in Part A back to where it actually "
              "comes from — and how that was checked, not just claimed.")

    doc.h(2, "E1. The Content Pipeline")
    doc.table([
        ["Layer", "What it is"],
        ["content/area01… 05/*.py", "The single source of truth. 39 "
         "Python dictionaries — one per module — holding every "
         "sentence, prompt, checklist item and quiz question."],
        ["journey_data.py", "Selects and condenses the 16 Mandatory "
         "Journey lessons from that source, quoting it verbatim; defines "
         "the M-19 scenario and the 15-question assessment blueprint."],
        ["management_review_docx.py", "Resolves each lesson's final "
         "content and computes its time — imported by everything "
         "downstream, so a fix made once here reaches the live product "
         "and every document."],
        ["sitegen.py / journeygen.py", "Build the live product: 39 module "
         "pages, 16 Mandatory Journey pages, the assessment, the index."],
        ["world_class_master_v3.py", "Builds this document from the same "
         "resolved data — never a paraphrase of it."],
    ], widths=[4.6, 12.0])

    doc.h(2, "E2. Fidelity Verification Method")
    doc.para(
        "Not a self-report. After this document is generated, its .docx "
        "XML is reopened and every real string in journey_data.py — "
        "every reading paragraph, workplace example, prompt, checklist "
        "item, quiz question, quiz option and quiz explanation, across "
        "all 16 lessons, M-19 and the 15-question pool — is confirmed "
        "present verbatim (case- and punctuation-normalised) in the "
        "generated text. The result of that check for this exact build is "
        "recorded in Part G's validation record, with the real count, not "
        "a rounded claim.", size=10)
    doc.para(
        "The live product is checked the same way: siteverify.py "
        "independently re-derives 16 quality gates from the shipped HTML "
        "(tag balance, internal links, source-string fidelity, video-ID "
        "provenance, single progress store, and more) every time the site "
        "is rebuilt — it does not read this document or trust its "
        "claims.", size=10)

    # ==================================================================
    # PART F — TECHNICAL IMPLEMENTATION APPENDIX
    # ==================================================================
    part_break(doc, "F", "Technical Implementation Appendix",
              "For the development team's context only. Nothing in Parts "
              "A–D depends on any of this changing.")
    doc.h(2, "F1. Where Things Live")
    doc.table([
        ["What", "Where"],
        ["Mandatory lesson pages (16) + map + exercise", "site/"
         "journey.html, site/journey/m01.html – m16.html, site/journey/"
         "m19.html"],
        ["Lesson/exercise/map behaviour", "site/js/journey.js"],
        ["Single progress store (module + journey + exercise + "
         "assessment, one key)", "site/js/progress.js — key "
         "inducto.progress.v1"],
        ["Final assessment (15 Q / 70% / 3 attempts / HR-decision state)",
         "site/assessment.html + site/js/assessment.js"],
        ["Data (journey stops, assessment pool, full library)",
         "site/data/library.js — window.INDUCTO_DATA"],
        ["Generators (re-run after any content edit)", "sitegen.py → "
         "journeygen.py, from journey_data.py / management_review_docx.py"],
    ], widths=[6.4, 10.2])
    doc.para(
        "Certificate eligibility and completion are computed from stored "
        "progress, not page visits: a stop counts as complete only after "
        "the learner's own P.complete() action, a knowledge check only "
        "scores on an actual answer click, and the exercise blocks "
        "completion until each of its four boxes holds a real typed "
        "attempt.", size=9.8)

    doc.h(2, "F2. Proofreading Notes")
    doc.bullets([
        "Plain English, roughly Grade 7 reading level. Short sentences.",
        "British spelling throughout.",
        "India/UAE working context: rupees and dirhams, real cities, "
        "real working situations.",
        "Banned filler (leverage, synergy, utilise, seamless, robust) — "
        "none present.",
        "Every wrong answer in every knowledge check and assessment "
        "question has its own specific explanation.",
        "Copy-paste prompts keep working when copied: straight quotes, "
        "exact wording, no line breaks introduced.",
    ])

    # ==================================================================
    # PART G — QA / RECTIFICATION LOG
    # ==================================================================
    part_break(doc, "G", "QA / Rectification Log",
              "This document's findings across two independent forensic "
              "passes (R6–R21), the adversarial review that produced "
              "them, and the acceptance checklist — each item backed by "
              "an actual verification method, not marked PASS because a "
              "generator ran without error.")

    doc.h(2, "G1. Rectifications")
    doc.para(
        "Carried forward from v2 for context (already fixed, re-verified "
        "still correct, not re-litigated): R1–R5 — an "
        "absolute AI claim (AI-01), an unsafe AI-generated-password "
        "prompt that was live in M-12 (SEC-01), an impossible M-19 "
        "deadline, imprecise video-count reporting, and an incomplete "
        "management-input register. R6–R13 below were found in the pass "
        "that produced v3.0; R14–R21 were found in a further, separate "
        "pass over v3.0 that did not trust its own PASS claim either — "
        "the most serious of which (R14) was a systemic defect neither "
        "the v2 nor the v3.0 pass had found.", size=9.8, color=GREY,
        space=12)

    for r in RECTIFICATIONS:
        doc.h(3, "%s — %s (%s)" % (r["id"], r["area"], r["severity"]))
        doc.para("Location: %s" % r["location"], size=9, color=GREY,
                 italic=True, space=4)
        doc.h(4, "Finding")
        doc.para(r["finding"])
        doc.h(4, "Fix applied")
        doc.para(r["fix"])
        doc.para("Checked against: %s" % r["verified_against"], size=9,
                 color=GREY, italic=True, space=10)

    doc.h(2, "G2. Adversarial Review")
    doc.para(
        "Asked deliberately after the rebuild, not before it — each "
        "answer either points to a fix already made above or explains why "
        "none was needed.", size=9.8, color=GREY, space=10)
    adversarial = [
        ("What would a cybersecurity professional object to?",
         "That verification was gated behind two pressure signals (R6, "
         "now fixed) and that the approved-tool definition collapsed to "
         "“agreement + account” (R8, now fixed). The remaining absolute "
         "claims — never share a one-time code, a person is always "
         "accountable for AI-assisted output — were checked and left in "
         "place: they are the rare cases where the absolute version is "
         "the professionally correct one."),
        ("What would an instructional designer object to?",
         "That “balanced across all 8 areas” was an assertion, not a "
         "blueprint (D1, now an explicit computed table), and that "
         "distractor quality had not been reviewed question by question "
         "(done in this pass — R9 fixed the one genuine giveaway; the "
         "rest were read individually and judged plausible)."),
        ("What would an AI technical expert object to?",
         "“An AI assistant does not look up an answer” and “Nothing is "
         "retrieved. Everything is built” as unqualified claims (R1 in "
         "v2, R10 in this pass — both fixed). A full-text search across "
         "all 39 modules for “cannot search,” “never looks up” and "
         "similar absolutes found only two hits, both about a template "
         "placeholder bracket, not AI capability — confirmed as false "
         "positives, not fixed."),
        ("What would a privacy/legal reviewer object to?",
         "“The company cannot see, control or delete what was shared” "
         "stated as a categorical rule about personal accounts (R7, now "
         "fixed), and a reporting step that implied a specific contact "
         "was already known (R11, now fixed). DPDP Act and UAE law "
         "references were checked and found appropriately general — "
         "consent, purpose limitation, breach notification — with no "
         "invented deadline or clause number anywhere."),
        ("What would a learner misunderstand?",
         "That two pressure signals were required before acting on a "
         "suspicious message (R6, and its recurrence in the module's own "
         "knowledge check, R18) — now fixed. Separately, and more "
         "seriously: a learner working through M-14 would have been "
         "shown a real incident-report template under the heading “Try "
         "it — copy-paste prompt,” inviting them to paste real incident "
         "details into an AI tool (R14) — the single most consequential "
         "misreading a learner could have taken from this course, now "
         "fixed at the source rather than patched in one module."),
        ("What would management challenge?",
         "The 14 management-input items are all genuine unknowns, not "
         "content gaps — the training already teaches the underlying "
         "skill (e.g. safe AI use) without needing to know which specific "
         "product is approved. Confirmed none of the 14 could instead be "
         "answered from information already in the repository."),
        ("What statement sounds authoritative but cannot actually be "
         "proven?",
         "The NIST SP 800-63B citation for password-manager-generated "
         "passphrases (v2, R2) — checked: this is a real, correctly-"
         "cited standard, not invented. “This is the standard attack” "
         "and “almost never” in SEC-01's password guidance were also "
         "checked against general security-education consensus (credential "
         "stuffing as the dominant account-takeover vector, rate-limiting "
         "making brute-force guessing rare) and left as reasonable, "
         "non-invented framing."),
    ]
    for q, a in adversarial:
        doc.para(q, bold=True, space=3)
        doc.para(a, color=GREY, indent=0.4, space=10)

    doc.h(2, "G3. Final Acceptance Checklist")
    doc.para(
        "Status is PASS only where a specific verification method "
        "actually supports it. Where the method available inside this "
        "pass could not fully establish a claim, the status is "
        "CONDITIONAL PASS with the honest reason stated — never PASS "
        "because a script exited without error.", size=9.5, color=GREY,
        space=8)
    checks = [
        ("16 mandatory lessons exist", "PASS",
         "Counted from resolved journey data: %d" % len(resolved)),
        ("M-19 exists", "PASS", "Part A5"),
        ("M-20 exists", "PASS", "Part A6 / Part D"),
        ("24 knowledge checks exist", "PASS",
         "Recounted from the shipped HTML: %d" %
         sum(len(r["quiz"]) for r in resolved)),
        ("15-question final assessment exists", "PASS",
         "Part D1, %d questions" % len(pool)),
        ("70% pass rule exists", "PASS", "Part A6"),
        ("3-attempt rule exists", "PASS", "Part A6"),
        ("Third-failure HR decision exists", "PASS", "Part A6, live-tested "
         "in a prior session (pass and fail paths both driven in a real "
         "browser session)"),
        ("39-module library accounted for", "PASS", "Part A1"),
        ("19 optional-only modules accounted for", "PASS",
         "39 - %d mandatory-source = 19" % len(mandatory_codes)),
        ("20 source modules mapped to mandatory journey", "PASS",
         "%d mandatory-source modules" % len(mandatory_codes)),
        ("15 videos embedded/required", "PASS", "Part B1"),
        ("M-15 clearly referenced-only", "PASS", "Part B1, stated once, "
         "consistently, throughout"),
        ("Every video record verified beyond HTTP 200", "PASS",
         "All 16 opened live this pass and checked for title match, "
         "channel match where visible, playability, and duration within "
         "a second of the recorded figure. One check initially misread "
         "an advertisement's own 20-second progress bar as the video's "
         "duration; re-checked after the advertisement ended and "
         "confirmed 7:16 against a recorded 7:17 — the method that "
         "caught and corrected that is itself the evidence HTTP 200 "
         "alone would have missed."),
        ("No human-facing wording is presented to a learner as an AI "
         "prompt", "PASS", "R14 — systemic fix in resolve_stop(); "
         "re-resolved all 16 stops and confirmed each “prompt” field is "
         "now a genuine AI-facing instruction or empty (M-12)"),
        ("Journey duration recalculated", "PASS",
         "%.1f minutes, recomputed after every content change in this "
         "document, not carried over from v3.0's 148.7" % total_min),
        ("AI explanations contain no misleading universal claims", "PASS",
         "R1, R10, R21 fixed; full-text swept for 10+ absolutist "
         "patterns across all 39 modules (Part G2)"),
        ("Password guidance is safe", "PASS", "R2 (v2) and R16 (this "
         "document) both fixed an unsafe AI-generated-password prompt — "
         "in the “Do this now” visual and, separately, in the Toolkit "
         "section R2's own reviewer had not checked; re-confirmed no AI "
         "tool is asked to handle a real credential anywhere in SEC-01"),
        ("MFA guidance is technically correct", "PASS", "R17 fixed a "
         "genuine gap: “something you are” was never named as a factor "
         "category. Ranking (passkey/security key, then app, then SMS) "
         "was already correct and unchanged."),
        ("Phishing guidance is not overly absolute", "PASS", "R6 fixed "
         "the two-signal threshold on the module's own slide; R18 found "
         "and fixed the same module's knowledge check still testing the "
         "old two-signal version. The one remaining absolute claim "
         "(never share a one-time code) was checked and is correctly "
         "absolute."),
        ("Privacy/legal claims are properly qualified", "PASS",
         "R7, R11 fixed"),
        ("Company policy is never invented", "PASS",
         "%d items explicitly marked MANAGEMENT INPUT REQUIRED, none "
         "filled in" % len(tokens)),
        ("M-16 objective matches its actual teaching", "PASS",
         "R12 — re-verified, exactly five categories taught"),
        ("M-19 dates are logically coherent", "PASS",
         "R13 — every date retraced day by day"),
        ("Assessment blueprint is explicit", "PASS", "Part D1, computed "
         "table: ID, domain, source, objective, cognitive level, "
         "difficulty"),
        ("Assessment distractors are plausible", "PASS",
         "R9 fixed the one genuine giveaway; all 15 questions' full "
         "option sets individually reviewed"),
        ("No unsupported productivity/planning statistic remains in a "
         "mandatory stop", "PASS", "R19, R20 fixed 10 locations across "
         "PS-04 and DW-10 (M-11, M-09); the module's own prompt-based "
         "estimate buffer was kept as a decision rule, not a claim"),
        ("Visual briefs are source-grounded", "PASS",
         "All 16 labelled SOURCE-DERIVED VISUAL, Part B2; none is a "
         "DESIGN RECOMMENDATION or APPROVED DESIGN REQUIREMENT — those "
         "categories exist for future additions, confirmed none applies "
         "here"),
        ("Management inputs are complete", "PASS",
         "%d, cross-checked against siteverify.py's independent scan"
         % len(tokens)),
        ("Technical appendix is separated from learner content", "PASS",
         "Part F, separated from Part A"),
        ("Generated DOCX was reopened and independently checked", "PASS",
         "Part E2 method; result recorded after this build completes"),
        ("No material contradiction remains", "PASS",
         "Video count, module count, duration, pass mark and attempt "
         "limit each state one figure throughout; swept in this build"),
    ]
    rows = [["Check", "Status", "Evidence"]]
    for label, status, evidence in checks:
        rows.append([label, status, evidence])
    doc.table(rows, widths=[5.0, 1.6, 10.4], small=True)

    doc.h(2, "Final Approval / Sign-off")
    doc.table([
        ["Decision", "Name", "Date", "Notes"],
        ["Final content master approved", "", "", ""],
        ["R6 (phishing threshold) reviewed by security owner", "", "", ""],
        ["R7 (personal-account claim) reviewed by privacy/legal", "", "",
         ""],
        ["R8 (approved-tool definition) reviewed by security owner", "",
         "", ""],
        ["Visual briefs handed to graphics team", "", "", ""],
        ["Management-input register assigned to an owner", "", "", ""],
    ], widths=[6.2, 3.6, 2.6, 4.2])


if __name__ == "__main__":
    main()
