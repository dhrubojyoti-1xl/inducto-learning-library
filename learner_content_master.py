# -*- coding: utf-8 -*-
"""
INDUCTO Learning & Knowledge — New Employee Learning Journey.

    python learner_content_master.py

This is the LEARNER CONTENT MASTER, not an audit. It answers one question:
what does a new employee actually see and do, lesson by lesson, in the
Inducto Learning & Knowledge module?

Every word of teaching content, every workplace example, every "try it"
prompt, every recap point and every knowledge-check question is pulled
verbatim from the same source dictionaries that build the live product
(content/areaNN/*.py -> journey_data.py / management_review_docx.py ->
journeygen.py). Nothing here is invented, reordered or renamed against
that source. Where a real company-specific fact is required and the
source does not supply one, the text says so plainly with
[COMPANY INPUT NEEDED: ...], exactly as the live product does.

Document shape, per lesson (matches the live Mandatory Journey page, in
the same order a learner actually sees it):

    What is this lesson about?  -> a 2-sentence hook (deck subtitle +
                                    the lesson's own first real outcome)
    What you'll learn           -> the module's own outcomes, verbatim
    Watch / Read                -> the real video record
    Learn                       -> the module's own "why this matters"
                                    story + its real handling rules
    Workplace example           -> the module's own scenario: situation,
                                    the real "good" choice, and why
    Try it                      -> the module's own AI-safe prompt,
                                    where the lesson has one
    Remember                    -> the module's own recap points
    Knowledge check             -> the lesson's real knowledge-check
                                    question(s), verbatim, with every
                                    answer's real explanation

Then M-19 (Integration Exercise) and M-20 (Final Assessment, described
only — not the question bank, so this document cannot be used to see
the graded answers in advance), then two short appendices.
"""

import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
sys.stdout.reconfigure(encoding="utf-8")

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm, Pt

import journey_data as J
import sitegen
import world_class_master_docx as W  # Doc, add_hyperlink, colours

ROOT = sitegen.ROOT
OUT_DOCX = os.path.join(ROOT, "docs",
                        "INDUCTO_World_Class_Content_Master_FINAL.docx")

INK, GREY, GOOD, BAD, ACCENT, GOLD = W.INK, W.GREY, W.GOOD, W.BAD, W.ACCENT, W.GOLD
Doc = W.Doc


def lc1(s):
    """Lower-case the first letter, for splicing a source phrase mid-sentence."""
    return (s[0].lower() + s[1:]) if s else s


def good_choice(scenario):
    for c in scenario.get("choices", []):
        if c.get("tone") == "good":
            return c
    return scenario.get("choices", [{}])[0] if scenario.get("choices") else {}


# ---------------------------------------------------------------------------
def build_lesson(doc, r, by_code, lesson_number, total_lessons):
    primary = by_code[r["sources"][0]]
    scenario = primary.get("scenario") or {}
    good = good_choice(scenario)
    recap_points = primary.get("recap", {}).get("points", [])[:5]
    outcomes = [t for _, t in primary.get("outcomes", [])][:5]
    video = r.get("video")

    doc.h(1, "%s — %s" % (r["code"], r["title"]), page_break=True)
    doc.para("%s · %s minutes · lesson %d of %d in the mandatory journey"
             % (r["stage"], r["time"]["total_min"], lesson_number,
                total_lessons), size=9.5, color=GREY, space=10)

    # ---- 1. What is this lesson about? ----
    doc.h(2, "What is this lesson about?")
    hook = primary.get("subtitle", "")
    if hook and not hook.endswith((".", "!", "?")):
        hook += "."
    lead_outcome = outcomes[0] if outcomes else ""
    doc.para(
        "%s%s" % (hook, (" By the end of this lesson, you can %s."
                        % lc1(lead_outcome)) if lead_outcome else ""),
        space=10)

    # ---- 2. What you'll learn ----
    doc.h(2, "What you'll learn")
    doc.para("By the end of this lesson, you can:", size=10, space=4)
    doc.bullets(outcomes)

    # ---- 3. Watch / Read ----
    doc.h(2, "Watch" if video else "Watch / Read")
    if video:
        if r.get("video_note_only"):
            doc.para(
                "This video is background material — useful, but not "
                "required to complete this lesson.", italic=True,
                color=GREY, size=9.8, space=4)
        else:
            doc.para("Watch this before you continue.", size=10, space=4)
        p = doc.d.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r_run = p.add_run(video["title"])
        r_run.font.bold = True
        r_run.font.size = Pt(10.5)
        doc.para("%s · %s" % (video["channel"], video["duration"]),
                 size=9.5, color=GREY, space=4)
        doc.link_para(video["url"], video["url"], size=9.5)
    else:
        doc.para("No video for this lesson — go straight to Learn below.",
                 italic=True, color=GREY, size=9.8, space=8)

    # ---- 4. Learn ----
    doc.h(2, "Learn")
    doc.para(r["reading"], space=8)
    if r.get("checklist"):
        doc.para("The rule, in practice:", bold=True, size=10, space=4)
        doc.bullets(r["checklist"])

    # ---- 5. Workplace example ----
    doc.h(2, "Workplace example")
    if scenario.get("situation"):
        doc.card("Situation", scenario["situation"], tone=None)
        if good.get("text"):
            doc.card("What you should do", good["text"], tone="good")
        why = good.get("consequence") or good.get("rule")
        if why:
            doc.card("Why", why, tone=None)
    elif r.get("example"):
        doc.para(r["example"])

    # ---- 6. Try it ----
    if r.get("prompt"):
        doc.h(2, "Try it")
        doc.para("Try this yourself, with your own real task in place of "
                 "the example:", size=10, space=4)
        doc.card("Prompt", r["prompt"], tone=None, accent=ACCENT)

    # ---- 7. Remember ----
    if recap_points:
        doc.h(2, "Remember")
        doc.bullets(["%s — %s" % (h, t) for h, t in recap_points])

    # ---- 8. Knowledge check ----
    doc.h(2, "Knowledge check")
    doc.para(
        "Practice only — this does not count towards your final "
        "assessment score, and you can try again as many times as you "
        "like.", italic=True, size=9.5, color=GREY, space=6)
    for qi, q in enumerate(r["quiz"], 1):
        doc.para("%d. %s" % (qi, q["q"]), bold=True, size=10.5, space=4)
        letters = "ABCD"
        for li, a in enumerate(q["answers"]):
            tag = " — CORRECT" if a["ok"] else ""
            doc.para("%s. %s%s" % (letters[li], a["text"], tag),
                     size=10, bold=a["ok"], color=GOOD if a["ok"] else INK,
                     space=2, indent=0.4)
            doc.para(a["why"], size=9.3, italic=True, color=GREY, space=4,
                     indent=0.8)

    doc.para(
        "Mark this lesson complete, then continue to the next lesson in "
        "the journey.", size=9.5, color=GREY, space=14)


# ---------------------------------------------------------------------------
def build_exercise(doc, ex):
    doc.h(1, "%s — %s" % (ex["code"], ex["title"]), page_break=True)
    doc.para(ex["stage"], size=9.5, color=GREY, space=10)
    doc.para(ex["intro"], space=10)

    doc.h(2, ex["scenario_title"])
    doc.card(None, ex["scenario"], tone=None)

    for i, step in enumerate(ex["steps"], 1):
        doc.h(2, step["title"])
        doc.para(step["instruction"], space=4)
        doc.para("Hint: %s" % step["hint"], italic=True, size=9.5,
                 color=GREY, space=6)
        doc.card("A strong example answer", step["model_answer"],
                 tone="good")

    doc.h(2, "Knowledge check")
    kc = ex["knowledge_check"]
    doc.para(kc["q"], bold=True, size=10.5, space=4)
    letters = "ABCD"
    for li, o in enumerate(kc["options"], 0):
        tag = " — CORRECT" if o["ok"] else ""
        doc.para("%s. %s%s" % (letters[li], o["text"], tag), size=10,
                 bold=o["ok"], color=GOOD if o["ok"] else INK, space=2,
                 indent=0.4)
        doc.para(o["why"], size=9.3, italic=True, color=GREY, space=4,
                 indent=0.8)

    doc.para(
        "Mark this exercise complete, then continue to the Final "
        "Assessment.", size=9.5, color=GREY, space=14)


# ---------------------------------------------------------------------------
def build_final_assessment(doc):
    doc.h(1, "M-20 — Final Assessment", page_break=True)
    doc.para(
        "You have completed the learning journey. The final assessment "
        "contains:", space=6)
    doc.bullets([
        "15 questions",
        "Pass mark: 70%",
        "Maximum attempts: 3",
    ])
    doc.para(
        "After three unsuccessful attempts: “Further action requires "
        "an HR decision.”", space=10)
    doc.para(
        "The final assessment draws on everything taught across the 16 "
        "lessons above. No question repeats a lesson's own knowledge "
        "check word for word, and none of its answers are printed in "
        "this document.", size=9.8, color=GREY)


def main():
    by_code, decks, resolved = J.load()
    total_min = sum(r["time"]["total_min"] for r in resolved) + 3 + 12
    pool = J.assessment_pool(by_code)

    # Internal verification (not printed in the learner document): the
    # final-assessment pool must not repeat a question already used as a
    # lesson knowledge check, and every pool question needs a real,
    # single correct answer.
    used_lesson_q = set()
    for r in resolved:
        primary = by_code[r["sources"][0]]
        for q in r["quiz"]:
            used_lesson_q.add((primary["module_code"], q["q"]))
    overlap = [q for q in pool if (q["source"], q["q"]) in used_lesson_q]
    for q in pool:
        n_ok = sum(1 for o in q["options"] if o["ok"])
        assert n_ok == 1, "assessment question %s has %d correct answers" % (
            q["id"], n_ok)
    print("Internal check: %d/%d assessment questions overlap a lesson "
          "knowledge check (expect 0)." % (len(overlap), len(pool)))

    tokens = {}
    import re
    import json
    for dk in decks:
        for mm in re.finditer(r"\[COMPANY INPUT NEEDED:([^\]]*)\]",
                              json.dumps(dk)):
            tokens.setdefault(mm.group(1).strip(), []).append(dk["module_code"])
    tokens.setdefault("name of the team that owns this training",
                      []).append("site-wide footer, every page")

    doc = Doc()
    # The base Doc class's footer says "Management Review" — wrong branding
    # for the learner content master; retitle it, keep the page-number field.
    doc.d.sections[0].footer.paragraphs[0].runs[0].text = (
        "Inducto Learning & Knowledge — New Employee Learning Journey · page ")
    build(doc, by_code, resolved, total_min, tokens)
    doc.d.save(OUT_DOCX)

    n_checks = sum(len(r["quiz"]) for r in resolved)
    print("Saved %s" % OUT_DOCX)
    print("  %d lessons, %.1f minutes, %d lesson-check questions, "
          "%d final-assessment questions, %d company inputs open"
          % (len(resolved), total_min, n_checks, len(pool), len(tokens)))
    return doc, resolved, pool, tokens


def build(doc, by_code, resolved, total_min, tokens):
    d = doc.d

    # ================================================================
    # COVER
    # ================================================================
    doc.h(1, "INDUCTO")
    doc.para("Learning & Knowledge", size=16, color=ACCENT, space=2)
    doc.para("New Employee Learning Journey", size=13, color=GREY,
             space=18)
    doc.rule()
    doc.para("Prepared for: Learning & Development · Content · Instructional "
             "Design · New employees", size=9.5, color=GREY)
    doc.para("Prepared by: Learning & Development — Dhrubojyoti "
             "(chetan@1xl.com)", size=9.5, color=GREY)
    doc.para("%d lessons · %.1f minutes · 5 September 2026" %
             (len(resolved), total_min), size=9.5, color=GREY, space=16)
    doc.para(
        "This document is the actual content of the Inducto Learning & "
        "Knowledge module's mandatory journey — what a new employee "
        "reads, watches and answers, lesson by lesson — drawn directly "
        "from the same source that builds the live product.", size=10.5)

    # ================================================================
    # 1. HOW THE JOURNEY WORKS
    # ================================================================
    doc.h(1, "1. How the Journey Works", page_break=True)
    doc.para(
        "Each lesson follows the same simple pattern, in this order:",
        space=6)
    doc.numbered([
        "Learn what the topic is and why it matters.",
        "Watch the lesson's video, or read its background note if it "
        "does not have one to watch.",
        "See a real workplace example: a situation, what to do, and why.",
        "Try it yourself, where a lesson has a practice activity.",
        "Answer a small knowledge check and see the explanation for "
        "each answer.",
        "Mark the lesson complete and move to the next one.",
    ])
    doc.para(
        "There are two different kinds of question in this journey, and "
        "they are not the same thing:", space=10, bold=True)
    doc.table([
        ["", "Lesson knowledge check", "Final assessment (M-20)"],
        ["When", "At the end of each of the 16 lessons", "After the last "
         "lesson and the integration exercise"],
        ["How many", "24 in total across the journey", "15, in one "
         "attempt"],
        ["Stakes", "Practice only — does not count towards your score",
         "Graded — needs a 70% pass mark"],
        ["Retries", "As many as you like", "Up to 3 attempts"],
        ["If you don't pass", "N/A — try again any time", "“Further "
         "action requires an HR decision” after the third "
         "unsuccessful attempt"],
    ], widths=[2.6, 6.8, 6.8])

    # ================================================================
    # 2. MANDATORY LEARNING JOURNEY
    # ================================================================
    doc.h(1, "2. Mandatory Learning Journey", page_break=True)
    doc.para(
        "The full Inducto library holds 39 modules across five tracks. "
        "This mandatory journey selects the required ones — some "
        "combined into a single stop — into the 16 lessons below, plus "
        "an integration exercise and a final assessment. Every other "
        "module in the library is an optional deep-dive, not part of "
        "the required path.", space=10)
    rows = [["Stage", "Code", "Lesson", "Minutes"]]
    for r in resolved:
        rows.append([r["stage"], r["code"], r["title"],
                    "%.1f" % r["time"]["total_min"]])
    rows.append(["8. Practice", "M-19", "Integration Exercise — Put It "
                "Together", "3"])
    rows.append(["9. Assessment", "M-20", "Final Assessment", "12"])
    doc.table(rows, widths=[3.6, 1.6, 8.2, 2.8])

    # ================================================================
    # 3-18. THE 16 LESSONS
    # ================================================================
    for i, r in enumerate(resolved, 1):
        build_lesson(doc, r, by_code, i, len(resolved))

    # ================================================================
    # M-19 / M-20
    # ================================================================
    build_exercise(doc, J.EXERCISE)
    build_final_assessment(doc)

    # ================================================================
    # APPENDIX A — LEARNING RESOURCES
    # ================================================================
    doc.h(1, "Appendix A — Learning Resources", page_break=True)
    doc.para(
        "Every video used in the mandatory journey. All are outside "
        "material, not company-produced — where a video differs from "
        "its lesson, the lesson is correct.", size=9.8, color=GREY,
        space=10)
    t = d.add_table(rows=1, cols=5)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0].cells
    for j, h in enumerate(["Lesson", "Title", "Creator", "Duration",
                          "Link"]):
        hdr[j].text = h
        hdr[j].paragraphs[0].runs[0].font.bold = True
        hdr[j].paragraphs[0].runs[0].font.size = Pt(9)
    for r in resolved:
        v = r.get("video")
        if not v:
            continue
        row = t.add_row().cells
        row[0].text = r["code"] + (" (referenced only)"
                                   if r.get("video_note_only") else "")
        row[1].text = v["title"]
        row[2].text = v["channel"]
        row[3].text = v["duration"]
        doc.link_cell(row[4], "Open", v["url"], size=9)
        for c in row[:4]:
            c.paragraphs[0].runs[0].font.size = Pt(9)
    d.add_paragraph().paragraph_format.space_after = Pt(4)
    for j, w in enumerate([2.0, 5.5, 3.4, 1.8, 1.5]):
        for row in t.rows:
            row.cells[j].width = Cm(w)

    # ================================================================
    # APPENDIX B — COMPANY INPUTS NEEDED
    # ================================================================
    doc.h(1, "Appendix B — Company Inputs Needed", page_break=True)
    doc.para(
        "This training deliberately does not invent company policy. "
        "Every fact below is a genuine, real gap this training depends "
        "on — none of it should be filled in without a management "
        "decision.", size=9.8, color=GREY, space=10)
    rows = [["What is needed", "Appears in"]]
    for tok, codes in sorted(tokens.items()):
        rows.append([tok, ", ".join(sorted(set(codes)))])
    doc.table(rows, widths=[7.0, 9.4], small=True)

    # ================================================================
    # SOURCE NOTE
    # ================================================================
    doc.h(1, "A note on where this content comes from", page_break=True)
    doc.para(
        "This document is generated directly from the Inducto source "
        "repository (content/area01-05, journey_data.py, "
        "management_review_docx.py) — the same data that builds the "
        "live Learning & Knowledge product, the PowerPoint library and "
        "this journey's own knowledge checks. It is not a separate, "
        "hand-written interpretation: a fix to a lesson in the source "
        "changes this document and the live product identically the "
        "next time both are generated.", size=9.8, color=GREY)


if __name__ == "__main__":
    main()
