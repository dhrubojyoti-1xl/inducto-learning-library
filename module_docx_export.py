# -*- coding: utf-8 -*-
"""
Export each module's full teaching content as its own clean Word document,
ready to upload or paste into Inducto (the OneWork admin console) or hand to
a proof reader.

    python module_docx_export.py

One .docx per module — 39 files — written under docs/upload-content/, plus
a manifest.json listing every file. Unlike the combined content blueprint,
these carry no reference codes: they read as a finished document, not a
markup sheet.

The text comes from content/areaNN/*.py, the same dictionaries that build
the PowerPoint decks and the web pages, so nothing here is retyped.
"""

import importlib
import io
import json
import os
import re
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

import build
import sitegen
import theme as T

ROOT = sitegen.ROOT
OUTDIR = os.path.join(ROOT, "docs", "upload-content")

INK = RGBColor(0x10, 0x18, 0x26)
GREY = RGBColor(0x5B, 0x67, 0x79)
GOOD = RGBColor(0x1B, 0x7F, 0x4B)
BAD = RGBColor(0xC6, 0x28, 0x28)

TRACK_ACCENT = {
    "01-ai-general": RGBColor(0x2F, 0x4B, 0xC4),
    "02-ai-daily-work": RGBColor(0x0E, 0x6E, 0x75),
    "03-prompt-engineering": RGBColor(0x6A, 0x2F, 0xA0),
    "04-professional-skills": RGBColor(0x8A, 0x5A, 0x00),
    "05-security-privacy": RGBColor(0x7A, 0x12, 0x20),
}

AREA_DIR = {
    "01-ai-general": "01-ai-courses-general",
    "02-ai-daily-work": "02-ai-for-day-to-day-work",
    "03-prompt-engineering": "03-prompt-engineering",
    "04-professional-skills": "04-professional-skills",
    "05-security-privacy": "05-cybersecurity-and-data-privacy",
}


def slug(code):
    return code.lower().replace("-", "")


def filename(d):
    title = re.sub(r"[^A-Za-z0-9]+", "-", d["title"]).strip("-")
    return "%s - %s.docx" % (d["module_code"], title)


# ---------------------------------------------------------------------------
class Doc(object):

    def __init__(self, accent):
        self.d = Document()
        self.accent = accent
        self._page()
        self._styles()

    def _page(self):
        s = self.d.sections[0]
        s.page_width, s.page_height = Cm(21.0), Cm(29.7)
        s.left_margin, s.right_margin = Cm(2.4), Cm(2.4)
        s.top_margin, s.bottom_margin = Cm(2.0), Cm(2.0)
        p = s.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("Inducto Learning & Knowledge Library · page ")
        r.font.size, r.font.color.rgb = Pt(8), GREY
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        p._p.append(fld)

    def _styles(self):
        st = self.d.styles
        n = st["Normal"]
        n.font.name, n.font.size, n.font.color.rgb = "Segoe UI", Pt(11), INK
        n.paragraph_format.space_after = Pt(8)
        n.paragraph_format.line_spacing = 1.2
        for name, size, col, before in (("Heading 1", 24, self.accent, 0),
                                        ("Heading 2", 15, INK, 22),
                                        ("Heading 3", 12, INK, 12)):
            s = st[name]
            s.font.name = "Segoe UI Semibold"
            s.font.size, s.font.color.rgb, s.font.bold = Pt(size), col, True
            s.paragraph_format.space_before = Pt(before)
            s.paragraph_format.space_after = Pt(6)
            s.paragraph_format.keep_with_next = True

    def h(self, level, text, color=None):
        p = self.d.add_paragraph(style="Heading %d" % level)
        r = p.add_run(text)
        if color is not None:
            r.font.color.rgb = color
        return p

    def para(self, text, size=11, color=None, italic=False, bold=False,
             space=8, indent=0.0, style=None):
        p = self.d.add_paragraph(style=style) if style else self.d.add_paragraph()
        p.paragraph_format.space_after = Pt(space)
        if indent:
            p.paragraph_format.left_indent = Cm(indent)
        r = p.add_run(text)
        r.font.size, r.font.italic, r.font.bold = Pt(size), italic, bold
        r.font.color.rgb = color if color is not None else INK
        return p

    def bullets(self, items, tone=None):
        for it in items:
            p = self.d.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(it)
            r.font.size = Pt(11)
            if tone:
                r.font.color.rgb = {"good": GOOD, "bad": BAD}[tone]

    def numbered(self, items):
        for it in items:
            p = self.d.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(4)
            p.add_run(it).font.size = Pt(11)

    def rule(self):
        p = self.d.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(10)
        pbdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
        bot.set(qn("w:color"), "D6DCE7")
        pbdr.append(bot)
        p._p.get_or_add_pPr().append(pbdr)

    def table(self, rows, widths=None):
        t = self.d.add_table(rows=0, cols=len(rows[0]))
        t.style = "Light Grid Accent 1"
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i, row in enumerate(rows):
            cells = t.add_row().cells
            for j, val in enumerate(row):
                cells[j].text = ""
                p = cells[j].paragraphs[0]
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(str(val))
                r.font.size = Pt(9.5)
                r.font.bold = (i == 0)
                if widths:
                    cells[j].width = Cm(widths[j])
        self.d.add_paragraph().paragraph_format.space_after = Pt(4)
        return t

    def card(self, title, text, tone=None):
        """A single-cell shaded box for a callout (cost / fix / rule)."""
        t = self.d.add_table(rows=1, cols=1)
        cell = t.rows[0].cells[0]
        shade = {"bad": "FDF3F3", "good": "F1F8F4"}.get(tone, "F3F6FB")
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), shade)
        tcPr.append(shd)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        if title:
            r = p.add_run(title.upper() + "\n")
            r.font.size, r.font.bold = Pt(8.5), True
            r.font.color.rgb = {"bad": BAD, "good": GOOD}.get(tone, self.accent)
        r = p.add_run(text)
        r.font.size = Pt(10.5)
        self.d.add_paragraph().paragraph_format.space_after = Pt(2)


# ---------------------------------------------------------------------------
def emit_visual(doc, v):
    t = v["type"]

    if t == "flow":
        doc.numbered(["%s — %s" % (title, sub) for title, sub in v["steps"]])

    elif t == "nested":
        for layer in v["layers"]:
            doc.para(layer["label"], bold=True, space=1)
            doc.para(layer["sub"], color=GREY, indent=0.6, space=8)
        if v.get("note"):
            doc.card(None, v["note"])

    elif t == "iconrow":
        for it in v["items"]:
            doc.para(it["label"], bold=True, space=1)
            doc.para(it["sub"], color=GREY, indent=0.6, space=8)

    elif t == "split":
        doc.para(v["left"]["title"], bold=True, space=2)
        doc.bullets(v["left"]["items"])
        doc.para(v["right"]["title"], bold=True, space=2)
        doc.bullets(v["right"]["items"])

    elif t == "tree":
        doc.para(v["question"], italic=True, space=6)
        for key in ("yes", "no"):
            b = v[key]
            tone = "good" if b.get("tone") == "good" else (
                "bad" if b.get("tone") == "bad" else None)
            doc.card(b["label"], b["detail"], tone=tone)

    elif t == "steps":
        doc.numbered(v["items"])
        if v.get("prompt"):
            doc.card("Copy-paste prompt", v["prompt"])
            if v.get("caption"):
                doc.para(v["caption"], color=GREY, italic=True, size=9.5, space=8)

    elif t == "beforeafter":
        doc.para(v["bad_tag"], bold=True, color=BAD, space=2)
        doc.bullets(v["bad"], tone="bad")
        doc.para(v["good_tag"], bold=True, color=GOOD, space=2)
        doc.bullets(v["good"], tone="good")
        if v.get("note"):
            doc.para(v["note"], bold=True, space=8)

    elif t == "prompt":
        doc.card(v.get("header") or "Copy this prompt", v["text"])
        if v.get("caption"):
            doc.para(v["caption"], color=GREY, italic=True, size=9.5, space=6)
        for w in v.get("why") or []:
            doc.bullets([w])

    elif t == "prompt_out":
        doc.card(v.get("header") or "Copy this prompt", v["text"])
        if v.get("caption"):
            doc.para(v["caption"], color=GREY, italic=True, size=9.5, space=6)
        doc.para(v.get("out_title") or "What comes back", bold=True, space=2)
        doc.bullets(v["out"])

    elif t == "checklist":
        doc.bullets(v["items"], tone="bad" if v.get("mark") == "ban" else None)

    elif t == "bandlist":
        doc.card(v["headline"], v.get("sub") or "",
                 tone="bad" if v.get("tone") not in ("good", "accent", "neutral")
                 else None)
        doc.bullets(v["items"])

    elif t == "mistakes":
        for what, why in v["items"]:
            doc.para(what, bold=True, color=BAD, space=1)
            doc.para(why, color=GREY, indent=0.6, space=8)

    else:
        raise SystemExit("unhandled visual type: %s" % t)


# ---------------------------------------------------------------------------
def emit_module(d):
    accent = TRACK_ACCENT[d["area"]]
    doc = Doc(accent)
    code = d["module_code"]

    # ---- cover ----
    doc.h(1, d["title"], color=accent)
    doc.para(d["subtitle"], size=13, color=GREY, space=14)
    doc.table([
        ["Field", "Value"],
        ["Module code", code],
        ["Track", T.AREAS[d["area"]]["name"]],
        ["Duration", "%d minutes" % d["duration_min"]],
        ["Audience", d["audience"]],
        ["Suggested passing score", "80%"],
    ], widths=[5.0, 11.0])
    doc.rule()

    # ---- why ----
    w = d["why"]
    doc.h(2, "Why this matters")
    doc.para(w["title"], bold=True, space=4)
    doc.para(w["scenario"])
    doc.card("The cost", w["cost"], tone="bad")
    doc.card("What changes", w["fix"], tone="good")

    # ---- outcomes ----
    doc.h(2, "What you'll be able to do")
    doc.bullets([text for icon, text in d["outcomes"]])

    # ---- objective, for the OneWork "Objective" field ----
    doc.h(2, "Objective")
    doc.para(
        "By the end of this module you will be able to: "
        + "; ".join(text[0].lower() + text[1:] for icon, text in d["outcomes"][:3])
        + ".", italic=True)

    # ---- content sections ----
    groups = sitegen.group_slides(d)
    for name, sub, anchor in d["sections"]:
        if anchor == "scenario":
            sc = d["scenario"]
            doc.h(2, name)
            doc.para(sub, color=GREY, italic=True, space=6)
            doc.para(sc["title"], bold=True, space=2)
            doc.para(sc["situation"])
            for c in sc["choices"]:
                tone = {"good": "good", "bad": "bad"}.get(c["tone"])
                doc.para(c["text"], bold=True, space=1)
                doc.card(c["headline"], c["consequence"] + "  Rule: " + c["rule"],
                         tone=tone)
            continue

        if anchor == "video":
            v = d["video"]
            doc.h(2, name)
            doc.para(sub, color=GREY, italic=True, space=6)
            doc.para(v.get("heading") or "", bold=True, space=2)
            doc.para("Video: “%s” — %s (%s)"
                     % (v["title"], v["channel"], v["duration"]), bold=True)
            doc.para(v["url"], color=accent)
            doc.para(v["note"], color=GREY, italic=True)
            doc.bullets(v["how"])
            continue

        doc.h(2, name)
        doc.para(sub, color=GREY, italic=True, space=6)
        for s in groups.get(anchor, []):
            doc.h(3, s["title"])
            if s.get("lead"):
                doc.para(s["lead"])
            if s.get("visual"):
                emit_visual(doc, s["visual"])
            if s.get("gloss"):
                doc.para("Terms used here: " + ", ".join(s["gloss"]),
                         color=GREY, size=9.5, space=8)

    # ---- knowledge check ----
    doc.h(2, "Knowledge check")
    doc.para(
        "Suggested weight: equal (100%) per question unless noted otherwise "
        "in the OneWork quiz builder.", color=GREY, italic=True, size=9.5)
    for qi, q in enumerate(d["quiz"], 1):
        doc.h(3, "Q%d. %s" % (qi, q["q"]))
        if q.get("stem"):
            doc.para(q["stem"], italic=True)
        for ai, a in enumerate(q["answers"]):
            mark = "✓ " if a["ok"] else "✗ "
            doc.para(mark + a["text"], bold=a["ok"],
                     color=GOOD if a["ok"] else BAD, space=1)
            doc.para(a["why"], color=GREY, indent=0.6, space=8)
        doc.card("Remember", q["remember"])

    # ---- recap ----
    r = d["recap"]
    doc.h(2, "Recap")
    doc.para(r["title"], bold=True, space=6)
    for t, dd in r["points"]:
        doc.para(t, bold=True, space=1)
        doc.para(dd, color=GREY, indent=0.6, space=8)
    doc.card(None, r["oneliner"], tone="good")

    # ---- toolkit ----
    tk = d["toolkit"]
    doc.h(2, "Toolkit")
    doc.para(tk["title"], bold=True, space=6)
    for icon, t, dd in tk["templates"]:
        doc.para(t, bold=True, space=1)
        doc.para(dd, color=GREY, indent=0.6, space=8)
    doc.para("Tools referenced: " + "; ".join(
        "%s (%s)" % (a, b) for a, b in tk["links"]), color=GREY, size=9.5)
    doc.para("Next module: " + tk["next"], italic=True)

    # ---- glossary ----
    doc.h(2, "Glossary")
    for term, definition in d["glossary"]:
        doc.para(term, bold=True, space=1)
        doc.para(definition, color=GREY, indent=0.6, space=8)

    return doc.d


def main():
    decks = [importlib.import_module(m).DECK for m in build.REGISTRY]
    order = {a: i for i, a in enumerate(sitegen.AREA_ORDER)}
    decks.sort(key=lambda d: (order[d["area"]], d["module_code"]))

    manifest = []
    for d in decks:
        sub = os.path.join(OUTDIR, AREA_DIR[d["area"]])
        os.makedirs(sub, exist_ok=True)
        fn = filename(d)
        path = os.path.join(sub, fn)
        doc = emit_module(d)
        doc.save(path)
        manifest.append({
            "module_code": d["module_code"],
            "title": d["title"],
            "track": T.AREAS[d["area"]]["name"],
            "path": os.path.relpath(path, ROOT).replace("\\", "/"),
        })
        print("  %-7s %s" % (d["module_code"], fn))

    with io.open(os.path.join(OUTDIR, "manifest.json"), "w", encoding="utf-8") as fh:
        json.dump(manifest, fh, ensure_ascii=False, indent=2)

    print("\n%d module documents written under %s" % (len(decks), OUTDIR))


if __name__ == "__main__":
    main()
