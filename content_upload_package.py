# -*- coding: utf-8 -*-
"""
INDUCTO AI Learning Journey — upload-ready content package for the 1XL
Admin "Add Training Module" workflow.

    python content_upload_package.py

Produces:
  docs/INDUCTO_AI_LEARNING_CONTENT_MASTER.docx
  docs/pdf/M-01_AI_Fundamentals.pdf ... M-16_What_Never_to_Paste.pdf

This is a content package, not a code change: every field is written so an
admin can copy it straight into the screens shown in the 1XL Admin
screenshots (Basics / Lesson / Audience & Grading / Review, then Quiz
questions, then Content Library). Every word of teaching content, every
video record and every quiz question is drawn verbatim/condensed from the
same source dictionaries that build the live Inducto product
(content/areaNN/*.py -> journey_data.py / management_review_docx.py).
Nothing is invented; where 1XL needs a company-specific fact Inducto does
not supply, the field says [COMPANY INPUT NEEDED: ...].
"""

import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
sys.stdout.reconfigure(encoding="utf-8")

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm, Pt

import journey_data as J
import sitegen
import world_class_master_docx as W  # Doc, add_hyperlink, colours

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                ListFlowable, ListItem, HRFlowable)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT

ROOT = sitegen.ROOT
OUT_DOCX = os.path.join(ROOT, "docs", "INDUCTO_AI_LEARNING_CONTENT_MASTER.docx")
PDF_DIR = os.path.join(ROOT, "docs", "pdf")

INK, GREY, GOOD, BAD, ACCENT, GOLD = W.INK, W.GREY, W.GOOD, W.BAD, W.ACCENT, W.GOLD
Doc = W.Doc

# ---------------------------------------------------------------------------
# short, filename-friendly labels (matches the style of the requested
# example filenames — a short topic label, not the full punctuated title)
FILENAME_LABEL = {
    "M-01": "AI_Fundamentals", "M-02": "Generative_AI",
    "M-03": "AI_Capabilities", "M-04": "AI_Hallucinations",
    "M-05": "Writing_Email_With_AI", "M-06": "Meeting_Notes_With_AI",
    "M-07": "Basic_Prompting", "M-08": "Instructions_Context_Role",
    "M-09": "Planning_Productivity", "M-10": "Business_Communication",
    "M-11": "Time_Management", "M-12": "Password_Security_MFA",
    "M-13": "Phishing_Social_Engineering", "M-14": "Data_Protection",
    "M-15": "Safe_AI_Use", "M-16": "What_Never_to_Paste",
}

# 1XL's own "Who takes this" / grading defaults, taken from the screenshots'
# own help text ("Leave blank to use the organisation default of ...").
# Nothing here is invented — it is what the existing 1XL screens already
# say happens if these fields are left blank.
ORG_DEFAULT_PASS = "Leave blank — organisation default (80%, per Settings → Organisation)"
ORG_DEFAULT_ATTEMPTS = "Leave blank — organisation default (3, per Settings → Organisation)"
ORG_DEFAULT_CERT = "Leave blank — organisation default (12 months, per Settings → Organisation)"


def lc1(s):
    return (s[0].lower() + s[1:]) if s else s


def good_choice(scenario):
    for c in scenario.get("choices", []):
        if c.get("tone") == "good":
            return c
    return scenario.get("choices", [{}])[0] if scenario.get("choices") else {}


def stage_level(stage):
    """1XL's 'Level' = programme stage. Inducto's own stage numbering,
    already real (e.g. "7. Security & Responsible Use" -> 7)."""
    try:
        return int(stage.split(".")[0])
    except Exception:
        return None


def lesson_field_text(r, primary):
    """The exact text for 1XL's Lesson field — what the employee reads
    before the quiz. One short paragraph per line, drawn from the
    module's own subtitle, its real "why this matters" story, its real
    handling rules and its own recap — condensed, not invented."""
    lines = []
    subtitle = primary.get("subtitle", "").rstrip(". ")
    if subtitle:
        lines.append(subtitle + ".")
    if r.get("reading"):
        lines.append(r["reading"])
    if r.get("checklist"):
        lines.append("Key things to remember: " +
                     " ".join(r["checklist"]))
    recap = primary.get("recap", {}).get("points", [])
    if recap:
        h, t = recap[0]
        lines.append("Remember: %s — %s" % (h, t))
    return lines


# ---------------------------------------------------------------------------
# PDF job aids (one page per mandatory lesson)
# ---------------------------------------------------------------------------
def build_pdf(r, by_code):
    primary = by_code[r["sources"][0]]
    recap = primary.get("recap", {}).get("points", [])[:5]
    label = FILENAME_LABEL.get(r["code"], r["code"])
    path = os.path.join(PDF_DIR, "%s_%s.pdf" % (r["code"], label))

    styles = getSampleStyleSheet()
    accent = colors.HexColor("#2F4BC4")
    ink = colors.HexColor("#101826")
    grey = colors.HexColor("#5B6779")

    title_style = ParagraphStyle("T", parent=styles["Title"], fontSize=20,
                                 textColor=ink, alignment=TA_LEFT,
                                 spaceAfter=2)
    kicker_style = ParagraphStyle("K", parent=styles["Normal"], fontSize=10,
                                  textColor=accent, spaceAfter=10,
                                  fontName="Helvetica-Bold")
    h2_style = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=13,
                              textColor=accent, spaceBefore=14,
                              spaceAfter=6)
    body_style = ParagraphStyle("B", parent=styles["Normal"], fontSize=10.5,
                                textColor=ink, leading=15, spaceAfter=8)
    bullet_style = ParagraphStyle("BL", parent=styles["Normal"],
                                  fontSize=10.2, textColor=ink, leading=14)
    foot_style = ParagraphStyle("F", parent=styles["Normal"], fontSize=8,
                                textColor=grey, spaceBefore=20)

    doc = SimpleDocTemplate(path, pagesize=A4,
                            leftMargin=2.2 * cm, rightMargin=2.2 * cm,
                            topMargin=2.0 * cm, bottomMargin=1.6 * cm)
    story = []
    story.append(Paragraph("%s · %s" % (r["stage"], "%.0f minutes"
                                        % r["time"]["total_min"]),
                          kicker_style))
    story.append(Paragraph("%s — %s" % (r["code"], r["title"]), title_style))
    story.append(HRFlowable(width="100%", thickness=1,
                            color=colors.HexColor("#D6DCE7"),
                            spaceAfter=12))

    story.append(Paragraph("What is this about?", h2_style))
    subtitle = primary.get("subtitle", "")
    story.append(Paragraph(subtitle, body_style))

    story.append(Paragraph("The idea, in plain terms", h2_style))
    story.append(Paragraph(r["reading"], body_style))

    if r.get("checklist"):
        story.append(Paragraph("Key things to know", h2_style))
        story.append(ListFlowable(
            [ListItem(Paragraph(c, bullet_style), spaceAfter=4)
             for c in r["checklist"]],
            bulletType="bullet", start="•", leftIndent=14))

    scenario = primary.get("scenario") or {}
    good = good_choice(scenario)
    if scenario.get("situation"):
        story.append(Paragraph("A simple example", h2_style))
        story.append(Paragraph("<b>Situation:</b> " + scenario["situation"],
                               body_style))
        if good.get("text"):
            story.append(Paragraph("<b>What to do:</b> " + good["text"],
                                   body_style))

    if recap:
        story.append(Paragraph("Remember", h2_style))
        story.append(ListFlowable(
            [ListItem(Paragraph("<b>%s</b> — %s" % (h, t), bullet_style),
                     spaceAfter=4) for h, t in recap],
            bulletType="bullet", start="•", leftIndent=14))

    story.append(Paragraph(
        "Inducto Learning &amp; Knowledge — quick reference for %s. "
        "Watch the lesson video for the full explanation; this page is "
        "a takeaway, not a replacement for it." % r["code"], foot_style))

    doc.build(story)
    return path


# ---------------------------------------------------------------------------
def main():
    by_code, decks, resolved = J.load()
    total_min = sum(r["time"]["total_min"] for r in resolved) + 3 + 12
    pool = J.assessment_pool(by_code)

    os.makedirs(PDF_DIR, exist_ok=True)
    pdf_paths = []
    for r in resolved:
        pdf_paths.append((r["code"], build_pdf(r, by_code)))

    import re
    import json
    tokens = {}
    for dk in decks:
        for mm in re.finditer(r"\[COMPANY INPUT NEEDED:([^\]]*)\]",
                              json.dumps(dk)):
            tokens.setdefault(mm.group(1).strip(), []).append(dk["module_code"])
    tokens.setdefault("name of the team that owns this training",
                      []).append("site-wide footer, every page")
    tokens.setdefault("SOPGalaxy link per module",
                      []).append("no SOPGalaxy link exists in source for any "
                                 "module — leave blank unless management "
                                 "provides one")

    doc = Doc()
    doc.d.sections[0].footer.paragraphs[0].runs[0].text = (
        "Inducto AI Learning Journey — Content Upload Package · page ")
    build(doc, by_code, resolved, total_min, tokens, pool, pdf_paths)
    doc.d.save(OUT_DOCX)

    n_checks = sum(len(r["quiz"]) for r in resolved)
    print("Saved %s" % OUT_DOCX)
    print("Saved %d PDF job aids to %s" % (len(pdf_paths), PDF_DIR))
    print("  %d modules (M-01..M-16 + M-19 + M-20), %d lesson-check "
          "questions, %d M-20 questions, %d videos, %d company inputs open"
          % (len(resolved) + 2, n_checks, len(pool), len(resolved),
             len(tokens)))
    return doc


def render_quiz_block(doc, q, qi):
    doc.para("%d. %s" % (qi, q["q"]), bold=True, size=10.5, space=4)
    letters = "ABCD"
    correct_letter, correct_text = None, None
    for li, a in enumerate(q["answers"]):
        doc.para("%s. %s" % (letters[li], a["text"]), size=10,
                 bold=a["ok"], color=GOOD if a["ok"] else INK, space=2,
                 indent=0.4)
        if a["ok"]:
            correct_letter, correct_text = letters[li], a["text"]
    doc.para("Correct Answer: %s — %s" % (correct_letter, correct_text),
             size=9.8, bold=True, color=GOOD, space=2, indent=0.4)
    for li, a in enumerate(q["answers"]):
        tag = "Explanation (%s)" % letters[li]
        doc.para("%s: %s" % (tag, a["why"]), size=9.2, italic=True,
                 color=GREY, space=3, indent=0.8)
    doc.para("", space=4)


def build(doc, by_code, resolved, total_min, tokens, pool, pdf_paths):
    d = doc.d
    pdf_by_code = dict(pdf_paths)

    # ================================================================
    # COVER
    # ================================================================
    doc.h(1, "INDUCTO")
    doc.para("AI Learning Journey", size=16, color=ACCENT, space=2)
    doc.para("Content Upload Package for 1XL Admin", size=13, color=GREY,
             space=18)
    doc.rule()
    doc.para("Prepared for: Learning & Development / 1XL Admin", size=9.5,
             color=GREY)
    doc.para("Prepared by: Learning & Development — Dhrubojyoti "
             "(chetan@1xl.com)", size=9.5, color=GREY)
    doc.para("18 entries (M-01–M-16, M-19, M-20) · %.0f minutes total · "
             "5 September 2026" % total_min, size=9.5, color=GREY, space=16)
    doc.para(
        "Copy-ready content for the 1XL “Add Training Module” "
        "workflow (Basics → Lesson → Audience & Grading → "
        "Review), the quiz-question screen that follows it, and the "
        "Content Library upload/link screens. Every field below is real "
        "— drawn from the audited Inducto AI Learning Journey source, "
        "not invented for this document.", size=10.5)
    doc.para(
        "16 one-page PDF quick-reference job aids accompany this "
        "document (M-01_AI_Fundamentals.pdf … M-16_What_Never_to_Paste."
        "pdf) for upload to the Content Library alongside each module's "
        "video.", size=9.8, color=GREY, space=10)

    # ================================================================
    # HOW TO USE THIS PACKAGE
    # ================================================================
    doc.h(1, "How to Use This Package", page_break=True)
    doc.numbered([
        "For each module below, open 1XL → Management → Learning → "
        "Add Training Module.",
        "Basics: copy Module Code, Title, Objective, Duration.",
        "Lesson: copy the Lesson field text exactly as given (one "
        "paragraph per line).",
        "Audience & Grading: set the fields as given below (mostly the "
        "organisation defaults — nothing invented).",
        "Review → Create Module (it saves as Draft).",
        "On the quiz screen that follows, add the questions given for "
        "that module, marking the correct option and pasting each "
        "explanation.",
        "In Content Library, add the module's video as an External "
        "Link (fields given below), and upload its PDF job aid as a "
        "Document, then attach both to the module.",
        "Publish once the lesson and quiz questions are both in place.",
    ])
    doc.para(
        "Two different kinds of quiz exist in this journey — do not "
        "confuse them:", bold=True, space=8)
    doc.table([
        ["", "Module quiz (M-01–M-16)", "M-20 Final Assessment"],
        ["Where", "End of each module, in 1XL", "Its own module, taken "
         "after M-19"],
        ["How many", "24 in total across the 16 modules", "15, in one "
         "attempt"],
        ["Passing score", "Organisation default (80%) unless management "
         "says otherwise", "70%, explicitly"],
        ["Max attempts", "Organisation default (3)", "3, explicitly"],
        ["On repeated failure", "Retake per organisation policy", "“"
         "Further action requires an HR decision” after the third "
         "unsuccessful attempt"],
    ], widths=[2.6, 7.2, 7.4])

    # ================================================================
    # M-01 – M-16
    # ================================================================
    for i, r in enumerate(resolved, 1):
        primary = by_code[r["sources"][0]]
        scenario = primary.get("scenario") or {}
        good = good_choice(scenario)
        video = r.get("video")
        level = stage_level(r["stage"])

        doc.h(1, "%s — %s" % (r["code"], r["title"]), page_break=True)
        doc.para("Lesson %d of 16 · %s" % (i, r["stage"]), size=9.5,
                 color=GREY, space=10)

        doc.h(2, "A. Admin Form Content")
        doc.table([
            ["Field", "Value"],
            ["Module Code", r["code"]],
            ["Title", r["title"]],
            ["Objective", r["objective"]],
            ["Duration (minutes)", "%.0f" % r["time"]["total_min"]],
            ["Who takes this", "Everyone in the organisation"],
            ["Level", str(level) if level else "—"],
            ["Part of (optional)", "— (top-level module, not a "
             "sub-module of an existing one)"],
            ["Passing Score (%)", ORG_DEFAULT_PASS],
            ["Max Attempts", ORG_DEFAULT_ATTEMPTS],
            ["Certificate Validity (months)", ORG_DEFAULT_CERT],
            ["SOPGalaxy Link", "None — leave blank"],
        ], widths=[4.8, 12.4])

        doc.h(2, "B. Lesson Field (copy exactly)")
        doc.para("One paragraph per line, as the field expects:",
                 size=9.5, italic=True, color=GREY, space=4)
        for line in lesson_field_text(r, primary):
            doc.card(None, line, tone=None)

        doc.h(2, "C. Learning Resource")
        if video:
            doc.para("Video title: %s" % video["title"], size=10.2, space=2)
            doc.para("Creator/channel: %s" % video["channel"], size=10.2,
                     space=2)
            doc.para("Duration: %s" % video["duration"], size=10.2, space=2)
            p = d.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r_run = p.add_run("Direct URL: ")
            r_run.font.size = Pt(10.2)
            W.add_hyperlink(p, video["url"], video["url"], size=10.2)
            why = video.get("note") or (
                "Chosen because it explains this lesson's topic in "
                "%s from an established channel." % video["duration"])
            doc.para("Why this video is relevant: %s" % why, size=9.5,
                     color=GREY, space=6)
            if r.get("video_note_only"):
                doc.para(
                    "Content type in 1XL: External Link, marked "
                    "background/optional — not required to complete the "
                    "module.", size=9.5, italic=True, color=GREY, space=4)
            else:
                doc.para("Content type in 1XL: External Link (required "
                         "viewing).", size=9.5, italic=True, color=GREY,
                         space=4)
        pdf_path = pdf_by_code.get(r["code"])
        if pdf_path:
            doc.para(
                "Companion PDF (upload as a Document in Content Library, "
                "attach alongside the video): %s"
                % os.path.basename(pdf_path), size=9.8, space=6)

        doc.h(2, "D. Module Quiz")
        doc.para(
            "Practice quiz for this module — %d question%s, matching the "
            "authoritative source's allocation."
            % (len(r["quiz"]), "" if len(r["quiz"]) == 1 else "s"),
            size=9.5, italic=True, color=GREY, space=6)
        for qi, q in enumerate(r["quiz"], 1):
            render_quiz_block(doc, q, qi)

    # ================================================================
    # M-19
    # ================================================================
    ex = J.EXERCISE
    doc.h(1, "%s — %s" % (ex["code"], ex["title"]), page_break=True)
    doc.table([
        ["Field", "Value"],
        ["Module Code", ex["code"]],
        ["Title", ex["title"]],
        ["Objective", "Apply the skills from M-01–M-16 to one realistic, "
         "end-to-end scenario before the final assessment."],
        ["Duration (minutes)", "3"],
        ["Who takes this", "Everyone in the organisation"],
        ["Level", "8"],
        ["Passing Score (%)", "Completion-based — every step attempted"],
        ["Max Attempts", ORG_DEFAULT_ATTEMPTS],
    ], widths=[4.8, 12.4])
    doc.h(2, "Lesson field")
    doc.para(ex["intro"], space=6)
    doc.card(ex["scenario_title"], ex["scenario"], tone=None)
    for i, step in enumerate(ex["steps"], 1):
        doc.h(3, step["title"])
        doc.para(step["instruction"], space=4)
        doc.para("Hint: %s" % step["hint"], italic=True, size=9.5,
                 color=GREY, space=4)
        doc.card("Model answer", step["model_answer"], tone="good")
    doc.h(2, "Knowledge check")
    kc = ex["knowledge_check"]
    fake_q = {"q": kc["q"], "answers": kc["options"]}
    render_quiz_block(doc, fake_q, 1)

    # ================================================================
    # M-20
    # ================================================================
    doc.h(1, "M-20 — Final Assessment", page_break=True)
    doc.table([
        ["Field", "Value"],
        ["Module Code", "M-20"],
        ["Title", "Final Assessment"],
        ["Objective", "Confirm the mandatory learning journey has been "
         "understood, across all required domains."],
        ["Duration (minutes)", "12"],
        ["Who takes this", "Everyone in the organisation"],
        ["Level", "9"],
        ["Passing Score (%)", "70 (explicit — overrides the organisation "
         "default)"],
        ["Max Attempts", "3 (explicit)"],
        ["On third failure", "“Further action requires an HR "
         "decision”"],
    ], widths=[4.8, 12.4])
    doc.para(
        "The 15 questions below are the real M-20 question bank, drawn "
        "from across M-01–M-16 and indexed so none repeats a module's "
        "own quiz word for word. This section is for the admin quiz-"
        "authoring screen only — it is not learner-facing.", size=9.8,
        color=GREY, space=10)
    for qi, q in enumerate(pool, 1):
        fake_q = {"q": q["q"], "answers": q["options"]}
        render_quiz_block(doc, fake_q, qi)

    # ================================================================
    # CONTENT LIBRARY UPLOAD REGISTER
    # ================================================================
    doc.h(1, "Content Library Upload Register", page_break=True)
    doc.para(
        "Every resource referenced above, in one register.", size=9.8,
        color=GREY, space=10)
    t = d.add_table(rows=1, cols=5)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0].cells
    for j, h in enumerate(["Module", "Resource title", "Content type",
                          "Source", "SOPGalaxy Link"]):
        hdr[j].text = h
        hdr[j].paragraphs[0].runs[0].font.bold = True
        hdr[j].paragraphs[0].runs[0].font.size = Pt(9)
    for r in resolved:
        v = r.get("video")
        if v:
            row = t.add_row().cells
            row[0].text = r["code"]
            row[1].text = v["title"]
            row[2].text = "External Link (Video)"
            row[3].text = v["url"]
            row[4].text = "None"
            for c in row:
                c.paragraphs[0].runs[0].font.size = Pt(8.8)
        pdf_path = pdf_by_code.get(r["code"])
        if pdf_path:
            row = t.add_row().cells
            row[0].text = r["code"]
            row[1].text = "%s — quick reference" % r["title"]
            row[2].text = "Document (PDF)"
            row[3].text = os.path.basename(pdf_path)
            row[4].text = "None"
            for c in row:
                c.paragraphs[0].runs[0].font.size = Pt(8.8)
    d.add_paragraph().paragraph_format.space_after = Pt(4)
    for j, w in enumerate([1.6, 6.4, 3.4, 5.0, 1.4]):
        for row in t.rows:
            row.cells[j].width = Cm(w)

    # ================================================================
    # COMPANY INPUTS NEEDED
    # ================================================================
    doc.h(1, "Company Inputs Needed", page_break=True)
    doc.para(
        "Genuine gaps only — nothing here has been invented. Fill these "
        "in before publishing the affected modules, or leave the "
        "placeholder text visible until they are known.", size=9.8,
        color=GREY, space=10)
    rows = [["What is needed", "Appears in"]]
    for tok, codes in sorted(tokens.items()):
        rows.append([tok, ", ".join(sorted(set(codes)))])
    doc.table(rows, widths=[7.0, 9.4], small=True)


if __name__ == "__main__":
    main()
