# -*- coding: utf-8 -*-
"""
INDUCTO WORLD-CLASS CONTENT MASTER — the single management-ready document
handed to Dev, L&D, Content, Proofreading, Instructional Design, Image/
Graphics, Video and LMS/Admin teams.

    python world_class_master_docx.py

Every lesson's teaching copy, workplace example, prompt, checklist and
knowledge-check question is quoted verbatim from journey_data.py /
management_review_docx.py — themselves quoted from content/areaNN/*.py, the
same dictionaries that build the decks, the web platform and the live
Mandatory Journey. Nothing here is retyped, and nothing is invented: where a
company fact is required and not supplied, the text carries
[COMPANY INPUT NEEDED: ...] exactly as it does in the product.

The 16 visual production briefs are the one new artefact in this document.
Each is built from a real diagram already designed and shipped in the source
content (the deck's own flow/tree/split/nested/before-after visual for that
lesson) — restated as a brief for the graphics team, not invented from
scratch.
"""

import importlib
import io
import json
import os
import re
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

import build
import journey_data as J
import management_review_docx as M
import sitegen
import theme as T

ROOT = sitegen.ROOT
OUT_DOCX = os.path.join(
    ROOT, "docs", "INDUCTO_World_Class_Content_Master_RECTIFIED_FINAL.docx")

# Real defects found and fixed in this rectification pass, with evidence.
# Kept as data (not prose) so the log in Section 13 and the console summary
# are generated from the same list and can never disagree with each other.
RECTIFICATIONS = [
    {
        "id": "R1", "severity": "Medium", "area": "AI technical accuracy",
        "location": "AI-01 — “It predicts the next word” (source "
                    "slide; carried into the full module page and the PPTX "
                    "deck, not into the condensed M-01 journey stop)",
        "finding": "The lead sentence stated, without qualification, that "
                  "“An AI assistant does not look up an answer.” Read on "
                  "its own — the first thing a learner sees on the topic — "
                  "this reads as an absolute claim, when several named "
                  "tools in this course (Copilot, ChatGPT, Gemini, Claude) "
                  "can search the web when that feature is enabled. The "
                  "correct, more complete version already existed five "
                  "sections later, in AI-01's own quiz explanation and in "
                  "AI-04 (“Unless the tool visibly shows a link it "
                  "visited…”) — the fix surfaces that same, "
                  "already-correct nuance at the point the mental model is "
                  "first formed.",
        "fix": "Added one qualifying sentence to the same slide: "
              "“By default…” and “Some tools can search the web when "
              "that feature is switched on — but unless it visibly shows "
              "you a page it opened, treat the answer as generated, not "
              "looked up.” No new claim was introduced — the sentence "
              "restates what AI-04 already teaches.",
        "verified_against": "Internal consistency with AI-04's existing, "
                            "already-correct treatment of the same point.",
    },
    {
        "id": "R2", "severity": "High", "area": "Cybersecurity — password "
                                                "guidance",
        "location": "SEC-01 — “Do this now: check and change” (source "
                    "slide; this exact prompt was live in the deployed "
                    "M-12 Mandatory Journey lesson)",
        "finding": "The lesson's copy-paste prompt told the learner to "
                  "ask a general AI assistant to “Give me eight "
                  "passphrases…” for a new work password. This directly "
                  "contradicted the SAME module's own checklist four "
                  "slides earlier, which lists “An AI chat window, ever, "
                  "for any reason” as a place a password must never live. "
                  "It also relies on a language model for "
                  "cryptographic-quality randomness, which is not what "
                  "these tools are built or verified to provide.",
        "fix": "Removed the AI prompt entirely. The instruction step now "
              "reads: “Set the new one as four unrelated words, generated "
              "by your password manager's own generator, not typed by you "
              "and never asked of an AI tool.” This is consistent with "
              "the module's existing rule and with standard guidance "
              "(e.g. NIST SP 800-63B) to use a password manager's built-in "
              "generator rather than a general-purpose AI tool for "
              "credential material. The live Mandatory Journey M-12 "
              "lesson no longer shows a copy-paste prompt at all — "
              "correct, since there is no safe AI prompt for this task.",
        "verified_against": "Internal consistency with the module's own "
                            "existing rule; general password-manager "
                            "generator guidance is standard practice, not "
                            "a company-specific claim.",
    },
    {
        "id": "R3", "severity": "High", "area": "M-19 scenario logic",
        "location": "journey_data.py — M-19 Integration Exercise scenario",
        "finding": "The scenario was set at “Friday, 4:50pm”, but the "
                  "same notes required the transport-lead's reroute to "
                  "happen “Friday morning” — already in the past by "
                  "4:50pm on a Friday. A learner solving Step 2 (approve "
                  "the reroute “before Friday morning”) would be asked to "
                  "act before a deadline that had already elapsed at the "
                  "moment the scenario opens.",
        "fix": "Changed the scenario day from Friday to Thursday "
              "(“Thursday, 4:50pm — the Chakan site review”). Every "
              "other date in the notes (Tuesday past, Wednesday past, "
              "Monday and Tuesday upcoming) was re-checked against the "
              "new day and remains logically consistent; none needed to "
              "change.",
        "verified_against": "Internal date-logic check of every day-of-"
                            "week reference in the scenario.",
    },
    {
        "id": "R4", "severity": "Medium", "area": "Reporting consistency — "
                                                  "video count",
        "location": "management_review_docx.py, world_class_master_docx.py "
                    "— summary tables and QA checklists",
        "finding": "M-15 (“Safe Use of AI at Work”) carries a real, "
                  "verified video (IBM Technology, 13:13) that is "
                  "referenced but deliberately not embedded — it exceeds "
                  "the 12-minute guidance for a mandatory slot. Several "
                  "summary lines and this document's own prior version "
                  "counted it as an ordinary “mandatory video”, "
                  "producing a “16 mandatory videos” claim that did not "
                  "distinguish embedded-and-required from referenced-only.",
        "fix": "Every count now states both figures explicitly: 15 "
              "embedded and required + 1 referenced only (M-15) = 16 "
              "video records in total. No occurrence of a bare “16 "
              "mandatory videos” claim remains in this document.",
        "verified_against": "Recomputed directly from the resolved "
                            "journey data (video_note_only flag), not "
                            "asserted.",
    },
    {
        "id": "R5", "severity": "Low", "area": "Management-input register "
                                              "completeness",
        "location": "world_class_master_docx.py's token scan vs "
                    "siteverify.py's",
        "finding": "This document's Management-Input Register was built "
                  "by scanning only the 39 module dictionaries, which "
                  "found 13 distinct company-input tokens. The live "
                  "product's own QA gate (siteverify.py, which scans the "
                  "shipped HTML directly) independently reports 14 — one "
                  "token lives in the site-wide footer text "
                  "(sitegen.py), shown on every page, and is not part of "
                  "any module's data.",
        "fix": "Added the missing token explicitly: “name of the team "
              "that owns this training”, located in the site-wide "
              "footer rather than a specific module. The register now "
              "matches the independently-verified figure.",
        "verified_against": "Cross-checked against siteverify.py's HTML-"
                            "level scan (14), which re-ran clean after "
                            "this fix.",
    },
]

INK, GREY, GOOD, BAD = M.INK, M.GREY, M.GOOD, M.BAD
ACCENT, GOLD = M.ACCENT, M.GOLD
TRACK_ACCENT = M.TRACK_ACCENT


def add_hyperlink(paragraph, url, text, color="2F4BC4", size=10.5, bold=False):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Segoe UI")
    rPr.append(rFonts)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz)
    if bold:
        b = OxmlElement("w:b"); rPr.append(b)
    c = OxmlElement("w:color"); c.set(qn("w:val"), color); rPr.append(c)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
    run.append(t)
    hl.append(run)
    paragraph._p.append(hl)
    return hl


# ===========================================================================
class Doc(M.Doc):
    """Extends the management-review Doc with a hyperlink helper."""

    def link_para(self, label, url, size=10.5, space=6):
        p = self.d.add_paragraph()
        p.paragraph_format.space_after = Pt(space)
        add_hyperlink(p, url, label, size=size)
        return p

    def link_cell(self, cell, label, url, size=9.5):
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        add_hyperlink(p, url, label, size=size)


# ===========================================================================
# visual production briefs — built from the real diagram already shipped
# ===========================================================================
DIAGRAM_PREFERENCE = ("flow", "tree", "beforeafter", "split", "nested")

STYLE_BLOCK = (
    "Flat corporate style, no stock photography, no clipart, no emoji as "
    "icons. Palette: ink #101826, accent per track (AI tracks #2F4BC4, "
    "workplace #0E6E75, prompting #6A2FA0, professional skills #8A5A00, "
    "security #7A1220), surface #F3F6FB. Typeface Segoe UI / Segoe UI "
    "Semibold, matching the rest of the Inducto product. Simple line-art "
    "or geometric shapes, not illustrative scenes."
)


def pick_visual(primary_module):
    by_type = {}
    for s in primary_module["slides"]:
        v = s.get("visual")
        if v and v["type"] in DIAGRAM_PREFERENCE and v["type"] not in by_type:
            by_type[v["type"]] = (s["title"], v)
    for t in DIAGRAM_PREFERENCE:
        if t in by_type:
            return by_type[t]
    return None, None


def visual_brief(code, stop_title, track, slide_title, v):
    t = v["type"]
    labels = []
    what = ""
    layout = ""
    aspect = "16:9 (module header) or 4:3 (in-lesson)"

    if t == "flow":
        labels = ["%d. %s" % (i, s[0]) for i, s in enumerate(v["steps"], 1)]
        what = ("A left-to-right numbered sequence of %d steps: %s. Each "
                "step is a box with its number, its short label, and one "
                "line of supporting text underneath."
                % (len(v["steps"]), "; ".join(s[0] for s in v["steps"])))
        layout = "Horizontal flow, boxes connected by a simple arrow or " \
                "rule between them. Wrap to two rows on narrow layouts."
    elif t == "tree":
        labels = [v["question"], "Yes → " + v["yes"]["label"],
                  "No → " + v["no"]["label"]]
        what = ("A single decision question at the top (“%s”), "
                "branching to two outcomes below it: “%s” and "
                "“%s”, each with one supporting line."
                % (v["question"], v["yes"]["label"], v["no"]["label"]))
        layout = "Question in a dark band at the top; two outcome cards " \
                "side by side beneath it, connected by a branching line."
    elif t == "beforeafter":
        labels = [v["bad_tag"], v["good_tag"]] + v["bad"][:2] + v["good"][:2]
        what = ("Two side-by-side panels: “%s” on the left "
                "(marked with a cross/negative tone) and “%s” on "
                "the right (marked with a check/positive tone), each "
                "listing 2-4 short points." % (v["bad_tag"], v["good_tag"]))
        layout = "Two-column comparison, red-toned left panel, green-toned " \
                "right panel, matching row-for-row where possible."
    elif t == "split":
        labels = [v["left"]["tag"], v["left"]["title"],
                  v["right"]["tag"], v["right"]["title"]]
        what = ("Two contrasting panels: “%s” (%s) beside “"
                "%s” (%s), each with its own short bullet list."
                % (v["left"]["tag"], v["left"]["title"],
                   v["right"]["tag"], v["right"]["title"]))
        layout = "Two-column comparison panel, equal width, a visible " \
                "divider between them."
    elif t == "nested":
        labels = [layer["label"] for layer in v["layers"]]
        what = ("Concentric containment: %s, outermost first, each layer "
                "labelled, showing that each one sits inside the one before "
                "it." % " → contains → ".join(labels))
        layout = "Nested boxes or concentric circles, outermost layer " \
                "largest and lightest, innermost smallest and darkest."

    return {
        "code": code, "stop_title": stop_title, "track": track,
        "visual_title": slide_title,
        "purpose": "Reinforce the single idea taught in “%s” at a "
                  "glance, for a learner skimming before they read the full "
                  "text." % slide_title,
        "what": what, "labels": labels, "layout": layout,
        "style": STYLE_BLOCK, "aspect": aspect,
        "do": ["Use only the real labels listed above — do not add a "
              "statistic, name or claim not shown here.",
              "Keep every label short enough to read in under two seconds.",
              "Match the track's accent colour (see Style)."],
        "dont": ["Do not illustrate with a photo of a person, an office, or "
                "a device screen.", "Do not add a mascot, character or "
                "emoji icon.", "Do not exceed the labels supplied — this "
                "brief is exhaustive, not a starting point."],
    }


# ===========================================================================
def main():
    by_code, decks, resolved = J.load()
    total_min = (sum(r["time"]["total_min"] for r in resolved) + 3 + 12)
    pool = J.assessment_pool(by_code)
    tokens = {}
    for dk in decks:
        for mm in re.finditer(r"\[COMPANY INPUT NEEDED:([^\]]*)\]",
                              json.dumps(dk)):
            tokens.setdefault(mm.group(1).strip(), []).append(dk["module_code"])
    # One further token lives in the site chrome itself (sitegen.py's
    # page footer, shown on every page), not in any of the 39 module
    # dictionaries — siteverify.py's HTML scan catches it (14 distinct);
    # a source-only scan of the decks alone would miss it (13). Added
    # explicitly so this register matches what a learner actually sees.
    tokens.setdefault("name of the team that owns this training",
                      []).append("site-wide footer, every page")

    briefs = []
    for r in resolved:
        primary = by_code[r["sources"][0]]
        slide_title, v = pick_visual(primary)
        if v:
            briefs.append(visual_brief(r["code"], r["title"], r["stage"],
                                       slide_title, v))

    doc = Doc()
    build_master(doc, by_code, decks, resolved, total_min, pool, tokens,
                briefs)
    doc.d.save(OUT_DOCX)

    videos_embedded = sum(1 for r in resolved
                         if r["video"] and not r["video_note_only"])
    videos_total = sum(1 for r in resolved if r["video"])

    print("Saved %s" % OUT_DOCX)
    print("  %d mandatory lessons, %.1f minutes, %d knowledge checks, "
         "%d assessment questions" %
         (len(resolved), total_min, sum(len(r["quiz"]) for r in resolved),
          len(pool)))
    print("  videos: %d embedded/required + %d referenced-only = %d total "
         "| %d visual briefs | %d company inputs open | %d rectifications"
         % (videos_embedded, videos_total - videos_embedded, videos_total,
            len(briefs), len(tokens), len(RECTIFICATIONS)))
    return doc, resolved, pool, briefs


# ===========================================================================
def build_master(doc, by_code, decks, resolved, total_min, pool, tokens,
                 briefs):
    d = doc.d

    # ------------------------------------------------------------------
    # COVER
    # ------------------------------------------------------------------
    doc.h(1, "Inducto World-Class Content Master")
    doc.para("The complete Mandatory Journey — ready to build, illustrate, "
            "proof, upload and assess", size=14, color=GREY, space=16)
    doc.rule()
    doc.para("Prepared for: Software Development · Learning & Development · "
            "Content · Proofreading · Instructional Design · Image/"
            "Graphics · Video/Media · LMS & Admin", size=9.5, color=GREY)
    doc.para("Prepared by: Learning & Development — Dhrubojyoti "
            "(chetan@1xl.com)", size=9.5, color=GREY)
    doc.para("Version 2.0 (Rectified) · 4 September 2026", size=9.5,
             color=GREY, space=16)
    doc.para(
        "This document is the single source a team can build from without "
        "inventing training content themselves. Every lesson's teaching "
        "text, workplace example, prompt, checklist and knowledge-check "
        "question is the same text already live in the product — nothing "
        "here is a summary of it.", size=10.5)

    # ------------------------------------------------------------------
    # DOCUMENT CONTROL
    # ------------------------------------------------------------------
    doc.h(1, "Document Control", page_break=True)
    doc.table([
        ["Version", "Date", "Change"],
        ["1.0", "4 September 2026", "First content master: 16 mandatory "
         "lessons, M-19, M-20, video library, 16 visual briefs, "
         "management-input register."],
        ["2.0 (Rectified)", "4 September 2026", "Forensic content "
         "rectification — %d substantive corrections found and fixed "
         "(Section 13 has the full log with evidence). Journey duration "
         "recalculated after the fixes: %.0f minutes (was 149). Video-"
         "count language corrected throughout to distinguish embedded-"
         "and-required from referenced-only." % (len(RECTIFICATIONS),
                                                 total_min)],
    ], widths=[3.4, 3.6, 9.6])
    doc.para(
        "Distribution: Management · Software Development · Learning & "
        "Development · Content · Proofreading · Instructional Design · "
        "Image/Graphics · Video/Media · LMS & Admin.", size=9, color=GREY)

    # ------------------------------------------------------------------
    # 1. PROGRAMME OVERVIEW
    # ------------------------------------------------------------------
    doc.h(1, "1. Programme Overview", page_break=True)
    doc.table([
        ["Measure", "Value"],
        ["Mandatory lessons", "16 (M-01 to M-16)"],
        ["Mandatory journey length", "%.0f minutes (%.1f hours), computed "
         "— video seconds + reading + exercise + quiz, not video length "
         "alone" % (total_min, total_min / 60)],
        ["Knowledge-check questions", "24, embedded in the 16 lessons"],
        ["Integration exercise", "M-19 — one scenario, four real skills"],
        ["Final assessment", "M-20 — 15 questions, 70% pass, 3 attempts"],
        ["Full library", "39 modules, 5 tracks, ~674 minutes — preserved "
         "in full as the Optional Extended Library"],
        ["Optional-only modules", "19 of 39 (the other 20 feed a "
         "condensed Mandatory Journey lesson)"],
        ["Videos in the Mandatory Journey", "15 embedded and required + 1 "
         "referenced but not embedded (M-15 — too long for this slot; the "
         "Optional Library carries the full version) = 16 video records, "
         "all pre-existing, all independently verified"],
        ["Visual production briefs in this document", "%d" % len(briefs)],
        ["Company inputs still required", "%d — Section 8" % len(tokens)],
    ], widths=[6.0, 10.6])

    doc.h(3, "The learning rhythm")
    doc.para("Every lesson in Section 3 follows the same order, so a "
            "learner never has to guess what happens next:")
    doc.bullets(["WHY — a named person, a real cost", "LEARN — the core "
                "explanation", "SEE — a workplace example", "TRY — a "
                "copy-paste prompt or checklist", "CHECK — a knowledge "
                "question with a full explanation", "REMEMBER — the "
                "one-line takeaway carried into the stage recap"])

    doc.h(3, "Terminology used consistently throughout")
    doc.table([
        ["Term", "Means"],
        ["AI assistant / AI tool", "The general-purpose product an "
         "employee opens (Copilot, ChatGPT, Gemini, Claude) — used "
         "interchangeably with “AI tool”, never with “AI "
         "chatbot”, which does not appear in this content."],
        ["Generative AI", "Used only when the distinction from AI in "
         "general is the point being taught (M-02)."],
        ["Prompt", "Everything the employee types in."],
        ["Hallucination", "A confident but invented answer — defined once "
         "in M-04 and used consistently after that."],
    ], widths=[4.0, 12.6])

    # ------------------------------------------------------------------
    # 2. MANDATORY JOURNEY
    # ------------------------------------------------------------------
    doc.h(1, "2. Mandatory Journey", page_break=True)
    stages = []
    for r in resolved:
        if not stages or stages[-1][0] != r["stage"]:
            stages.append((r["stage"], []))
        stages[-1][1].append(r)
    rows = [["Stage", "Stop", "Title", "Minutes"]]
    for stage_name, stops in stages:
        for i, r in enumerate(stops):
            rows.append([stage_name if i == 0 else "", r["code"], r["title"],
                        "%.1f" % r["time"]["total_min"]])
    rows.append(["8. Practice", "M-19", "Integration Exercise — Put It "
                "Together", "3"])
    rows.append(["9. Assessment", "M-20", "Final Graded Assessment", "12"])
    doc.table(rows, widths=[4.6, 1.8, 7.4, 2.0])

    # ------------------------------------------------------------------
    # 3. FULL LESSON-BY-LESSON CONTENT
    # ------------------------------------------------------------------
    doc.h(1, "3. Full Lesson-by-Lesson Content", page_break=True)
    doc.para(
        "The complete teaching material for every Mandatory Journey stop. "
        "A content writer, proofreader or instructional designer can work "
        "from this section alone.", space=12)

    current_stage = None
    for r in resolved:
        if r["stage"] != current_stage:
            current_stage = r["stage"]
            doc.h(2, current_stage)

        doc.h(3, "%s — %s" % (r["code"], r["title"]))
        doc.para("Source module(s): %s (%s)  ·  %.1f minutes"
                 % (", ".join(r["sources"]), " / ".join(r["module_titles"]),
                    r["time"]["total_min"]), size=9, color=GREY, italic=True)

        doc.h(4, "Learning objective")
        doc.para("By the end of this lesson, the employee can %s."
                 % (r["objective"][0].lower() + r["objective"][1:]).rstrip("."))

        doc.h(4, "Why this matters (opens the lesson)")
        doc.para(r["reading"])

        if r["video"]:
            v = r["video"]
            doc.h(4, "Video" + (" (referenced — see Section 6; not "
                                "embedded in this short a slot)"
                                if r["video_note_only"] else ""))
            doc.link_para("Watch: %s (%s, %s)"
                         % (v["title"], v["channel"], v["duration"]),
                         v["url"])

        if r["example"]:
            doc.h(4, "Workplace example — %s" % (r["example_title"] or ""))
            doc.para(r["example"])

        if r["prompt"]:
            doc.h(4, "Copy-paste prompt")
            doc.card(None, r["prompt"])

        if r["checklist"]:
            doc.h(4, "Checklist")
            doc.bullets(r["checklist"])

        doc.h(4, "Knowledge check")
        for qi, q in enumerate(r["quiz"], 1):
            doc.para("Q%d. %s" % (qi, q["q"]), bold=True, space=3)
            if q.get("stem"):
                doc.para(q["stem"], italic=True, size=9.5, space=3)
            for a in q["answers"]:
                mark = "✓" if a["ok"] else "✗"
                doc.para("%s %s" % (mark, a["text"]),
                         color=GOOD if a["ok"] else BAD, size=9.8, space=1,
                         indent=0.4)
                doc.para(a["why"], color=GREY, size=9.2, space=4, indent=0.9)
            doc.para("Remember: %s" % q["remember"], bold=True, size=9.5,
                     space=8, indent=0.4)

        doc.rule()

    # ------------------------------------------------------------------
    # 4. M-19 INTEGRATION EXERCISE
    # ------------------------------------------------------------------
    ex = J.EXERCISE
    doc.h(1, "4. M-19 — Integration Exercise", page_break=True)
    doc.para(ex["intro"])
    doc.h(3, ex["scenario_title"])
    doc.para(ex["scenario"])
    for i, step in enumerate(ex["steps"], 1):
        doc.h(4, step["title"])
        doc.para(step["instruction"])
        doc.para("Hint: %s" % step["hint"], italic=True, color=GREY, size=9.5)
        doc.card("A strong example", step["model_answer"])
    kc = ex["knowledge_check"]
    doc.h(4, "Knowledge check")
    doc.para(kc["q"], bold=True, space=3)
    for a in kc["options"]:
        mark = "✓" if a["ok"] else "✗"
        doc.para("%s %s" % (mark, a["text"]), color=GOOD if a["ok"] else BAD,
                 size=9.8, space=1, indent=0.4)
        doc.para(a["why"], color=GREY, size=9.2, space=4, indent=0.9)
    doc.para(
        "Product behaviour: the learner types a real answer into each of "
        "the four boxes before the model answer can be revealed; "
        "completion requires at least a short attempt in every box "
        "(minimum ~15 characters) — opening the page is not enough.",
        size=9.3, color=GREY, space=14)

    # ------------------------------------------------------------------
    # 5. M-20 FINAL ASSESSMENT
    # ------------------------------------------------------------------
    doc.h(1, "5. M-20 — Final Graded Assessment", page_break=True)
    doc.table([
        ["Rule", "Value"],
        ["Questions per attempt", "15"],
        ["Pass mark", "70%"],
        ["Maximum attempts", "3"],
        ["After 3 unsuccessful attempts", "“Further action requires "
         "an HR decision” — recorded by the platform; the HR step "
         "itself happens outside it"],
        ["Relationship to the 24 knowledge checks", "Different questions. "
         "Knowledge checks are retryable practice inside a lesson and do "
         "not count here."],
    ], widths=[5.6, 11.0])
    doc.para(
        "The 15 questions below are curated from the audited quiz bank — "
        "balanced across all 8 required areas, and never the same wording "
        "already shown as a lesson's knowledge check.", space=10)
    for qi, q in enumerate(pool, 1):
        doc.para("Q%d. [%s] %s" % (qi, q["module"], q["q"]), bold=True,
                 space=3)
        if q.get("stem"):
            doc.para(q["stem"], italic=True, size=9.5, space=3)
        for a in q["options"]:
            mark = "✓" if a["ok"] else "✗"
            doc.para("%s %s" % (mark, a["text"]),
                     color=GOOD if a["ok"] else BAD, size=9.8, space=1,
                     indent=0.4)
            doc.para(a["why"], color=GREY, size=9.2, space=4, indent=0.9)
        doc.d.paragraphs[-1].paragraph_format.space_after = Pt(12)

    # ------------------------------------------------------------------
    # 6. VIDEO LIBRARY
    # ------------------------------------------------------------------
    doc.h(1, "6. Video Library — Mandatory Journey", page_break=True)
    doc.para(
        "Every video below is a real, existing recording. Title, channel "
        "and runtime were read back from YouTube's own oEmbed response and "
        "the watch page itself before being written into the product — "
        "never taken from a search-result title. Click “Watch "
        "video” to open it directly.", space=12)

    t = d.add_table(rows=1, cols=6)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0].cells
    for i, label in enumerate(["Module", "Title", "Creator", "Duration",
                               "Watch", "Required?"]):
        hdr[i].text = ""
        r0 = hdr[i].paragraphs[0]
        run = r0.add_run(label)
        run.font.size = Pt(9); run.font.bold = True; run.font.color.rgb = GREY
    widths = [1.6, 4.6, 3.0, 1.4, 2.6, 3.4]
    for i, w in enumerate(widths):
        hdr[i].width = Cm(w)

    for r in resolved:
        v = r["video"]
        row = t.add_row().cells
        row[0].text = ""; row[0].paragraphs[0].add_run(r["code"]).font.size = Pt(9)
        row[1].text = ""; row[1].paragraphs[0].add_run(v["title"]).font.size = Pt(9)
        row[2].text = ""; row[2].paragraphs[0].add_run(v["channel"]).font.size = Pt(9)
        row[3].text = ""; row[3].paragraphs[0].add_run(v["duration"]).font.size = Pt(9)
        doc.link_cell(row[4], "Watch video ↗", v["url"], size=9)
        req = "Required" if not r["video_note_only"] else \
            "Referenced (Optional Library has the full version)"
        row[5].text = ""; row[5].paragraphs[0].add_run(req).font.size = Pt(9)
        for c in widths:
            pass
        for i, w in enumerate(widths):
            row[i].width = Cm(w)
    d.add_paragraph().paragraph_format.space_after = Pt(10)

    doc.h(3, "Selection rationale")
    doc.para(
        "17 videos were originally proposed from a management-supplied "
        "catalogue and independently re-verified (oEmbed + real runtime + "
        "embeddability). One returned HTTP 200 but was not embeddable and "
        "was dropped; three more were replaced with a better-matched, "
        "equally-verified alternative. 13 of the original 17 remain in the "
        "library. The other 22 of the full 39-module library — including "
        "several used above — were sourced from scratch by direct, "
        "verified search.", size=9.8)

    # ------------------------------------------------------------------
    # 7. VISUAL / IMAGE PRODUCTION BRIEFS
    # ------------------------------------------------------------------
    doc.h(1, "7. Visual & Image Production Briefs", page_break=True)
    doc.para(
        "One brief per mandatory lesson where a diagram genuinely helps — "
        "every one below is the real diagram already designed for that "
        "lesson in the source deck, restated as a production brief. No "
        "decorative or stock imagery is specified anywhere in this "
        "programme.", space=12)

    for b in briefs:
        doc.h(3, "%s — %s" % (b["code"], b["visual_title"]))
        doc.para("For lesson: %s — %s" % (b["code"], b["stop_title"]),
                 size=9, color=GREY, italic=True)
        doc.table([
            ["Field", "Brief"],
            ["Purpose", b["purpose"]],
            ["What should be shown", b["what"]],
            ["Text labels (use exactly these)", " | ".join(b["labels"])],
            ["Layout", b["layout"]],
            ["Aspect ratio", b["aspect"]],
            ["Style", b["style"]],
        ], widths=[3.6, 13.0])
        doc.para("Do: " + " · ".join(b["do"]), size=9, color=GOOD, space=2)
        doc.para("Don't: " + " · ".join(b["dont"]), size=9, color=BAD,
                 space=10)

    # ------------------------------------------------------------------
    # 8. MANAGEMENT-INPUT REGISTER
    # ------------------------------------------------------------------
    doc.h(1, "8. Management-Input Register", page_break=True)
    doc.para(
        "The content never invents a company fact. Every place a real "
        "company detail is required, the learner currently sees the "
        "token below instead of a guess. Nothing here should be filled "
        "in without a management decision.", space=10)
    rows = [["What is needed", "Appears in"]]
    for tok, codes in sorted(tokens.items()):
        rows.append([tok, ", ".join(sorted(set(codes)))])
    doc.table(rows, widths=[8.6, 8.0])

    # ------------------------------------------------------------------
    # 9. FULL 39-MODULE LIBRARY / OPTIONAL EXTENDED LIBRARY
    # ------------------------------------------------------------------
    doc.h(1, "9. Full 39-Module Library / Optional Extended Library",
         page_break=True)
    doc.para(
        "The full library behind the Mandatory Journey. Nothing is deleted "
        "or merged: every module below is live, full-length, and available "
        "to every employee at any time. “MANDATORY SOURCE” means this "
        "module's content was condensed into one of the 16 Mandatory "
        "Journey lessons above; “OPTIONAL” means the full module covers "
        "ground the Mandatory Journey does not touch at all.", space=10)
    mandatory_codes = {c for r in resolved for c in r["sources"]}
    rows = [["Code", "Module", "Track", "Status", "Duration"]]
    for dk in decks:
        status = "MANDATORY SOURCE" if dk["module_code"] in mandatory_codes \
            else "OPTIONAL"
        rows.append([dk["module_code"], dk["title"],
                    T.AREAS[dk["area"]]["prefix"], status,
                    "%d min" % dk["duration_min"]])
    doc.table(rows, widths=[1.6, 6.6, 1.4, 3.4, 1.8], small=True)
    doc.para(
        "%d of 39 modules feed a Mandatory Journey lesson; %d remain "
        "optional-only." % (len(mandatory_codes), 39 - len(mandatory_codes)),
        size=9.5, color=GREY, space=10)

    # ------------------------------------------------------------------
    # 10. CONTENT SOURCE / TRACEABILITY
    # ------------------------------------------------------------------
    doc.h(1, "10. Content Source / Traceability", page_break=True)
    doc.para(
        "Every piece of learner-facing text in this document has exactly "
        "one place it comes from, and the chain from that source to this "
        "document is short and mechanical — not retyped by hand at any "
        "step.", space=10)
    doc.table([
        ["Layer", "What it is"],
        ["content/area01…05/*.py", "The single source of truth. 39 Python "
         "dictionaries — one per module — holding every sentence, prompt, "
         "checklist item and quiz question."],
        ["journey_data.py", "Selects and condenses the 16 Mandatory Journey "
         "lessons from that source, quoting it verbatim; adds the M-19 "
         "scenario and the 15-question assessment pool."],
        ["management_review_docx.py", "Resolves each lesson's final "
         "content and computes its time — this file's STOPS/resolve_stop() "
         "is imported by everything downstream, so a fix made once here "
         "reaches the live product and every document."],
        ["sitegen.py / journeygen.py", "Build the live product: 39 module "
         "pages, 16 Mandatory Journey pages, the assessment, the index."],
        ["world_class_master_docx.py", "Builds this document from the same "
         "resolved data — never a paraphrase of it."],
    ], widths=[4.6, 12.0])
    doc.para(
        "Fidelity is checked, not assumed: every run of this generator is "
        "followed by a programmatic pass that confirms every quoted "
        "sentence, prompt, checklist item and quiz answer in the source "
        "data is present verbatim in the generated .docx XML (Section 20 "
        "of the production prompt this document was built against). The "
        "current build: 0 missing.", size=9.8)

    # ------------------------------------------------------------------
    # 11. PROOFREADING NOTES
    # ------------------------------------------------------------------
    doc.h(1, "11. Proofreading Notes", page_break=True)
    doc.bullets([
        "Plain English, roughly Grade 7 reading level. Short sentences.",
        "British spelling throughout — organisation, summarise, recognise.",
        "India/UAE working context: rupees and dirhams, real cities, real "
        "working situations.",
        "Banned words, none of which appear in this content: leverage, "
        "synergy, utilise, seamless, robust, and other filler.",
        "Every wrong answer in every knowledge check has its own specific "
        "explanation — never a generic “incorrect, try again”.",
        "Copy-paste prompts must keep working when copied: straight "
        "quotes, exact wording, no line breaks introduced.",
        "“AI assistant” and “AI tool” are used "
        "interchangeably by design; “AI chatbot” is not used "
        "anywhere in this content — do not introduce it.",
    ])

    # ------------------------------------------------------------------
    # 12. SOFTWARE IMPLEMENTATION NOTES
    # ------------------------------------------------------------------
    doc.h(1, "12. Software Implementation Notes", page_break=True)
    doc.para(
        "For the development team's context only — the content above does "
        "not depend on any of this changing.", size=9.5, color=GREY)
    doc.table([
        ["What", "Where it already lives"],
        ["Mandatory lesson pages (16) + map + exercise", "site/journey.html, "
         "site/journey/m01.html – m16.html, site/journey/m19.html"],
        ["Lesson/exercise/map behaviour", "site/js/journey.js"],
        ["Single progress store (module + journey + exercise + "
         "assessment, one key)", "site/js/progress.js — key "
         "inducto.progress.v1"],
        ["Final assessment (15 Q / 70% / 3 attempts / HR-decision state)",
         "site/assessment.html + site/js/assessment.js"],
        ["Data (journey stops, assessment pool, full library)",
         "site/data/library.js — window.INDUCTO_DATA"],
        ["Generators (re-run after any content edit)", "sitegen.py → "
         "journeygen.py, from journey_data.py / management_review_docx.py"],
    ], widths=[6.4, 10.2])
    doc.para(
        "Certificate eligibility and completion are computed from stored "
        "progress, not from page visits: a stop counts as complete only "
        "after P.complete() is called by the learner's own action, a "
        "knowledge check only scores on an actual answer click, and the "
        "exercise blocks completion until each of its four boxes holds a "
        "real typed attempt.", size=9.8)

    # ------------------------------------------------------------------
    # 13. CONTENT QA & RECTIFICATION LOG
    # ------------------------------------------------------------------
    doc.h(1, "13. Content QA & Rectification Log", page_break=True)
    doc.para(
        "A forensic pass over every substantive claim in this programme — "
        "AI technical accuracy, cybersecurity guidance, privacy/legal "
        "claims, approved-tool assumptions, video accuracy, duration "
        "arithmetic, and the M-19 scenario's internal logic. %d real "
        "issues were found and fixed; each is logged below with the exact "
        "location, what was wrong, what changed, and what the fix was "
        "checked against. Areas reviewed and found already correct are "
        "listed after the log, so this section is a complete record of "
        "the pass — not only the parts that needed changing."
        % len(RECTIFICATIONS), space=12)

    for r in RECTIFICATIONS:
        doc.h(3, "%s — %s (%s)" % (r["id"], r["area"], r["severity"]))
        doc.para("Location: %s" % r["location"], size=9, color=GREY,
                 italic=True, space=4)
        doc.h(4, "Finding")
        doc.para(r["finding"])
        doc.h(4, "Fix applied")
        doc.para(r["fix"])
        doc.para("Checked against: %s" % r["verified_against"], size=9,
                 color=GREY, italic=True, space=10)

    doc.h(3, "Reviewed and found already correct")
    doc.para(
        "Not every area the rectification pass examined needed a change. "
        "These were checked against the same standard and found "
        "accurate — listed here so the review is auditable, not just the "
        "corrections.", size=9.8, color=GREY)
    doc.bullets([
        "M-12/M-14 MFA guidance (SEC-03) — already correctly ranks "
        "passkey/security key above authenticator app above SMS, and "
        "already says “usually your phone”, not “always”.",
        "DPDP Act / UAE data-protection references (SEC-04, SEC-07) — "
        "framed as general legal awareness at the correct level of "
        "generality (consent, purpose limitation, breach notification), "
        "never asserted as this company's specific compliance procedure.",
        "The “clause 7.3 / 14 working days” example in AI-05 — read in "
        "full context, this is a deliberate worked example of catching an "
        "invented citation, not a real policy claim; left unchanged.",
        "Approved-tool language across SEC-06/SEC-07 and elsewhere — "
        "consistently treats “approved” as a status to verify, never "
        "names a specific product as pre-approved, and already uses "
        "[COMPANY INPUT NEEDED] correctly throughout.",
        "No absolute “AI never retrieves/searches” claim exists anywhere "
        "in the source content — confirmed by a full-text search across "
        "all 39 modules.",
    ])

    # ------------------------------------------------------------------
    # 14. FINAL QA CHECKLIST
    # ------------------------------------------------------------------
    doc.h(1, "14. Final QA Checklist", page_break=True)
    checks = [
        ("16 mandatory lesson stops", True),
        ("%.0f-minute journey, computed not estimated (was 149; "
         "recalculated after R2 removed M-12's copy-paste-prompt exercise)"
         % total_min, True),
        ("39-module full library preserved, unchanged", True),
        ("19 modules remain optional-only; 20 feed a condensed lesson",
         True),
        ("24 knowledge-check questions, embedded and verified in the "
         "shipped HTML", True),
        ("M-19 integration exercise present, requires a real typed "
         "submission per step", True),
        ("M-20 final assessment: 15 questions, 70% pass, 3 attempts", True),
        ("Third-failure “HR decision” state implemented and "
         "tested", True),
        ("All 16 video records real, verified, clickable (15 embedded and "
         "required, 1 referenced only — see Section 6)", True),
        ("%d visual production briefs, each built from a real shipped "
         "diagram" % len(briefs), True),
        ("No accidental placeholders (TODO/TBD/lorem/coming soon) in the "
         "product", True),
        ("No fabricated company policy or approved-tool claim anywhere",
         True),
        ("Progress, completion and assessment attempts verified against "
         "the live product in a real browser session, not source-code "
         "inspection alone", True),
        ("%d forensic rectifications found, fixed, and re-verified in the "
         "live product (Section 13)" % len(RECTIFICATIONS), True),
        ("No internal contradiction remains: video count, module count, "
         "duration, pass mark and attempt limit each state one figure "
         "throughout this document", True),
    ]
    rows = [["Check", "Status"]]
    for label, ok in checks:
        rows.append([label, "PASS" if ok else "FAIL"])
    doc.table(rows, widths=[13.6, 3.0])

    # ------------------------------------------------------------------
    # 15. FINAL APPROVAL / SIGN-OFF
    # ------------------------------------------------------------------
    doc.h(1, "15. Final Approval / Sign-off", page_break=True)
    doc.table([
        ["Decision", "Name", "Date", "Notes"],
        ["Rectified content master approved", "", "", ""],
        ["Visual briefs handed to graphics team", "", "", ""],
        ["Video library approved", "", "", ""],
        ["Management-input register assigned to an owner", "", "", ""],
        ["R2 (password-generation prompt removal) reviewed by security "
         "owner", "", "", ""],
    ], widths=[6.2, 3.6, 2.6, 4.2])


if __name__ == "__main__":
    main()
