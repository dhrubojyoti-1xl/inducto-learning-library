# -*- coding: utf-8 -*-
"""
Build the full content blueprint as a Word document, for proof reading.

    python blueprint_docx.py

Every sentence a learner can see — in the PowerPoint decks and on the web
platform — is written into this document once, in reading order, with a
stable reference code so a proof reader can cite an exact line.

The text is read from content/areaNN/*.py, the same dictionaries that build
both outputs. Nothing is retyped or paraphrased here.
"""

import datetime
import importlib
import io
import json
import os
import re
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

import build
import sitegen
import theme as T

ROOT = sitegen.ROOT
OUT_DOCX = os.path.join(ROOT, "docs",
                        "Inducto-Learning-Library-Content-Blueprint.docx")

INK = RGBColor(0x10, 0x18, 0x26)
GREY = RGBColor(0x5B, 0x67, 0x79)
ACCENT = RGBColor(0x2F, 0x4B, 0xC4)
ALERT = RGBColor(0xC6, 0x28, 0x28)

TRACK_ACCENT = {
    "01-ai-general": RGBColor(0x2F, 0x4B, 0xC4),
    "02-ai-daily-work": RGBColor(0x0E, 0x6E, 0x75),
    "03-prompt-engineering": RGBColor(0x6A, 0x2F, 0xA0),
    "04-professional-skills": RGBColor(0x8A, 0x5A, 0x00),
    "05-security-privacy": RGBColor(0x7A, 0x12, 0x20),
}


# ---------------------------------------------------------------------------
# document helpers
# ---------------------------------------------------------------------------
class Doc(object):

    def __init__(self):
        self.d = Document()
        self._page()
        self._styles()
        self.ref_n = 0
        self.code = ""
        self.refs = []          # (ref, text) for the coverage check

    def _page(self):
        s = self.d.sections[0]
        s.page_width, s.page_height = Cm(21.0), Cm(29.7)
        s.left_margin, s.right_margin = Cm(2.2), Cm(3.4)
        s.top_margin, s.bottom_margin = Cm(2.0), Cm(2.0)
        self._footer(s)

    def _footer(self, section):
        p = section.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("Inducto Learning & Knowledge Library · content blueprint · page ")
        r.font.size, r.font.color.rgb = Pt(8), GREY
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        p._p.append(fld)

    def _styles(self):
        st = self.d.styles
        n = st["Normal"]
        n.font.name = "Segoe UI"
        n.font.size = Pt(10.5)
        n.font.color.rgb = INK
        n.paragraph_format.space_after = Pt(6)
        n.paragraph_format.line_spacing = 1.15
        for name, size, col, before in (("Heading 1", 20, INK, 22),
                                        ("Heading 2", 14, INK, 16),
                                        ("Heading 3", 11.5, INK, 12),
                                        ("Heading 4", 10.5, GREY, 10)):
            s = st[name]
            s.font.name = "Segoe UI Semibold"
            s.font.size = Pt(size)
            s.font.color.rgb = col
            s.font.bold = True
            s.paragraph_format.space_before = Pt(before)
            s.paragraph_format.space_after = Pt(4)
            s.paragraph_format.keep_with_next = True

    # -- structure ---------------------------------------------------------
    def h(self, level, text, color=None, page_break=False):
        if page_break:
            self.d.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        p = self.d.add_paragraph(style="Heading %d" % level)
        r = p.add_run(text)
        if color is not None:
            r.font.color.rgb = color
        return p

    def para(self, text, size=10.5, color=None, italic=False, bold=False,
             space=6, indent=0.0):
        p = self.d.add_paragraph()
        p.paragraph_format.space_after = Pt(space)
        if indent:
            p.paragraph_format.left_indent = Cm(indent)
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.italic = italic
        r.font.bold = bold
        r.font.color.rgb = color if color is not None else INK
        return p

    def rule(self):
        p = self.d.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(8)
        pbdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "6")
        bot.set(qn("w:color"), "D6DCE7")
        pbdr.append(bot)
        p._p.get_or_add_pPr().append(pbdr)

    def bullets(self, items, style="List Bullet"):
        for it in items:
            p = self.d.add_paragraph(style=style)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(it)
            r.font.size = Pt(10.5)

    def table(self, rows, widths=None, header=True):
        t = self.d.add_table(rows=0, cols=len(rows[0]))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i, row in enumerate(rows):
            cells = t.add_row().cells
            for j, val in enumerate(row):
                cells[j].text = ""
                p = cells[j].paragraphs[0]
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(str(val))
                r.font.size = Pt(9)
                r.font.bold = header and i == 0
                r.font.color.rgb = GREY if (header and i == 0) else INK
                if widths:
                    cells[j].width = Cm(widths[j])
        self.d.add_paragraph().paragraph_format.space_after = Pt(4)
        return t

    # -- the referenced content block --------------------------------------
    def block(self, label, text, tone=None, indent=0.0):
        """One proof-readable line: reference code, label, then the text."""
        self.ref_n += 1
        ref = "%s.%03d" % (self.code, self.ref_n)
        self.refs.append((ref, text))

        p = self.d.add_paragraph()
        p.paragraph_format.space_after = Pt(7)
        p.paragraph_format.left_indent = Cm(indent)
        p.paragraph_format.keep_together = True

        r = p.add_run(ref + "  ")
        r.font.size, r.font.bold = Pt(7.5), True
        r.font.color.rgb = RGBColor(0x9A, 0xA6, 0xB8)

        r = p.add_run(label.upper() + "   ")
        r.font.size, r.font.bold = Pt(7.5), True
        r.font.color.rgb = {"bad": ALERT, "good": RGBColor(0x1B, 0x7F, 0x4B)}.get(
            tone, ACCENT)

        r = p.add_run(text)
        r.font.size = Pt(10.5)
        return ref

    def toc(self, instr):
        """A real Word TOC field. Word fills it in on first update."""
        p = self.d.add_paragraph()
        r = p.add_run()._r
        for typ, extra in (("begin", None), (None, instr), ("separate", None),
                           (None, "Right-click and choose Update Field."),
                           ("end", None)):
            if typ:
                f = OxmlElement("w:fldChar")
                f.set(qn("w:fldCharType"), typ)
                r.append(f)
            else:
                t = OxmlElement("w:instrText" if extra is instr else "w:t")
                t.set(qn("xml:space"), "preserve")
                t.text = extra
                r.append(t)

    def new_module(self, code):
        self.code = code
        self.ref_n = 0


# ---------------------------------------------------------------------------
# visual → labelled blocks
# ---------------------------------------------------------------------------
def emit_visual(doc, v):
    t = v["type"]

    if t == "flow":
        for i, (title, sub) in enumerate(v["steps"], 1):
            doc.block("Step %d · heading" % i, title)
            doc.block("Step %d · text" % i, sub, indent=0.6)

    elif t == "nested":
        for i, layer in enumerate(v["layers"], 1):
            doc.block("Layer %d · heading" % i, layer["label"])
            doc.block("Layer %d · text" % i, layer["sub"], indent=0.6)
        if v.get("note"):
            doc.block("Note", v["note"])

    elif t == "iconrow":
        for i, it in enumerate(v["items"], 1):
            doc.block("Tile %d · heading" % i, it["label"])
            doc.block("Tile %d · text" % i, it["sub"], indent=0.6)

    elif t == "split":
        for side, letter in (("left", "A"), ("right", "B")):
            s = v[side]
            doc.block("Panel %s · tag" % letter, s["tag"], tone=s.get("tone"))
            doc.block("Panel %s · heading" % letter, s["title"], tone=s.get("tone"))
            for i, it in enumerate(s["items"], 1):
                doc.block("Panel %s · point %d" % (letter, i), it, indent=0.6)

    elif t == "tree":
        doc.block("Decision question", v["question"])
        for key in ("yes", "no"):
            b = v[key]
            doc.block("Branch · path", b["path"], tone=b.get("tone"), indent=0.6)
            doc.block("Branch · heading", b["label"], tone=b.get("tone"), indent=0.6)
            doc.block("Branch · text", b["detail"], indent=1.2)

    elif t == "steps":
        for i, it in enumerate(v["items"], 1):
            doc.block("Instruction %d" % i, it)
        if v.get("prompt"):
            doc.block("Copy-paste prompt", v["prompt"])
        if v.get("caption"):
            doc.block("Prompt caption", v["caption"], indent=0.6)

    elif t == "beforeafter":
        doc.block("Before · tag", v["bad_tag"], tone="bad")
        for i, it in enumerate(v["bad"], 1):
            doc.block("Before · point %d" % i, it, tone="bad", indent=0.6)
        doc.block("After · tag", v["good_tag"], tone="good")
        for i, it in enumerate(v["good"], 1):
            doc.block("After · point %d" % i, it, tone="good", indent=0.6)
        if v.get("note"):
            doc.block("Closing note", v["note"])

    elif t == "prompt":
        if v.get("header"):
            doc.block("Prompt card heading", v["header"])
        doc.block("Copy-paste prompt", v["text"])
        if v.get("caption"):
            doc.block("Prompt caption", v["caption"], indent=0.6)
        for i, wy in enumerate(v.get("why") or [], 1):
            doc.block("Why it works %d" % i, wy, indent=0.6)

    elif t == "prompt_out":
        if v.get("header"):
            doc.block("Prompt card heading", v["header"])
        doc.block("Copy-paste prompt", v["text"])
        if v.get("caption"):
            doc.block("Prompt caption", v["caption"], indent=0.6)
        doc.block("Output panel heading", v.get("out_title") or "What comes back")
        for i, o in enumerate(v["out"], 1):
            doc.block("Output point %d" % i, o, indent=0.6)

    elif t == "checklist":
        tone = "bad" if v.get("mark") == "ban" else None
        for i, it in enumerate(v["items"], 1):
            doc.block("Checklist %d" % i, it, tone=tone)

    elif t == "bandlist":
        doc.block("Band headline", v["headline"], tone=v.get("tone") or "bad")
        if v.get("sub"):
            doc.block("Band text", v["sub"], indent=0.6)
        for i, it in enumerate(v["items"], 1):
            doc.block("Point %d" % i, it, indent=0.6)

    elif t == "mistakes":
        for i, (what, why) in enumerate(v["items"], 1):
            doc.block("Mistake %d · what" % i, what, tone="bad")
            doc.block("Mistake %d · consequence" % i, why, indent=0.6)

    else:
        raise SystemExit("unhandled visual type: %s" % t)


# ---------------------------------------------------------------------------
# one module
# ---------------------------------------------------------------------------
def emit_module(doc, d, n, total):
    code = d["module_code"]
    accent = TRACK_ACCENT[d["area"]]
    doc.new_module(code)

    doc.h(1, "%s — %s" % (code, d["title"]), color=accent, page_break=True)
    doc.para("Module %d of %d · %s · %d minutes · %s"
             % (n, total, T.AREAS[d["area"]]["name"], d["duration_min"],
                d["audience"]), size=9, color=GREY)
    doc.table([
        ["Where it appears", "Reference"],
        ["PowerPoint deck", "output/%s" % d["filename"]],
        ["Web page", "site/modules/%s.html" % sitegen.slug(code)],
        ["Live page", "https://dhrubojyoti-1xl.github.io/inducto-learning-"
                      "library/modules/%s.html" % sitegen.slug(code)],
        ["Source of this text", "content/…/%s" % module_file(code)],
    ], widths=[4.6, 10.8])

    # ---- cover ----
    doc.h(2, "1. Cover")
    doc.block("Module title", d["title"])
    doc.block("Subtitle", d["subtitle"])

    # ---- why ----
    doc.h(2, "2. Why this matters")
    w = d["why"]
    doc.block("Heading", w["title"])
    doc.block("Story", w["scenario"])
    doc.block("The cost", w["cost"], tone="bad")
    doc.block("What changes", w["fix"], tone="good")

    # ---- outcomes ----
    doc.h(2, "3. What you'll be able to do")
    for i, (icon, text) in enumerate(d["outcomes"], 1):
        doc.block("Outcome %d" % i, text)

    # ---- content sections ----
    groups = sitegen.group_slides(d)
    num = 4
    for name, sub, anchor in d["sections"]:
        if anchor == "scenario":
            doc.h(2, "%d. %s" % (num, name))
            sc = d["scenario"]
            doc.block("Section subtitle", sub)
            doc.block("Scenario heading", sc["title"])
            doc.block("The situation", sc["situation"])
            for i, c in enumerate(sc["choices"]):
                letter = "ABCDE"[i]
                tone = {"good": "good", "bad": "bad"}.get(c["tone"])
                doc.h(3, "Choice %s (%s)" % (letter, c["tone"]))
                doc.block("Choice %s · text" % letter, c["text"], tone=tone)
                doc.block("Choice %s · feedback heading" % letter, c["headline"],
                          tone=tone, indent=0.6)
                doc.block("Choice %s · what happens" % letter, c["consequence"],
                          indent=0.6)
                doc.block("Choice %s · the rule" % letter, c["rule"], indent=0.6)
            num += 1
            continue

        if anchor == "video":
            doc.h(2, "%d. %s" % (num, name))
            v = d["video"]
            doc.block("Section subtitle", sub)
            doc.block("Section heading", v["heading"])
            doc.para("Video: “%s” · %s · %s · %s"
                     % (v["title"], v["channel"], v["duration"], v["url"]),
                     size=9, color=GREY)
            doc.block("Note to the learner", v["note"])
            for i, h in enumerate(v["how"], 1):
                doc.block("How to use it %d" % i, h, indent=0.6)
            num += 1
            continue

        doc.h(2, "%d. %s" % (num, name))
        doc.block("Section subtitle", sub)
        for s in groups.get(anchor, []):
            doc.h(3, s["title"])
            if s.get("lead"):
                doc.block("Lead", s["lead"])
            if s.get("visual"):
                emit_visual(doc, s["visual"])
            if s.get("gloss"):
                doc.para("Glossary terms used here: %s" % ", ".join(s["gloss"]),
                         size=8.5, color=GREY)
        num += 1

    # ---- knowledge check ----
    doc.h(2, "%d. Knowledge check" % num)
    num += 1
    for qi, q in enumerate(d["quiz"], 1):
        doc.h(3, "Question %d" % qi)
        doc.block("Q%d · question" % qi, q["q"])
        if q.get("stem"):
            doc.block("Q%d · set-up" % qi, q["stem"])
        for ai, a in enumerate(q["answers"]):
            letter = "ABCDE"[ai]
            tone = "good" if a["ok"] else "bad"
            mark = "correct" if a["ok"] else "incorrect"
            doc.block("Q%d %s · answer (%s)" % (qi, letter, mark), a["text"],
                      tone=tone)
            doc.block("Q%d %s · feedback" % (qi, letter), a["why"], indent=0.6)
        doc.block("Q%d · remember" % qi, q["remember"])

    # ---- recap ----
    doc.h(2, "%d. Recap" % num)
    num += 1
    r = d["recap"]
    doc.block("Recap heading", r["title"])
    for i, (t, dd) in enumerate(r["points"], 1):
        doc.block("Recap %d · heading" % i, t)
        doc.block("Recap %d · text" % i, dd, indent=0.6)
    doc.block("One-liner", r["oneliner"], tone="good")

    # ---- toolkit ----
    doc.h(2, "%d. Toolkit" % num)
    num += 1
    tk = d["toolkit"]
    doc.block("Toolkit heading", tk["title"])
    for i, (icon, t, dd) in enumerate(tk["templates"], 1):
        doc.block("Template %d · heading" % i, t)
        doc.block("Template %d · text" % i, dd, indent=0.6)
    doc.para("Tools referenced: %s"
             % "; ".join("%s — %s" % (a, b) for a, b in tk["links"]),
             size=9, color=GREY)
    doc.block("Next module", tk["next"])

    # ---- glossary ----
    doc.h(2, "%d. Glossary" % num)
    for term, definition in d["glossary"]:
        doc.block(term, definition)


def module_file(code):
    for mod in build.REGISTRY:
        if importlib.import_module(mod).DECK["module_code"] == code:
            return mod.split(".")[-1] + ".py"
    return "?"


# ---------------------------------------------------------------------------
# front matter
# ---------------------------------------------------------------------------
def front_matter(doc, decks, tokens):
    today = datetime.date.today().strftime("%d %B %Y")

    doc.h(1, "Inducto Learning & Knowledge Library")
    doc.para("Full content blueprint, for proof reading", size=14, color=GREY)
    doc.rule()
    doc.para("Version 1.0 · %s" % today, size=9, color=GREY)
    gloss_entries = sum(len(d["glossary"]) for d in decks)
    gloss_distinct = len({t.lower() for d in decks for t, _ in d["glossary"]})
    doc.para("%d modules · %d lessons · about %d hours of learning · "
             "%d glossary entries (%d distinct terms)"
             % (len(decks), sum(len(d["slides"]) for d in decks),
                round(sum(d["duration_min"] for d in decks) / 60.0),
                gloss_entries, gloss_distinct), size=10, color=INK)
    doc.para("")

    doc.h(2, "What this document is")
    doc.para(
        "Every sentence a learner can see, in reading order, exactly as it "
        "appears in the training. It is generated straight from the source "
        "content files — the same files that build the 40 PowerPoint decks "
        "and the 39 web pages — so what you read here is what learners read. "
        "Nothing has been retyped, shortened or paraphrased for this "
        "document.")
    doc.para(
        "That also means corrections have one place to go. A change made to "
        "the source content appears in the deck, on the web page and in the "
        "next version of this document at the same time.")

    doc.h(2, "How to mark it up")
    doc.para(
        "Every line of learner-facing text carries a reference code in grey "
        "at the start of the line — for example AI-01.014. Quote that code "
        "when you report a change and there is no ambiguity about which line "
        "you mean, even where the same sentence appears twice.")
    doc.para("The label after the code says where the line sits on the page:",
             space=3)
    doc.bullets([
        "LEAD — the opening sentence under a lesson heading",
        "STEP / INSTRUCTION — a numbered instruction the learner follows",
        "COPY-PASTE PROMPT — text the learner copies into an AI tool, so "
        "wording and punctuation are load-bearing",
        "PANEL A / PANEL B — the two sides of a comparison",
        "BRANCH — one arm of a decision tree",
        "CHOICE A · FEEDBACK — what the learner is told after picking that "
        "option in the scenario",
        "Q1 B · ANSWER (INCORRECT) — one wrong answer, followed by the "
        "explanation written for that specific wrong answer",
        "MISTAKE · WHAT / CONSEQUENCE — a common mistake and what it costs",
        "RECAP / ONE-LINER / NEXT MODULE — the closing summary lines",
    ])
    doc.para(
        "Please leave the reference codes in place when you comment. They are "
        "generated, not typed, and they are how a correction gets found again "
        "in the source.", space=10)

    doc.h(2, "House rules this content is written to")
    doc.para("Worth checking against, because a well-meant edit can break one:",
             space=3)
    doc.bullets([
        "Plain English at roughly a Grade 7 reading level. Short sentences. "
        "No jargon that is not defined on the same page.",
        "British spelling — organisation, summarise, recognise.",
        "Indian and UAE working context. Rupees and dirhams, Indian and UAE "
        "law, real places and real working situations.",
        "No invented company detail. Where a real company fact is needed the "
        "text carries a [COMPANY INPUT NEEDED: …] token instead of a guess. "
        "These are listed below and are deliberate — please do not fill them "
        "in while proof reading.",
        "No filler. Words such as leverage, synergy, utilise, seamless and "
        "robust are banned throughout and should not be introduced.",
        "Every wrong answer gets its own explanation. If you see feedback "
        "that would work for any wrong answer, that is a defect worth "
        "flagging.",
        "Copy-paste prompts must keep working when copied. Straight quotes, "
        "exact wording, no line breaks added.",
    ])

    doc.h(2, "The shape every module follows", page_break=True)
    doc.para(
        "All %d modules run the same spine, so a learner who has done one "
        "knows where to find things in the next." % len(decks), space=3)
    doc.table([
        ["#", "Part", "What it does"],
        ["1", "Cover", "Title and one-line promise"],
        ["2", "Why this matters", "One named person, the cost of the problem, "
                                  "what changes"],
        ["3", "What you'll be able to do", "Five outcomes, each a thing the "
                                           "learner can do afterwards"],
        ["4+", "Teaching sections", "Lessons: a heading, a lead, and one "
                                    "diagram, comparison, checklist, prompt "
                                    "or list of mistakes"],
        ["", "Choose what you'd do", "A branching workplace decision. Each "
                                     "choice has its own consequence and rule"],
        ["", "Watch this", "One outside video, under 15 minutes"],
        ["", "Knowledge check", "Five questions. Every answer, right or "
                               "wrong, has its own explanation"],
        ["", "Recap", "Six points and a one-line summary"],
        ["", "Toolkit", "Three things to take away, the tools referenced, "
                        "and the next module"],
        ["", "Glossary", "Every term used in the module, defined"],
    ], widths=[1.2, 4.4, 9.8])

    doc.h(2, "The five tracks")
    rows = [["Track", "Modules", "What it covers"]]
    for a in sitegen.AREA_ORDER:
        mods = [d for d in decks if d["area"] == a]
        rows.append([T.AREAS[a]["name"], str(len(mods)), sitegen.AREA_BLURB[a]])
    doc.table(rows, widths=[4.8, 1.8, 8.8])

    doc.h(2, "Module index")
    rows = [["Code", "Module", "Track", "Min", "Lessons", "Page"]]
    for d in decks:
        rows.append([d["module_code"], d["title"], T.AREAS[d["area"]]["prefix"],
                     str(d["duration_min"]), str(len(d["slides"])), ""])
    doc.table(rows, widths=[1.8, 6.6, 1.4, 1.2, 1.8, 2.6])
    doc.para("Page numbers are for this version of the document. The Contents "
             "list on the next page is clickable.", size=8.5, color=GREY)

    doc.h(2, "Contents", page_break=True)
    doc.para("Page numbers fill in when Word updates the field. If the list "
             "below is empty, click it and press F9.", size=8.5, color=GREY)
    doc.toc(r'TOC \o "1-2" \h \z \u')

    doc.h(2, "Company inputs still required", page_break=True)
    doc.para(
        "The content never invents a company fact. Where one is needed it "
        "carries a token, and the token is what learners currently see. "
        "There are %d distinct inputs outstanding. These are not typos — "
        "please leave them alone and let the training owner fill them in."
        % len(tokens), space=6)
    rows = [["What is needed", "Appears in"]]
    for tok, codes in sorted(tokens.items()):
        rows.append([tok, ", ".join(sorted(set(codes)))])
    doc.table(rows, widths=[8.6, 6.8])

    doc.h(2, "Video register")
    doc.para(
        "One outside video per module. Every title, channel and runtime below "
        "was read back from YouTube itself rather than copied from a list. "
        "These are third-party videos: where one disagrees with the module, "
        "the module wins, and each page says so.", space=6)
    rows = [["Module", "Video", "Channel", "Length"]]
    for d in decks:
        v = d["video"]
        rows.append([d["module_code"], v["title"], v["channel"], v["duration"]])
    doc.table(rows, widths=[1.8, 7.4, 4.4, 1.6])

    doc.h(2, "Sign-off")
    doc.para("One row per pass. Please note the version at the top of this "
             "document.", size=9, color=GREY)
    doc.table([
        ["Pass", "Reviewer", "Date", "Modules covered", "Outcome"],
        ["Language and spelling", "", "", "", ""],
        ["Factual accuracy", "", "", "", ""],
        ["Prompts tested by copying", "", "", "", ""],
        ["Tone and reading level", "", "", "", ""],
        ["Final approval", "", "", "", ""],
    ], widths=[4.0, 3.2, 2.2, 3.0, 3.0])


# ---------------------------------------------------------------------------
def main():
    decks = [importlib.import_module(m).DECK for m in build.REGISTRY]
    order = {a: i for i, a in enumerate(sitegen.AREA_ORDER)}
    decks.sort(key=lambda d: (order[d["area"]], d["module_code"]))

    tokens = {}
    for d in decks:
        for m in re.finditer(r"\[COMPANY INPUT NEEDED:([^\]]*)\]", json.dumps(d)):
            tokens.setdefault(m.group(1).strip(), []).append(d["module_code"])

    doc = Doc()
    front_matter(doc, decks, tokens)

    doc.h(1, "The content", page_break=True)
    doc.para("Every module in full, in the order a learner meets it.",
             size=11, color=GREY)

    for i, d in enumerate(decks, 1):
        emit_module(doc, d, i, len(decks))
        print("  %s  %-38s %4d referenced lines"
              % (d["module_code"], d["title"][:38], doc.ref_n))

    os.makedirs(os.path.dirname(OUT_DOCX), exist_ok=True)
    doc.d.save(OUT_DOCX)

    total_refs = len(doc.refs)
    words = sum(len(t.split()) for _, t in doc.refs)
    print("\nSaved %s" % OUT_DOCX)
    print("  %d referenced lines, %d words of learner-facing text"
          % (total_refs, words))
    print("  %.1f MB" % (os.path.getsize(OUT_DOCX) / 1024.0 / 1024.0))

    json.dump([{"ref": r, "text": t} for r, t in doc.refs],
              io.open(os.path.join(ROOT, "blueprint_refs.json"), "w",
                      encoding="utf-8"), ensure_ascii=False, indent=1)

    if "--no-word" not in sys.argv:
        finish_in_word([d["module_code"] for d in decks])


def finish_in_word(codes):
    """Let Word build the contents list and fill in the index page numbers.

    python-docx can write the TOC field but only Word can resolve it, so this
    opens the file once, updates the field, reads the page numbers back out of
    the finished contents list, writes them into the module index, and saves.
    """
    try:
        import win32com.client as win32
    except ImportError:
        print("\nWord not available - contents list will fill in when "
              "someone opens the file and presses F9.")
        return

    app = win32.Dispatch("Word.Application")
    app.Visible = False
    app.DisplayAlerts = 0
    d = app.Documents.Open(OUT_DOCX)
    try:
        d.Fields.Update()
        toc = d.TablesOfContents(1)
        toc.Update()
        d.Repaginate()

        pages = {}
        for line in toc.Range.Text.splitlines():
            m = re.match(r"\s*([A-Z]{2,3}-\d\d)\s+—.*?(\d+)\s*$", line)
            if m:
                pages[m.group(1)] = m.group(2)

        index = d.Tables(3)          # Code | Module | Track | Min | Lessons | Page
        filled = 0
        for i, code in enumerate(codes, start=2):
            if code in pages:
                index.Cell(i, 6).Range.Text = pages[code]
                filled += 1

        toc.Update()
        d.Repaginate()
        total = d.ComputeStatistics(2)
        words = d.ComputeStatistics(0)
        d.Save()
        print("\nWord pass: contents list built, %d of %d index page "
              "numbers filled." % (filled, len(codes)))
        print("  %d pages, %d words in the finished document." % (total, words))
    finally:
        d.Close(True)
        app.Quit()


if __name__ == "__main__":
    main()
